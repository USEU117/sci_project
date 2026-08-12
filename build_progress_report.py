from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUT = Path(r"D:\保研\SCI\少样本工业异常检测项目进度报告_2026-07-30.docx")


def set_run_font(run, name="Calibri", size=11, color="000000", bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            set_cell_width(cell, widths[i])
            cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_table(table, widths, header=True):
    set_table_geometry(table, widths)
    for r, row in enumerate(table.rows):
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.line_spacing = 1.05
                for run in p.runs:
                    set_run_font(run, size=9.2, bold=(header and r == 0))
            if header and r == 0:
                shade(cell, "E8EEF5")


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    style_table(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_para(doc, text="", bold_prefix=None, style=None, after=6, align=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    if align is not None:
        p.alignment = align
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5 + 0.2 * level)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.10
    r = p.add_run(text)
    set_run_font(r)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.10
    r = p.add_run(text)
    set_run_font(r)
    return p


def add_callout(doc, label, text, fill="F4F6F9"):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    shade(cell, fill)
    cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.10
    r = p.add_run(label + "：")
    set_run_font(r, bold=True, color="1F4D78")
    r = p.add_run(text)
    set_run_font(r)
    set_table_geometry(table, [9360])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def setup_styles(doc):
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        st = doc.styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.font.bold = True
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True
    header = sec.header.paragraphs[0]
    header.text = "少样本工业异常检测项目 · 进度报告"
    for run in header.runs:
        set_run_font(run, size=9, color="666666")
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("项目进度报告 | 2026-07-30")
    set_run_font(run, size=9, color="666666")


def build():
    doc = Document()
    setup_styles(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("少样本工业异常检测项目")
    set_run_font(r, size=25, color="0B2545", bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("从基准复现到文本—视觉证据融合的完整进度报告")
    set_run_font(r, size=14, color="555555")
    add_callout(
        doc,
        "项目主题",
        "本项目研究的是：在工业异常检测中，只提供很少几张正常样本时，如何利用视觉信息和语言信息判断异常，并进一步研究不同信息之间如何互相补充。第一阶段先复现已有方法，建立公平、可重复的基准；第二阶段才设计动态融合方法。",
    )
    add_table(
        doc,
        ["报告信息", "内容"],
        [
            ("报告日期", "2026年7月30日"),
            ("项目目录", r"D:\STUDY\My_github\sci_project"),
            ("数据集", "MVTec AD（15类）和 VisA（12类）"),
            ("第一阶段目标", "完成1/2/4-shot、3-seed基准复现和统一评测"),
            ("当前阶段", "PromptAD VisA 4-shot、seed 0正在运行"),
        ],
        [2100, 7260],
    )

    doc.add_heading("一、先用一句话说明项目现在走到哪里", level=1)
    add_para(
        doc,
        "项目已经从“搭建环境、准备数据、验证单个方法”进入“补齐完整实验矩阵”的阶段。数据、固定抽样清单、统一评测程序和三种主要方法的 VisA 完整矩阵已经完成；目前正在补 PromptAD 的剩余配置，同时准备 ReMP-AD 和 AdaptCLIP 的 Gate A 测试。"
    )
    add_callout(
        doc,
        "当前最重要的事实",
        "现在还没有开始动态融合模块。这样安排是为了先保证所有基准方法的输入、抽样、指标和结果都可靠，否则后面的融合结果没有可信的比较基础。",
        fill="FFF8E8",
    )

    doc.add_heading("二、项目从开始到现在做了什么", level=1)
    doc.add_heading("1. 建立项目骨架和管理方式", level=2)
    add_para(doc, "项目最开始没有现成的代码和实验记录，因此先建立了统一的项目结构。主要内容包括：")
    for x in [
        "建立 PLAN.md，明确第一阶段目标、实验门槛、方法顺序和后续融合入口。",
        "建立 PROJECT_STATUS.md，持续记录已完成、正在运行、阻塞项和下一步。",
        "建立 experiments/registry.csv，每一行对应一组实验，记录数据集、类别、shot、seed、指标、checkpoint和代码版本。",
        "把第三方方法代码、项目自己的脚本、数据文件和实验输出分开管理。",
        "对第三方源码的必要改动保存为 patches/*.patch，避免直接修改后无法追溯。",
        "建立独立 Python 环境，避免不同方法的 PyTorch、CLIP、FAISS版本互相冲突。",
    ]:
        add_bullet(doc, x)

    doc.add_heading("2. 准备和检查数据", level=2)
    add_para(doc, "数据准备是整个项目的基础。数据不正确，后面的指标即使算出来也不能用于论文。")
    add_table(
        doc,
        ["数据集", "完成内容", "结果"],
        [
            ("VisA", "下载、解压、整理成统一单类目录、生成metadata", "12类数据可正常读取"),
            ("MVTec AD", "通过官方许可流程取得压缩包，计算SHA256并解压", "15类数据可正常读取"),
            ("MVTec完整性", "检查train/test/ground_truth、图像和掩码对应关系", "错误数为0"),
            ("统一抽样", "生成1/2/4-shot、seed 0/1/2嵌套manifest", "315个MVTec引用项校验通过"),
        ],
        [1800, 5200, 2360],
    )
    add_para(
        doc,
        "MVTec压缩包大小为5,264,982,680字节，SHA256为CF4313B13603BEC67ABB49CA959488F7EEDCE2A9F7795EC54446C649AC98CD3D。MVTec manifest的SHA256为0a04260becf73635dd1ffdbe6fb8f16047e6086a9d431dc973a70f2b258fe59f。"
    )
    add_callout(
        doc,
        "为什么要固定manifest",
        "每种方法都必须使用同一批正常参考图。这样比较时，差异主要来自方法本身，而不是某个方法随机抽到了更容易的图片。",
    )

    doc.add_heading("3. 建立统一输出和评测程序", level=2)
    add_para(doc, "不同论文的代码通常使用不同的文件格式和指标计算方式。项目因此增加了一层方法无关的统一评测。")
    add_table(
        doc,
        ["指标", "简单解释", "用途"],
        [
            ("图像 AUROC", "把整张图判断为正常或异常的排序能力", "判断异常检测能力"),
            ("图像 AP", "在异常样本较少时，更关注准确率和召回率", "补充AUROC"),
            ("图像 F1-max", "在所有阈值中取F1最高值", "观察分类阈值效果"),
            ("像素 AUROC", "异常热力图定位缺陷像素的能力", "判断定位质量"),
            ("像素 AP", "异常像素很少时的区域检测质量", "补充像素AUROC"),
            ("AUPRO", "限制误报范围后，衡量异常区域覆盖能力", "工业缺陷定位核心指标"),
        ],
        [1800, 5000, 2560],
    )
    add_para(doc, "所有统一结果都以NPZ保存，并进行样本数、标签、掩码、异常图尺寸和sample_id检查。当前实验注册表共有52条记录，字段校验错误为0。")

    doc.add_page_break()
    doc.add_heading("三、已经完成的基准实验", level=1)
    doc.add_heading("1. AnomalyCLIP", level=2)
    add_para(doc, "AnomalyCLIP主要用于建立项目早期基准，也用于验证图像分数和像素热力图能否正确缓存、读取和重新评测。")
    for x in [
        "完成官方源码、权重、环境和VisA/MVTec数据路径检查。",
        "完成VisA全12类推理和预测缓存。",
        "完成MVTec全15类、1725张测试图的推理和统一评测。",
        "修复官方代码在慢速AUPRO计算前不保存预测的问题，使长任务可以恢复。",
        "完成公共NPZ评测与官方结果的交叉核对。",
    ]:
        add_bullet(doc, x)
    add_table(
        doc,
        ["数据集/配置", "图像AUROC", "图像AP", "像素AUROC", "AUPRO"],
        [
            ("MVTec 15类官方推理", "93.89%", "97.05%", "94.23%", "88.33%"),
            ("VisA 12类缓存评测", "81.97%", "85.34%", "93.93%", "83.60%"),
        ],
        [3200, 1500, 1500, 1500, 1660],
    )

    doc.add_heading("2. PatchCore", level=2)
    add_para(doc, "PatchCore作为传统视觉特征库方法，主要用来提供一个不依赖语言模型的视觉基线。")
    for x in [
        "完成官方源码和独立环境检查。",
        "使用CPU FAISS，降低Windows GPU FAISS不稳定的风险。",
        "完成VisA 12类×1/2/4-shot×3-seed完整矩阵。",
        "完成MVTec bottle、1-shot、seed 0 Gate A。",
        "统一使用WideResNet50的layer2和layer3特征，并记录128px/256维工程协议。",
    ]:
        add_bullet(doc, x)
    add_table(
        doc,
        ["VisA shot", "图像AUROC mean±std", "像素AUROC mean±std", "AUPRO mean±std"],
        [
            ("1-shot", "68.03±1.19%", "85.96±0.66%", "50.60±0.29%"),
            ("2-shot", "72.91±0.47%", "90.04±0.21%", "57.96±0.37%"),
            ("4-shot", "78.68±1.11%", "91.95±0.27%", "62.21±1.32%"),
        ],
        [1600, 2800, 2800, 2160],
    )
    add_para(doc, "MVTec bottle Gate A结果：图像AUROC 92.54%，像素AUROC 92.71%，AUPRO 69.59%。")

    doc.add_heading("3. WinCLIP+", level=2)
    add_para(doc, "WinCLIP+是CLIP视觉和窗口级异常定位方法，也是项目中较重要的视觉—语言基线。")
    for x in [
        "完成VisA 12类×1/2/4-shot×3-seed完整矩阵。",
        "加入统一manifest读取逻辑，避免官方代码重新随机抽样。",
        "缓存每个类别的测试特征，后续不同shot和seed可以复用。",
        "完成MVTec 15类、1-shot、seed 0 Gate B。",
    ]:
        add_bullet(doc, x)
    add_table(
        doc,
        ["VisA shot", "图像AUROC mean±std", "像素AUROC mean±std", "AUPRO mean±std"],
        [
            ("1-shot", "69.96±0.19%", "89.98±0.24%", "67.34±0.73%"),
            ("2-shot", "71.57±0.34%", "90.50±0.25%", "68.32±0.75%"),
            ("4-shot", "72.58±0.11%", "90.88±0.08%", "69.05±0.51%"),
        ],
        [1600, 2800, 2800, 2160],
    )
    add_para(doc, "MVTec 15类Gate B结果：图像AUROC 76.79%，像素AUROC 86.51%，AUPRO 70.64%。")

    doc.add_heading("4. AnomalyDINO", level=2)
    add_para(doc, "AnomalyDINO是DINOv2视觉特征方法，代表较新的纯视觉少样本基线。")
    for x in [
        "完成DINOv2权重、缓存、Windows路径和CPU FAISS适配。",
        "完成VisA 12类×1/2/4-shot×3-seed完整矩阵。",
        "所有VisA运行都包含2162张测试图，schema错误为0。",
        "完成MVTec bottle、1-shot、seed 0 Gate A。",
        "修正MVTec掩码命名兼容：MVTec使用xxx_mask.png。",
    ]:
        add_bullet(doc, x)
    add_table(
        doc,
        ["VisA shot", "图像AUROC mean±std", "像素AUROC mean±std", "AUPRO mean±std"],
        [
            ("1-shot", "89.40±0.98%", "97.97±0.12%", "92.21±0.64%"),
            ("2-shot", "91.40±0.69%", "98.28±0.04%", "93.10±0.36%"),
            ("4-shot", "92.58±0.24%", "98.45±0.04%", "93.69±0.14%"),
        ],
        [1600, 2800, 2800, 2160],
    )
    add_para(doc, "MVTec bottle Gate A结果：图像AUROC 99.92%，像素AUROC 98.96%，AUPRO 96.46%。")

    doc.add_page_break()
    doc.add_heading("四、PromptAD的完成情况", level=1)
    add_para(doc, "PromptAD与PatchCore等方法不同，它会利用目标数据集中的正常样本学习Prompt。因此报告中单独标记target_normal_tuning=true，不能和完全训练-free的方法混为一谈。")
    add_heading = doc.add_heading
    doc.add_heading("1. 已完成的PromptAD工作", level=2)
    for x in [
        "固定PromptAD官方源码commit 0f86ce0dc1ed59007d51348d8d566aed31360cf9。",
        "建立独立环境.venv-promptad，CUDA可用。",
        "改造VisA数据加载，使其读取项目数据目录和冻结manifest。",
        "分别处理分类分支和分割分支。",
        "增加可选NPZ导出，将分类分数和分割异常图合并成统一预测。",
        "建立断点恢复逻辑：分类完成、分割完成、NPZ导出和统一评测分别有标记。",
    ]:
        add_bullet(doc, x)
    doc.add_heading("2. PromptAD VisA已完成配置", level=2)
    add_table(
        doc,
        ["配置", "类别/样本", "图像AUROC", "像素AUROC", "AUPRO", "状态"],
        [
            ("1-shot seed 0", "12类/2162张", "80.25%", "96.20%", "81.73%", "已完成"),
            ("2-shot seed 0", "12类/2162张", "81.15%", "96.79%", "82.25%", "已完成"),
            ("4-shot seed 0", "当前训练中", "—", "—", "—", "进行中"),
        ],
        [1900, 2200, 1600, 1600, 1200, 860],
    )
    add_para(doc, "1-shot和2-shot的全部12类结果都已经完成，且两个配置都通过了统一schema检查。")
    doc.add_heading("3. 当前实时训练", level=2)
    add_callout(
        doc,
        "当前GPU任务",
        "截至2026年7月30日09:35，PromptAD VisA 4-shot、seed 0正在运行。candle类别已完成，当前capsules类别正在进行分类训练，之后还会继续分割训练和NPZ合并。训练命令使用100 epochs，单个类别通常需要较长时间。",
        fill="FFF8E8",
    )
    add_para(doc, "当前训练日志和监控脚本：")
    add_bullet(doc, r"日志：outputs\logs\promptad\visa\seed_0_shot_4\matrix.log")
    add_bullet(doc, r"监控：powershell -ExecutionPolicy Bypass -File scripts\watch_project_progress.ps1 -Seed 0 -Shot 4")

    doc.add_heading("五、MVTec AD当前进度", level=1)
    add_table(
        doc,
        ["方法", "当前完成内容", "下一道门槛"],
        [
            ("AnomalyCLIP", "15类官方推理和统一评测已完成", "少样本版本待补"),
            ("PatchCore", "bottle、1-shot、seed 0 Gate A", "15类1-shot Gate B"),
            ("WinCLIP+", "15类、1-shot、seed 0 Gate B", "2/4-shot和其他seed"),
            ("AnomalyDINO", "bottle、1-shot、seed 0 Gate A", "15类1-shot Gate B"),
            ("PromptAD", "尚未开始MVTec Gate A", "bottle、1-shot、seed 0"),
            ("ReMP-AD", "源码和环境准备完成", "VisA或MVTec Gate A"),
            ("AdaptCLIP", "源码审计完成，待checkpoint", "下载权重后Gate A"),
        ],
        [1700, 4800, 2860],
    )
    add_para(doc, "MVTec数据本身已经不是阻塞项。当前主要工作是把通过VisA Gate的各个方法，依次扩展到MVTec完整矩阵。")

    doc.add_heading("六、ReMP-AD和AdaptCLIP", level=1)
    doc.add_heading("1. ReMP-AD", level=2)
    for x in [
        "官方仓库已下载，commit为d3fbc46adfd91406859b90dece65c221343096c7。",
        "独立环境.venv-rempad已建立。",
        "PyTorch 2.0.0+cu118、torchvision 0.15.1+cu118、NumPy 1.24.4已验证。",
        "CUDA导入成功，OpenCV、pandas、scikit-image、scikit-learn、timm、transformers等依赖已安装。",
        "官方README没有提供可直接使用的预训练权重，默认流程是训练后测试。",
    ]:
        add_bullet(doc, x)
    add_para(doc, "下一步先检查ReMP-AD能否读取项目数据，再完成bottle单类别Gate A。")
    doc.add_heading("2. AdaptCLIP", level=2)
    for x in [
        "官方仓库已下载，commit为354d9e3332ec5348b3d8e4439111d34f8e94c0a9。",
        "已审计README、requirements和官方测试脚本。",
        "官方权重需要从Hugging Face下载到adaptclip_checkpoints目录。",
        "官方shell脚本中存在重复的数据集测试块，需要本地修复并记录补丁。",
        "由于本机只有6 GB显存，首先要用batch size 1验证权重能否加载。",
    ]:
        add_bullet(doc, x)

    doc.add_heading("七、曾经出现的中断问题和解决方式", level=1)
    add_table(
        doc,
        ["现象", "实际原因", "处理结果"],
        [
            ("长时间没有日志", "PromptAD官方代码不打印每个epoch进度", "用进程、GPU、checkpoint和marker判断真实状态"),
            ("后台只跑了第一个类别", "PowerShell数组参数传递不正确", "改为单类别断点恢复或使用默认类别列表"),
            ("GPU显存高但没有输出", "测试集评测和CPU预处理占用时间", "确认CPU时间和GPU进程仍在增长"),
            ("ReMP-AD环境警告", "最新OpenCV拉取NumPy 2.x", "固定OpenCV 4.8.1.78和NumPy 1.24.4"),
            ("AUPRO长时间不结束", "官方实现阈值计算很慢且最后才保存", "先缓存预测，再使用统一评测器"),
        ],
        [2300, 3800, 3260],
    )
    add_callout(
        doc,
        "结论",
        "目前没有发现数据损坏、显存不足或模型本身无法运行的问题。之前的“中断”主要是日志不明显、长任务耗时和PowerShell参数传递问题；现在已经用marker、NPZ缓存和监控脚本把这些问题变成可恢复流程。",
        fill="E8EEF5",
    )

    doc.add_page_break()
    doc.add_heading("八、接下来要完成的内容和每一步的目的", level=1)
    doc.add_heading("阶段A：完成PromptAD VisA矩阵", level=2)
    for x in [
        "完成4-shot、seed 0的12类运行，确认shot增加后结果是否稳定变化。",
        "完成seed 1和seed 2的1/2/4-shot，得到mean±std，而不是只看单个随机种子。",
        "检查每个配置的2162张测试图、NPZ数组、样本顺序和指标。",
        "形成PromptAD和PatchCore、WinCLIP+、AnomalyDINO的VisA对比表。",
    ]:
        add_number(doc, x)
    add_para(doc, "目的：完成PromptAD在VisA上的公平少样本比较，并把目标域Prompt学习方法与训练-free方法分开解释。")

    doc.add_heading("阶段B：完成PromptAD MVTec Gate A", level=2)
    for x in [
        "先跑MVTec bottle、1-shot、seed 0。",
        "检查train/good引用是否来自manifest。",
        "检查MVTec异常掩码的_mask.png命名是否正确。",
        "完成分类、分割、NPZ合并和统一评测。",
    ]:
        add_number(doc, x)
    add_para(doc, "目的：先用单类别验证PromptAD的MVTec适配，避免直接投入15类长时间训练。")

    doc.add_heading("阶段C：补齐ReMP-AD和AdaptCLIP", level=2)
    for x in [
        "每个方法先完成源码、环境、权重和数据入口审计。",
        "每个方法先做一个类别的Gate A。",
        "Gate A通过后再做15类Gate B。",
        "只有输出格式和指标都稳定，才进入完整1/2/4-shot、3-seed矩阵。",
    ]:
        add_number(doc, x)
    add_para(doc, "目的：加入老师要求的2025—2026年较新的方法，同时避免因为依赖或权重问题反复重跑大矩阵。")

    doc.add_heading("阶段D：补齐MVTec完整矩阵", level=2)
    add_para(doc, "每个方法都按照下面的三级门槛执行：")
    add_table(
        doc,
        ["门槛", "实验范围", "通过条件"],
        [
            ("Gate A", "bottle × 1-shot × seed 0", "能运行、能导出NPZ、schema和指标正确"),
            ("Gate B", "15类 × 1-shot × seed 0", "所有类别完成、样本数正确、无schema错误"),
            ("Gate C", "15类 × 1/2/4-shot × 3 seeds", "得到完整mean±std和可追溯日志"),
        ],
        [1500, 4000, 3860],
    )
    add_para(doc, "建议顺序：PromptAD → ReMP-AD → AdaptCLIP → PatchCore/AnomalyDINO/WinCLIP+的MVTec扩展 → 总结果汇总。")

    doc.add_heading("阶段E：统计汇总和论文材料", level=2)
    for x in [
        "生成按方法、数据集、shot分组的mean±std主表。",
        "生成逐类别结果、失败案例和效率表。",
        "对比论文原始数字、官方本地数字和统一协议数字。",
        "记录输入尺寸、源域、是否目标域Prompt训练、显存和运行时间。",
        "建立文本分支和视觉分支的互补性统计。",
    ]:
        add_bullet(doc, x)
    add_para(doc, "目的：形成可以直接用于论文实验章节的证据，而不是只保存几组零散的高分结果。")

    doc.add_heading("阶段F：第二阶段动态融合", level=2)
    add_para(doc, "动态融合必须等第一阶段基准矩阵稳定后再开始。计划包括：")
    for x in [
        "分别保存文本分支和视觉分支的图像分数、异常图和置信度。",
        "分析两条分支在哪些类别和缺陷区域互相补充。",
        "先建立固定加权融合基线。",
        "再设计基于置信度、一致性和分数间隔的动态路由。",
        "完成仅文本、仅视觉、固定融合、动态融合的消融实验。",
        "检查路由模块是否使用了测试异常信息，防止数据泄漏。",
    ]:
        add_number(doc, x)
    add_para(doc, "目的：回答论文真正的研究问题：在少样本工业异常检测中，什么时候应该更相信语言证据，什么时候应该更相信视觉参考，动态选择是否确实带来稳定提升。")

    doc.add_heading("九、项目完成的最终标准", level=1)
    for x in [
        "MVTec AD和VisA都有固定、可复查的1/2/4-shot、3-seed清单。",
        "主要方法都至少完成Gate A；可运行方法完成完整矩阵，无法复现的方法有明确的阻塞证据。",
        "每组实验都有命令、环境、commit、权重、日志和预测缓存。",
        "所有结果都能从NPZ重新计算，实验注册表无结构错误。",
        "报告中清楚区分zero-shot、few-shot reference和target normal tuning。",
        "动态融合只建立在已经通过基准质量检查的结果之上。",
    ]:
        add_bullet(doc, x)
    add_callout(
        doc,
        "当前判断",
        "项目基础工作已经比较扎实，下一阶段的主要工作量来自长时间训练和完整矩阵扩展，而不是重新搭建基础设施。当前最合理的做法是让PromptAD 4-shot继续完成，同时并行准备ReMP-AD和AdaptCLIP的Gate A条件。",
    )
    add_para(doc, "本报告依据项目目录中的PLAN.md、PROJECT_STATUS.md、实验注册表、统一评测结果、运行日志和预测缓存整理。报告中的“正在运行”状态以2026年7月30日09:35的进程检查为准。", after=0)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()

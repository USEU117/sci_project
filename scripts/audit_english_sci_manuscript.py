from __future__ import annotations

import hashlib
import json
import re
import sys
import zipfile
from pathlib import Path

from docx import Document


def word_count(text: str) -> int:
    return len(re.findall(r"[A-Za-z0-9]+(?:[-'][A-Za-z0-9]+)*", text))


def main() -> None:
    path = Path(sys.argv[1]).resolve()
    document = Document(path)
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs]
    full_text = " ".join(text for text in paragraphs if text)

    abstract_heading_index = paragraphs.index("Abstract")
    abstract = next(text for text in paragraphs[abstract_heading_index + 1:] if text)
    keywords = next(text for text in paragraphs if text.startswith("Keywords:"))
    reference_heading_index = paragraphs.index("References")
    body_text = " ".join(paragraphs[:reference_heading_index])

    citations = sorted(set(map(int, re.findall(r"\[(\d+)\]", body_text))))
    references = sorted(
        int(match.group(1))
        for text in paragraphs
        if (match := re.match(r"\[(\d+)\]", text))
    )

    with zipfile.ZipFile(path) as archive:
        corrupt_member = archive.testzip()

    report = {
        "path": str(path),
        "file_bytes": path.stat().st_size,
        "sha256": hashlib.sha256(path.read_bytes()).hexdigest(),
        "paragraphs": len(document.paragraphs),
        "tables_including_equation_layout_tables": len(document.tables),
        "inline_images": len(document.inline_shapes),
        "total_word_tokens": word_count(full_text),
        "abstract_words": word_count(abstract),
        "keyword_count": len(keywords.split(":", 1)[1].split(";")),
        "citations": citations,
        "references": references,
        "citation_reference_match": citations == references,
        "docx_zip_integrity": corrupt_member or "OK",
    }
    print(json.dumps(report, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()

from __future__ import annotations

import csv
import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "outputs" / "project_progress_report_20260811"
DOCX_PATH = OUT_DIR / "少样本工业异常检测项目全程进展与后续工作说明_20260811.docx"

DOC_SKILL = Path(
    r"C:\Users\lynle\.codex\plugins\cache\openai-primary-runtime\documents"
    r"\26.805.11740\skills\documents"
)
sys.path.insert(0, str(DOC_SKILL / "scripts"))
from table_geometry import apply_table_geometry  # noqa: E402


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "16324F"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
PALE_BLUE = "F4F8FC"
GREEN = "2E7D32"
PALE_GREEN = "EAF4EA"
GOLD = "7A5A00"
PALE_GOLD = "FFF4CE"
RED = "9B1C1C"
PALE_RED = "FCE8E6"
GRAY = "666666"
WHITE = "FFFFFF"


def set_run_font(run, size=None, bold=None, color=None, italic=None):
    run.font.name = "Calibri"
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    r_fonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tag = OxmlElement("w:tblHeader")
    tag.set(qn("w:val"), "true")
    tr_pr.append(tag)


def set_cell_text(cell, text, *, bold=False, color=None, size=9.2, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.08
    run = p.add_run(str(text))
    set_run_font(run, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell)


def add_table(doc, headers, rows, widths, *, header_fill=LIGHT_GRAY, font_size=9.0,
              aligns=None, first_col_bold=False):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    apply_table_geometry(table, widths, table_width_dxa=sum(widths), indent_dxa=120)
    for i, value in enumerate(headers):
        shade_cell(table.rows[0].cells[i], header_fill)
        set_cell_text(table.rows[0].cells[i], value, bold=True, color=NAVY, size=font_size,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
    repeat_header(table.rows[0])
    for row_values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_values):
            align = aligns[i] if aligns else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[i], value, bold=(first_col_bold and i == 0), size=font_size, align=align)
    for row in table.rows:
        row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    return table


def add_body(doc, text, *, bold_lead=None, color=None, keep=False):
    p = doc.add_paragraph()
    p.paragraph_format.keep_together = keep
    if bold_lead and text.startswith(bold_lead):
        first = p.add_run(bold_lead)
        set_run_font(first, bold=True, color=color)
        rest = p.add_run(text[len(bold_lead):])
        set_run_font(rest)
    else:
        run = p.add_run(text)
        set_run_font(run, color=color)
    return p


def add_bullet(doc, text, *, level=0, status=None):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.left_indent = Inches(0.5 if level == 0 else 0.75)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(5)
    if status:
        run = p.add_run(f"{status} ")
        color = GREEN if status == "已完成" else GOLD if status in {"待完成", "部分完成"} else RED
        set_run_font(run, bold=True, color=color)
    run = p.add_run(text)
    set_run_font(run)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run)
    return p


def add_callout(doc, label, text, *, fill=PALE_BLUE, color=NAVY):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    apply_table_geometry(table, [9360], table_width_dxa=9360, indent_dxa=120)
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.12
    r = p.add_run(f"{label}：")
    set_run_font(r, bold=True, color=color, size=10.5)
    r = p.add_run(text)
    set_run_font(r, size=10.5)
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=GRAY)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, value, end])
    tail = paragraph.add_run(" 页")
    set_run_font(tail, size=9, color=GRAY)


def setup_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    normal.paragraph_format.widow_control = True

    title = styles["Title"]
    title.font.name = "Calibri"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    title.font.size = Pt(25)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(NAVY)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(8)

    settings = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in settings.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Bullet 2", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.line_spacing = 1.167

    if "Small Note" not in styles:
        note = styles.add_style("Small Note", WD_STYLE_TYPE.PARAGRAPH)
    else:
        note = styles["Small Note"]
    note.font.name = "Calibri"
    note._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    note.font.size = Pt(9)
    note.font.color.rgb = RGBColor.from_string(GRAY)
    note.paragraph_format.space_after = Pt(4)


def load_current_data():
    visa_path = ROOT / "experiments" / "summaries" / "visa_baseline_main_table_20260803.csv"
    visa_rows = list(csv.DictReader(visa_path.open("r", encoding="utf-8-sig")))
    btad = json.loads((ROOT / "experiments" / "dynamic_fusion" / "v2" / "btad_frozen_evaluation" / "report.json").read_text(encoding="utf-8"))
    freeze = json.loads((ROOT / "experiments" / "dynamic_fusion" / "v2" / "parameter_freeze" / "manifest.json").read_text(encoding="utf-8"))
    queue = json.loads((ROOT / "outputs" / "logs" / "promptad_mvtec_resumable_queue" / "status.json").read_text(encoding="utf-8-sig"))
    return visa_rows, btad, freeze, queue


def pct(value):
    return f"{float(value) * 100:.2f}%"


def build():
    visa_rows, btad, freeze, queue = load_current_data()
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    setup_styles(doc)

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hrun = header.add_run("少样本工业异常检测项目｜全程进展说明")
    set_run_font(hrun, size=9, color=GRAY)
    add_page_number(section.footer.paragraphs[0])

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.add_run("少样本工业异常检测项目\n全程进展与后续工作说明")
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(16)
    run = subtitle.add_run("从基准复现、动态融合 V1/V2 到论文初稿的完整大白话说明")
    set_run_font(run, size=13.5, color=GRAY)

    metadata = [
        ("项目主题", "基于不确定性路由的少样本工业异常视觉—语言证据融合"),
        ("更新日期", "2026 年 8 月 11 日"),
        ("当前机器", "Windows；NVIDIA RTX 3060 Laptop，6 GB 显存"),
        ("报告用途", "帮助快速理解从项目开始到现在做了什么、做到了什么、还缺什么"),
    ]
    for label, value in metadata:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(f"{label}：")
        set_run_font(r, bold=True, color=NAVY)
        r = p.add_run(value)
        set_run_font(r)

    add_callout(
        doc,
        "一句话结论",
        "项目已经完成了可审计的实验基础、两套主要数据集的大部分基线、动态融合 V1 的失败分析、DynamicFusion V2 的代码与冻结验证，以及中文和英文论文初稿。现在不是“项目刚开始”，而是已经进入收尾和论文定稿阶段；但还不能说全部结束，因为 PromptAD 的 MVTec 还缺 5 组，ReMP-AD 和 AdaptCLIP 尚未通过 Gate A，而且 V2 最终退回纯视觉方案，动态文本融合的稳定提升目前没有被证明。",
        fill=PALE_GOLD,
        color=GOLD,
    )

    doc.add_heading("阅读提示", level=1)
    add_body(doc, "这份报告故意不用复杂论文语言。为了避免误解，全文把项目状态分成四类：")
    add_bullet(doc, "已经完整跑完、有日志、有审计、有结果，可以用于论文。", status="已完成")
    add_bullet(doc, "做过实验，也得到了重要结论，但结果没有达到原先预期。", status="已完成")
    add_bullet(doc, "已有代码、环境或断点，但矩阵还没有全部补齐。", status="部分完成")
    add_bullet(doc, "还没有通过最小实验门槛，不能写成本地已比较方法。", status="待完成")

    doc.add_page_break()
    doc.add_heading("目录式概览", level=1)
    for item in [
        "1. 项目最初想解决什么问题",
        "2. 从项目开始到现在的时间线",
        "3. 数据、协议、评价和工程基础做了什么",
        "4. 第一阶段基线复现完成到什么程度",
        "5. 第二阶段 DynamicFusion V1 做了什么、为什么没有达到预期",
        "6. DynamicFusion V2 做了什么、最终得到什么结论",
        "7. 论文、文献和图表材料完成到什么程度",
        "8. 项目当前实时状态与完成度判断",
        "9. 后续还要做哪些工作、预计耗时和 GPU 需求",
        "10. 推荐的收尾路线和最终判断",
    ]:
        add_bullet(doc, item)

    doc.add_heading("0. 先把几个词说清楚", level=1)
    glossary = [
        ("异常检测", "判断一张工业图片是否有缺陷，并尽量在图上指出缺陷位置。"),
        ("少样本 / K-shot", "每个产品类别只给模型 K 张正常图片作参考；本项目用 K=1、2、4。"),
        ("seed", "随机种子。相当于换一组正常参考图，检查结果是否稳定。本项目固定 seed 0、1、2。"),
        ("基线", "用来比较的已有方法。只有基线先跑可靠，后面才能判断新方法到底有没有提升。"),
        ("分支", "同一张图由两种方法分别给出证据。本项目的视觉分支主要是 AnomalyDINO，文本分支主要是 AnomalyCLIP。"),
        ("动态融合", "不是一直使用固定比例，而是根据每张图或每个像素的情况，决定更相信视觉证据还是文本证据。"),
        ("Gate A", "最小门槛实验。先跑一个类别、一个 seed、一个 shot，确认代码、显存、输出和指标都正确，再决定是否跑完整矩阵。"),
        ("冻结", "参数和规则确定后不再根据最终测试结果修改，防止为了让数字好看而反复调参。"),
        ("审计", "检查样本顺序、标签、掩码、配置、哈希和禁止信息，证明结果不是错位或数据泄漏造成的。"),
    ]
    add_table(doc, ["名词", "大白话解释"], glossary, [1900, 7460], first_col_bold=True, font_size=9.4)

    doc.add_heading("1. 项目最初想解决什么问题", level=1)
    add_body(doc, "工业异常检测有一个很现实的问题：新产品刚上线时，通常只有少量正常图片，异常图片很少甚至完全没有。项目因此采用 1/2/4-shot 设定，只允许每个类别使用 1、2 或 4 张正常图作为参考。")
    add_body(doc, "项目同时观察到两类方法各有长处：纯视觉方法擅长抓细小纹理和结构变化，视觉语言方法能够利用“正常/异常”的语义信息。最初设想是把两者结合：视觉分支负责细节，文本分支负责语义，再由一个路由器动态决定每张图、每个像素该信谁。")
    add_callout(doc, "原始研究问题", "能不能在不使用测试标签、不偷看测试集整体统计量的情况下，只根据少量正常参考图和两条分支自己的输出，安全地完成视觉—文本动态融合，并稳定超过最强单分支？")
    add_body(doc, "为了回答这个问题，项目被分成两大阶段：第一阶段先建立公平、可重复的基线；第二阶段再做动态融合。后来又因为 V1 暴露出严重问题，继续开发了 V2。")

    doc.add_heading("2. 从项目开始到现在的时间线", level=1)
    timeline = [
        ("7 月 24 日前后", "立项与第一阶段规划", "确定 MVTec AD、VisA、1/2/4-shot、3 seeds、统一指标和方法清单。"),
        ("7 月 25—28 日", "环境、数据和最早基线", "下载 VisA、建立独立虚拟环境、完成 AnomalyCLIP/WinCLIP/PatchCore 冒烟与统一评价。"),
        ("7 月 28—31 日", "VisA 完整矩阵", "PatchCore、WinCLIP+、AnomalyDINO、PromptAD 逐步完成；建立断点续跑和自动 GPU 队列。"),
        ("7 月 30—8 月 5 日", "DynamicFusion V1", "完成缓存对齐、正常参考校准、图像/像素动态权重、温度与 margin 消融、独立最终验证。"),
        ("8 月 3—9 日", "MVTec 与状态纠偏", "MVTec 数据到位并完成主要矩阵；重新审计旧记录，修正过期状态，完成 V1 科学失败分析和可视化。"),
        ("8 月 10 日", "论文与文献", "完成中文论文 V0.1、英文 SCI 风格 V0.2/V0.3，改进主图，扩展到 37 篇参考文献并整理 2026 年相关工作。"),
        ("8 月 10—11 日", "DynamicFusion V2", "实现保序校准、超支持范围检测、安全回退、独立图像/像素路由；准备 MPDD/BTAD 并跑完整冻结验证。"),
        ("8 月 11 日当前", "进入收尾阶段", "V2 已冻结，BTAD 9/9 完成；进入基线补齐、结果整理和论文定稿。"),
    ]
    add_table(doc, ["时间", "主要阶段", "完成内容"], timeline, [1500, 2200, 5660], header_fill=LIGHT_BLUE, font_size=8.8, first_col_bold=True)

    doc.add_heading("3. 数据、协议、评价和工程基础做了什么", level=1)
    doc.add_heading("3.1 项目不是简单地下载代码运行", level=2)
    add_body(doc, "从最开始就建立了自己的项目骨架，把第三方方法、研究代码、数据、清单、日志、预测缓存、图表和论文材料分开保存。每个方法尽量使用独立虚拟环境，避免不同 PyTorch、CLIP、FAISS 和 CUDA 版本互相冲突。")
    add_bullet(doc, "建立 PLAN.md、PROJECT_STATUS.md、NEXT_ACTIONS.md 和实验注册表。", status="已完成")
    add_bullet(doc, "记录方法源码来源、commit、checkpoint 来源与 SHA256。", status="已完成")
    add_bullet(doc, "建立统一 NPZ 输出，让不同方法都能交给同一个评价程序。", status="已完成")
    add_bullet(doc, "建立断点续跑、任务状态 JSON、完成标记和自动切换队列。", status="已完成")
    add_bullet(doc, "失败运行不覆盖成功结果，失败日志保留作为证据。", status="已完成")

    doc.add_heading("3.2 数据集已经准备到什么程度", level=2)
    data_rows = [
        ("VisA", "12 类，2,162 张测试图", "官方数据已下载、解压、校验；1/2/4-shot × 3 seeds 清单完成", "主要基线和 V1 已完成"),
        ("MVTec AD", "15 类，1,725 张测试图", "通过官方许可流程取得；压缩包哈希、目录、掩码、metadata 和嵌套清单均完成", "主要基线完成，PromptAD 仍缺 5 组"),
        ("MPDD", "6 类，458 张测试图", "作为 V2 开发集；归档、类别、掩码、1/2/4-shot × 3 seeds 全部审计", "V2 选择和参数冻结已完成"),
        ("BTAD", "3 类，741 张测试图", "作为 V2 最终保持集；含 290 张异常图；清单和掩码审计完成", "冻结后 9/9 最终验证已完成"),
    ]
    add_table(doc, ["数据集", "规模", "准备工作", "现在用途"], data_rows, [1400, 1800, 3600, 2560], font_size=8.6, header_fill=LIGHT_BLUE, first_col_bold=True)

    doc.add_heading("3.3 统一实验协议为什么重要", level=2)
    add_body(doc, "不同官方方法原本会自己随机选正常图、使用不同尺寸、不同指标和不同输出格式。如果直接拿官方数字横向比较，容易出现“看起来公平，实际不是同一条件”的问题。因此项目固定了以下规则：")
    for text in [
        "统一使用 1/2/4-shot 和 seed 0/1/2。",
        "同一个 seed 下，1-shot 是 2-shot 的子集，2-shot 是 4-shot 的子集。",
        "各方法必须读取同一份正常参考图清单，不能自己重新随机抽样。",
        "测试异常图和掩码只用于最终评价，不能用于校准、训练或路由。",
        "图像级指标包括 AUROC、AP、F1-max；像素级包括 AUROC、AP、AUPRO。",
        "先逐类别计算，再做宏平均，不能把所有像素粗暴拼起来冒充类别均值。",
        "PromptAD 使用目标类别正常图学习 prompt，必须标记 target_normal_tuning=true。",
        "AnomalyCLIP 的 zero-shot 结果单独报告，不能冒充 1/2/4-shot 矩阵。",
    ]:
        add_bullet(doc, text)

    doc.add_heading("3.4 统一评价和测试", level=2)
    add_body(doc, "项目已经实现方法无关的评价程序，能检查分数方向、数组形状、样本数量、掩码和预测图对应关系。随着 V2 功能增加，当前项目自己的 CPU 测试为 49/49 通过。")
    add_callout(doc, "为什么这很重要", "异常检测中最危险的错误不一定会让程序报错。例如样本顺序错一位、掩码双线性缩放、分数方向相反，都可能产生一组看似正常但实际无效的数字。项目通过统一 schema 和审计尽量把这些问题在汇总前拦住。", fill=PALE_GREEN, color=GREEN)

    doc.add_heading("4. 第一阶段基线复现完成到什么程度", level=1)
    add_body(doc, "第一阶段的主要目的不是发明新算法，而是证明数据、代码、评价和实验管理都可靠。现在 VisA 的四个主要少样本基线已经完整闭合；MVTec 的三个主要基线已经闭合，PromptAD 还差 5 组。")

    doc.add_heading("4.1 VisA：四种方法全部完成", level=2)
    visa_status = [
        ("PatchCore", "9/9", "完成", "传统视觉记忆库基线；shot 增加时提升明显"),
        ("WinCLIP+", "9/9", "完成", "CLIP 少样本窗口基线；结果稳定但总体弱于 AnomalyDINO"),
        ("AnomalyDINO", "9/9", "完成", "当前最强视觉分支，也是融合必须保护的基线"),
        ("PromptAD", "9/9", "完成", "目标正常样本 prompt 学习；必须单独标记 target_normal_tuning=true"),
    ]
    add_table(doc, ["方法", "3 seeds × 3 shots", "状态", "简单说明"], visa_status, [1500, 1800, 1300, 4760], font_size=8.8, first_col_bold=True)

    # Compact metric summary: I-AUROC / P-AUROC / AUPRO for each shot.
    metric_rows = []
    for row in visa_rows:
        metric_rows.append((
            row["method"],
            row["shot"],
            pct(row["image_auroc_mean"]),
            pct(row["pixel_auroc_mean"]),
            pct(row["aupro_mean"]),
        ))
    add_table(doc, ["方法", "shot", "图像 AUROC", "像素 AUROC", "AUPRO"], metric_rows,
              [2000, 1000, 2100, 2100, 2160], font_size=8.5,
              aligns=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER])
    add_body(doc, "大白话解读：在 VisA 上，AnomalyDINO 是目前最强且最稳定的视觉基线。PatchCore 会随正常参考图增多明显变好；WinCLIP+ 提升较慢；PromptAD 的像素指标较好，但它使用目标正常样本进行 prompt 学习，协议身份与完全冻结方法不同。")

    doc.add_heading("4.2 MVTec AD：主要矩阵基本完成，PromptAD 还没闭合", level=2)
    mvtec_status = [
        ("PatchCore", "9/9", "完成", "每组 15 类、1,725 样本、零 schema 错误"),
        ("WinCLIP+", "9/9", "完成", "每组 15 类、1,725 样本、零 schema 错误"),
        ("AnomalyDINO", "9/9", "完成", "最强已完成视觉基线；旧 8/9 记录已纠正为 9/9"),
        ("PromptAD", "4/9", "暂停待续", "已完成 s0/k1、s0/k2、s0/k4、s1/k1；还差 5 组"),
        ("DynamicFusion V1", "9/9", "完成", "结果可报告，但明显没有超过 AnomalyDINO"),
        ("AnomalyCLIP", "zero-shot", "单独报告", "MVTec 全 15 类 zero-shot 已完成，不属于少样本矩阵"),
        ("ReMP-AD", "0/9", "待 Gate A", "环境存在，但 manifest/NPZ 适配和最小实验未完成"),
        ("AdaptCLIP", "0/9", "阻塞/待 Gate A", "仍需官方 checkpoint、配置修复和 6 GB 显存检查"),
    ]
    add_table(doc, ["方法", "完成度", "当前状态", "说明"], mvtec_status, [1500, 1200, 1400, 5260], font_size=8.5, header_fill=LIGHT_BLUE, first_col_bold=True)
    add_callout(doc, "PromptAD 当前断点", f"正式队列状态是 {queue['state']}。已完成 {', '.join(queue['completed'])}；待完成 {', '.join(queue['pending'])}。s1/k2 已有 5/30 个阶段标记，可从 capsule 分割继续，不需要从头重跑。", fill=PALE_GOLD, color=GOLD)

    doc.add_heading("4.3 各基线在项目中的作用", level=2)
    roles = [
        ("PatchCore", "最基础的视觉记忆库参照", "说明只靠预训练视觉 patch + 少量正常图能做到什么程度。"),
        ("WinCLIP+", "较早的 CLIP 少样本方法", "说明图文预训练加正常参考的基本水平。"),
        ("AnomalyDINO", "强视觉主线", "是当前最强分支；任何融合如果破坏它，都很难证明有价值。"),
        ("PromptAD", "目标正常样本 prompt 学习", "与本项目都用正常样本，但它是学习 prompt，不是冻结分支可靠性路由。"),
        ("AnomalyCLIP", "冻结文本证据", "给出语义分数和热图，是 V1/V2 的文本分支。"),
        ("ReMP-AD / AdaptCLIP", "近期多模态对比", "对论文相关，但目前还不能写成本地完成比较。"),
    ]
    add_table(doc, ["方法", "角色", "为什么需要它"], roles, [1900, 2600, 4860], font_size=8.8, first_col_bold=True)

    doc.add_heading("5. DynamicFusion V1 做了什么", level=1)
    add_body(doc, "V1 是项目第二阶段的第一版动态融合。它把 AnomalyDINO 作为视觉分支，把 AnomalyCLIP 作为文本分支。每张测试图都有两套图像分数和两张像素热图，路由器根据正常参考图校准后的分数、不确定性和分支差异，为图像级和像素级分别生成权重。")
    add_body(doc, "V1 完成的不是一个简单公式，而是一整套可审计流程：")
    for text in [
        "跨分支 sample_id、顺序、标签和像素图尺寸对齐。",
        "只用正常参考图做分支校准，拒绝使用测试标签或测试集整体统计量。",
        "支持视觉单分支、文本单分支、固定权重和动态权重对照。",
        "图像温度和像素温度分开设计，最终冻结为图像 0.50、像素 0.20。",
        "完成 temperature、decision margin、K=1/2/4、图像/像素权重和类别差异消融。",
        "完成 17/17 冻结运行审计：VisA 8 个、MVTec 9 个。",
        "生成成功案例、失败案例、热图、路由统计、逐类别图和 9-sheet Excel 数据包。",
    ]:
        add_bullet(doc, text, status="已完成")

    doc.add_heading("5.1 V1 为什么没有达到预期", level=2)
    add_callout(doc, "核心问题", "少量正常参考图估计出的分数尺度太小，sigmoid 校准把大量不同分数都压到接近 1，原本很强的 AnomalyDINO 排序被破坏。随后熵又把“接近 1”误认为高置信，路由器没能发现这是超出参考范围，而不是可靠判断。", fill=PALE_RED, color=RED)
    failure_chain = [
        ("第一步", "正常参考太少", "中位数和 MAD 尺度估计不稳定，特别容易过小。"),
        ("第二步", "校准饱和", "大量不同测试分数都变成约 0.999，出现很多并列值。"),
        ("第三步", "不确定性误判", "二元熵看到接近 1 会认为很自信，却不知道这是超出支持范围。"),
        ("第四步", "动态权重破坏排序", "每张图使用不同权重，正常图和异常图的全局顺序被重新打乱。"),
        ("最后", "融合低于强视觉分支", "文本偶尔能补充局部区域，但总体不能抵消视觉排序损失。"),
    ]
    add_table(doc, ["环节", "发生了什么", "大白话解释"], failure_chain, [1200, 2200, 5960], font_size=8.8, header_fill=PALE_RED, first_col_bold=True)
    add_body(doc, "量化证据非常明确：MVTec 中视觉校准分数落在 0.999 以上的平均比例约为 99.99%，VisA 约为 91.54%。这意味着很多本来不同的视觉分数被挤成了几乎相同的值。")

    doc.add_heading("5.2 V1 最终结果应该怎样理解", level=2)
    v1_rows = [
        ("1-shot", "95.71%", "79.43%", "-16.29 个百分点", "82.33%"),
        ("2-shot", "96.86%", "86.37%", "-10.49 个百分点", "90.58%"),
        ("4-shot", "97.46%", "89.52%", "-7.94 个百分点", "91.81%"),
    ]
    add_table(doc, ["MVTec", "AnomalyDINO 图像 AUROC", "V1 图像 AUROC", "差距", "V1 AUPRO"], v1_rows,
              [1100, 2300, 2100, 1900, 1960], font_size=8.8,
              aligns=[WD_ALIGN_PARAGRAPH.LEFT] + [WD_ALIGN_PARAGRAPH.CENTER] * 4)
    add_body(doc, "结论不是“V1 完全没用”。它确实在部分 zipper 等样本的定位上利用了文本热图，也验证了图像级和像素级应该使用不同路由强度。但从整体结果看，V1 没有超过最强视觉分支，不能把论文写成“动态融合全面提升”。")
    add_bullet(doc, "可以写：完成了严格的无泄漏动态融合流程，发现并解释了校准饱和、熵误判和排序破坏。", status="已完成")
    add_bullet(doc, "可以写：部分样本、部分定位指标存在互补收益，图像级和像素级需要分开。", status="已完成")
    add_bullet(doc, "不能写：文本分支能够稳定接管视觉失败样本。", status="不能宣称")
    add_bullet(doc, "不能写：动态融合全面优于 AnomalyDINO。", status="不能宣称")

    doc.add_heading("6. DynamicFusion V2 做了什么", level=1)
    add_body(doc, "V2 不是继续搜索更好的温度，而是针对 V1 的根本问题重写安全机制。为了避免反复使用已经看过的 MVTec/VisA 调参，项目新增 MPDD 作为开发集，BTAD 作为最终保持集。")
    v2_modules = [
        ("保序校准", "把分数变到可比较范围时，尽量保留原视觉分数的高低顺序，避免大量 0.999 并列。"),
        ("超出支持范围检测", "测试分数远离少量正常参考分布时，标记为“不知道”，不再当成高置信。"),
        ("视觉默认安全回退", "文本证据没通过安全条件时，直接保留 AnomalyDINO，不让融合随便伤害强分支。"),
        ("图像/像素彻底分开", "图像分数和像素图使用不同判断、不同权重上限和不同消融开关。"),
        ("禁止信息检查", "路由接口不接收测试标签、掩码、类别测试标签或测试集整体统计量。"),
        ("证据冻结", "RunId、配置、哈希、审计和失败结果全部保留；冻结前不能读取 BTAD 指标。"),
    ]
    add_table(doc, ["V2 模块", "解决什么问题"], v2_modules, [2500, 6860], font_size=9.0, header_fill=LIGHT_BLUE, first_col_bold=True)

    doc.add_heading("6.1 V2 的数据准备和工程验收", level=2)
    for text in [
        "MPDD 6 类、BTAD 3 类完成归档、SHA256、解压、类别和掩码审计。",
        "两个数据集都生成 1/2/4-shot × 3 seeds 嵌套 manifest。",
        "18 组正常参考视图、36 个两分支缓存任务全部完成。",
        "缓存审计 36/36、校准审计 18/18 通过，失败 0。",
        "MPDD 双分支 Gate A 通过：6 类、458 张测试图。",
        "MPDD 完整预测矩阵 9/9 通过。",
        "V2 代码修复后，项目测试 49/49 通过。",
    ]:
        add_bullet(doc, text, status="已完成")

    doc.add_heading("6.2 MPDD 上的动态融合候选发生了什么", level=2)
    add_body(doc, "V2 在 MPDD 上先尝试小范围候选。早期候选有的完全回退为纯视觉，有的只在 seed 0 的一个类别产生很小收益。随后发现像素平滑代码错误地把图像级 out-of-support 条件又用到了像素级，导致图像和像素没有真正独立；该错误已修复并增加回归测试。")
    add_body(doc, "修复后，像素文本辅助确实启动了，但不同 seed 的方向仍不一致。最后使用统一 200 阈值 AUPRO 比较时，pixel_only_w15 相对纯视觉的总体 AUPRO 只增加 0.00324；逐 seed 分别为 -0.00315、+0.01770、-0.00482，只有 1/3 seed 为正，没有通过“至少 2/3 seed 重复”的预设门槛。")
    add_callout(doc, "V2 最终冻结决定", "选择 visual_only_safe_fallback。图像和像素文本权重上限都冻结为 0，动态文本辅助关闭。这个决定保护了稳定的视觉分支，也如实说明：现有文本辅助在 MPDD 上没有证明可重复收益。", fill=PALE_GOLD, color=GOLD)
    add_body(doc, f"参数冻结清单包含 {len(freeze['evidence_files'])} 份代码、测试和实验文件的 SHA256。冻结时明确记录：BTAD 预测、标签和指标均未读取；冻结后才允许一次性进入 BTAD 最终验证。")

    doc.add_heading("6.3 BTAD 最终保持集验证", level=2)
    btad_summary = btad["summary"]
    btad_rows = [
        ("Image AUROC", pct(btad_summary["image_auroc"]["macro_mean"]), pct(btad_summary["image_auroc"]["macro_std"])),
        ("Image AP", pct(btad_summary["image_ap"]["macro_mean"]), pct(btad_summary["image_ap"]["macro_std"])),
        ("Pixel AUROC", pct(btad_summary["pixel_auroc"]["macro_mean"]), pct(btad_summary["pixel_auroc"]["macro_std"])),
        ("Pixel AP", pct(btad_summary["pixel_ap"]["macro_mean"]), pct(btad_summary["pixel_ap"]["macro_std"])),
        ("AUPRO", pct(btad_summary["aupro"]["macro_mean"]), pct(btad_summary["aupro"]["macro_std"])),
    ]
    add_table(doc, ["BTAD 冻结方案", "27 个单元宏平均", "总体标准差"], btad_rows, [3100, 3100, 3160], font_size=9.0,
              aligns=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER], first_col_bold=True)
    add_bullet(doc, "BTAD 双分支 Gate A：3 类、741 张图，审计通过。", status="已完成")
    add_bullet(doc, "3 seeds × 1/2/4-shot：9/9 预测组合、9/9 配对审计通过，失败 0。", status="已完成")
    add_bullet(doc, "最终只评估已冻结方案，没有在 BTAD 比较候选或修改参数。", status="已完成")
    add_body(doc, "最重要的解释：这组 BTAD 数字证明的是“冻结流程和纯视觉安全回退能够稳定运行”，不是“动态文本融合带来了提升”。因为冻结方案的文本权重为 0，结果本质上等于受协议约束的视觉分支。")

    doc.add_heading("6.4 V2 到底算成功还是失败", level=2)
    add_body(doc, "如果从工程和科学规范看，V2 是成功的：它修复了数值饱和，建立了新开发/保持边界，完成参数冻结和独立验证，并在收益不稳定时拒绝选择动态融合。")
    add_body(doc, "如果从“提出一个能够稳定提升指标的新动态融合算法”看，V2 目前没有成功。最终关闭文本辅助意味着核心动态融合增益没有被证明。这一点必须在论文中诚实说明，也决定了后续有两条路线：把论文收紧为安全路由与失效分析，或继续开发更强的 V3/文本分支后再追求性能提升。")

    doc.add_heading("7. 论文、文献和图表材料完成到什么程度", level=1)
    paper_rows = [
        ("中文论文初稿 V0.1", "完成", "包含方法、实验、结果、消融、失败案例和局限性；33 页视觉检查完成。"),
        ("英文 SCI 初稿 V0.2", "完成", "期刊中性单栏格式；摘要、关键词、声明、数据/代码可用性等齐全。"),
        ("英文 SCI 初稿 V0.3", "完成", "重新制作主图，扩充并融入代表性与 2026 年文献，当前共 37 篇参考文献。"),
        ("动态融合分析材料", "完成", "科学分析报告、消融说明、成功/失败案例、热图、路由统计和 Excel 数据包齐全。"),
        ("目标期刊格式", "未完成", "尚未选定具体期刊并迁移到该刊 Word/LaTeX 模板。"),
        ("最终结果表", "未完成", "等待 PromptAD 剩余矩阵和 Gate A 接纳结果后重建。"),
        ("作者信息", "未完成", "作者、单位、ORCID、基金、CRediT、利益冲突和仓库编号仍需补。"),
    ]
    add_table(doc, ["材料", "状态", "说明"], paper_rows, [2300, 1200, 5860], font_size=8.8, header_fill=LIGHT_BLUE, first_col_bold=True)
    add_body(doc, "文献工作已经不再停留在 19 篇早期参考。项目系统筛选了 2024—2026 年少样本视觉、视觉语言、动态路由和校准方法，重点包括 SubspaceAD、FastRef、AnoPLe、AdaptCLIP、DCP-SFR、PGAD、MoECLIP、PAPL 和 TGRF-CLIP 等。")
    add_callout(doc, "论文当前能不能写", "能，而且已经有英文 V0.3。但是它仍是“可继续补实验的完整初稿”，不是可立即投稿的最终稿。最终稿至少要同步 V2/BTAD 新结果、补齐公平基线状态、选定目标期刊并更新结论定位。")

    doc.add_heading("8. 项目当前实时状态与完成度判断", level=1)
    dashboard = [
        ("数据与协议", "约 95%", "四个数据集和统一清单已建立；主要剩复现入口整理。"),
        ("主要基线复现", "约 80%", "VisA 主基线完成；MVTec PromptAD 缺 5 组；ReMP-AD/AdaptCLIP 未过 Gate A。"),
        ("DynamicFusion V1", "100%", "实现、冻结、最终验证、失败分析、消融和可视化均完成。"),
        ("DynamicFusion V2", "约 90%", "代码、MPDD 冻结和 BTAD 独立验证完成；待最终分析和论文同步。"),
        ("论文材料", "约 70%", "中英文初稿和文献完成；最终实验表、目标期刊格式和作者信息待补。"),
        ("整个项目", "约 75%—80%", "已进入收尾期，但新算法稳定增益和近期方法对比仍是关键缺口。"),
    ]
    add_table(doc, ["部分", "管理估算", "判断依据"], dashboard, [2200, 1500, 5660], font_size=8.8, header_fill=LIGHT_BLUE, first_col_bold=True)
    add_body(doc, "上面的百分比只是项目管理估算，不是数学公式。它反映“距离一篇较完整 SCI 投稿稿还差多少”，不是代码文件数量。")

    doc.add_heading("8.1 当前电脑和后台状态", level=2)
    add_bullet(doc, "当前没有本项目的训练或推理进程。", status="当前")
    add_bullet(doc, "BTAD 的 9/9 GPU 预测和最终 CPU 指标已经完成。", status="已完成")
    add_bullet(doc, "PromptAD MVTec 队列仍是 paused_by_schedule，不是正在训练。", status="当前")
    add_bullet(doc, "GPU 监控可能显示桌面或其他程序占用，但没有项目 Python 训练命令。", status="当前")
    add_bullet(doc, "工作区存在大量未提交研究文件和历史脚本；这些属于项目证据，后续应整理，不应直接批量删除。", status="风险")

    doc.add_heading("8.2 现在已经可以确定的科学结论", level=2)
    for text in [
        "AnomalyDINO 是当前已完成实验中的强视觉基线，shot 增加通常能继续提升。",
        "V1 失败的主要原因是少样本正常参考校准饱和和不确定性误判，不只是温度没调好。",
        "图像级排序和像素级定位需要不同的路由逻辑。",
        "文本热图在少数局部案例中有补充价值，但当前证据不支持稳定接管视觉分支。",
        "V2 的安全回退避免了为了平均小收益而选择跨 seed 不稳定的动态方案。",
        "目前不能声称 DynamicFusion 全面优于 AnomalyDINO；论文必须保留负结果和限制。",
    ]:
        add_bullet(doc, text)

    doc.add_heading("9. 后续还要做哪些工作", level=1)
    add_body(doc, "后续工作可以分成“必须完成”和“增强论文”两档。必须完成项决定项目能不能干净收尾；增强项决定论文是否更有竞争力。")

    doc.add_heading("9.1 必须完成的工作", level=2)
    required = [
        ("P0", "BTAD/V2 结果整理", "生成 shot、seed、逐类别表；分析 AUPRO 和 Pixel AP 方差；同步英文论文。", "0.5—1 天", "否"),
        ("P0", "PromptAD MVTec 剩余 5 组", "从 s1/k2 的 capsule 分割断点恢复，完成 s1/k4、s2/k1、s2/k2、s2/k4。", "约 1—3 天", "是，长时"),
        ("P0", "MVTec 最终公平表", "PromptAD 完成后重新生成 3-seed 均值±标准差和逐类别表。", "0.5 天", "否"),
        ("P1", "ReMP-AD Gate A", "完成 manifest、统一 NPZ、单类别最小运行和显存检查。", "1—6 小时起", "是，小到中"),
        ("P1", "AdaptCLIP Gate A", "取得并校验 checkpoint，修复配置，batch=1 做 6 GB 检查。", "0.5—2 天", "是，小到中"),
        ("P0", "论文定位与重写", "把 V2 的视觉回退、BTAD 结果和“不稳定动态增益”融入方法、实验和讨论。", "2—4 天", "否"),
        ("P0", "选定目标期刊", "核对分区、范围、篇幅、图表和模板，再迁移格式。", "0.5—1 天", "否"),
        ("P0", "补作者与声明", "姓名、单位、ORCID、基金、CRediT、利益冲突、数据和代码地址。", "数小时", "否"),
    ]
    add_table(doc, ["优先级", "任务", "要做什么", "预计时间", "GPU"], required,
              [900, 1900, 4060, 1400, 1100], font_size=8.2, header_fill=LIGHT_BLUE,
              aligns=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER])

    doc.add_heading("9.2 推荐但不是最低限度必须的增强实验", level=2)
    enhanced = [
        ("SubspaceAD", "2026 强视觉训练免费方法", "最高优先级新增 Gate A；它直接检验复杂融合是否真的必要。"),
        ("FastRef + AnomalyDINO", "2026 查询原型细化", "高优先级；需核对逐图独立适配和测试统计边界。"),
        ("AnoPLe", "2026 双向多模态提示", "与视觉—文本和图像/像素分离高度相关，资源允许再做。"),
        ("AA-CLIP 或 DLVP-CLIP", "更强文本分支", "二选一做小规模替换实验，判断问题有多少来自 AnomalyCLIP 文本分支。"),
        ("DynamicFusion V3", "真正恢复文本辅助", "只有在新文本分支或新可靠性特征有独立开发数据时才考虑；不能回头用 BTAD 调参。"),
    ]
    add_table(doc, ["候选", "作用", "建议"], enhanced, [2200, 2500, 4660], font_size=8.6, first_col_bold=True)
    add_body(doc, "如果某个新方法通过 Gate A，完整双数据集 1/2/4-shot × 3 seeds 可能需要约 2—7 天 GPU 时间/方法。这里是管理估计，实际受 checkpoint、依赖、分辨率和 6 GB 显存影响很大。")

    doc.add_heading("9.3 后续验收标准", level=2)
    for text in [
        "每个完整 MVTec 组合必须有 15 类、1,725 样本、零 schema 错误。",
        "每个完整 VisA 组合必须有 12 类、2,162 样本、零 schema 错误。",
        "每个新方法先过 Gate A，失败时保留日志，不直接进入长矩阵。",
        "所有主表数字必须由逐运行报告脚本生成，不能手工填。",
        "目标域调参、源域训练、单查询适配、测试集统计使用情况必须单独标注。",
        "BTAD 已经是保持集，后续只能报告和分析，不能用于 V2 回头调参。",
    ]:
        add_bullet(doc, text)

    doc.add_heading("10. 推荐的项目收尾路线", level=1)
    doc.add_heading("路线 A：尽快形成可投稿版本", level=2)
    add_body(doc, "适合希望尽快完成一篇客观、规范的 SCI 稿。重点不再追求“动态融合一定提高”，而是把论文主线收紧为：少样本异构分数校准为什么会失败，如何建立无泄漏、安全回退和独立验证流程。")
    for text in [
        "完成 PromptAD MVTec 剩余 5 组。",
        "ReMP-AD 和 AdaptCLIP 至少完成 Gate A；如果失败，给出可复核原因。",
        "整理 BTAD 逐类别结果和 V2 冻结链。",
        "重写英文 V0.3，突出 V1 失败机制、V2 安全拒绝和 BTAD holdout。",
        "选择目标期刊，完成最终图表、格式、作者信息和语言检查。",
    ]:
        add_number(doc, text)

    doc.add_heading("路线 B：继续追求更强算法贡献", level=2)
    add_body(doc, "适合希望让论文主线仍然是“动态融合带来稳定提升”。这条路线风险和耗时都更高。")
    for text in [
        "至少完成 SubspaceAD、FastRef 和一个更强文本分支 Gate A。",
        "在新的开发数据上设计 V3，不能继续使用 BTAD 调参。",
        "重新设定可重复性门槛，并准备新的独立保持集或交叉数据验证。",
        "只有当至少 2/3 seed 的定位收益重复、图像指标无材料性退化时，才把动态融合恢复为论文主结果。",
        "如果仍不能重复，及时停止，回到路线 A，不无限追加实验。",
    ]:
        add_number(doc, text)

    doc.add_heading("10.1 我的客观建议", level=2)
    add_callout(doc, "建议", "先走路线 A，把已经完成的扎实工程、严格协议、V1 失效链、V2 安全冻结和 BTAD 独立验证整理成一篇诚实、可复核的论文；同时并行做 SubspaceAD/FastRef Gate A。只有新证据非常明确时，再把路线升级为更强算法稿。不要为了追求一个漂亮平均数，继续在已经看过的测试集上调动态融合。", fill=PALE_GREEN, color=GREEN)

    doc.add_heading("11. 最终判断", level=1)
    add_body(doc, "项目已经跨过了最难的“从零到可复现系统”阶段，也完成了完整的 DynamicFusion V1 和较规范的 V2 冻结验证。当前手头材料足以形成一篇结构完整的英文论文初稿，并且已经有 V0.3。")
    add_body(doc, "但项目的核心新算法结论还需要谨慎：V1 明显低于强视觉分支；V2 为避免不稳定收益，最终选择纯视觉安全回退。因此，现在最可靠的贡献是严格少样本协议、异构分数校准失效分析、无泄漏路由设计、安全回退和负结果的完整证据链，而不是稳定性能领先。")
    add_body(doc, "后续只要完成 PromptAD 剩余矩阵、处理 ReMP-AD/AdaptCLIP Gate A、整理 V2/BTAD 结果并重写英文稿，就可以进入目标期刊格式化和投稿准备。若要显著提高论文算法竞争力，则还需至少一个 2026 强基线和一个新的、独立验证的 V3 方案。")

    doc.add_heading("附录 A：当前权威证据位置", level=1)
    evidence = [
        ("当前待办", "NEXT_ACTIONS.md"),
        ("项目状态", "PROJECT_STATUS.md 最新章节"),
        ("VisA 主表", "experiments/summaries/visa_baseline_main_table_20260803.csv"),
        ("MVTec 完整性矩阵", "experiments/summaries/mvtec_method_seed_shot_completeness_20260809.csv"),
        ("V1 科学分析", "docs/dynamic_fusion_scientific_analysis_20260809.md"),
        ("V2 参数冻结", "experiments/dynamic_fusion/v2/parameter_freeze/manifest.json"),
        ("MPDD 完整矩阵", "experiments/dynamic_fusion/v2/mpdd_prediction_matrix/runtime/status.json"),
        ("BTAD 完整矩阵", "experiments/dynamic_fusion/v2/btad_prediction_matrix/runtime/status.json"),
        ("BTAD 最终指标", "experiments/dynamic_fusion/v2/btad_frozen_evaluation/report.json / report.csv"),
        ("PromptAD 队列", "outputs/logs/promptad_mvtec_resumable_queue/status.json"),
        ("英文初稿", "outputs/paper_draft_20260810/Leakage-Safe_Uncertainty_Routing_English_SCI_Draft_V0.3.docx"),
        ("文献与验证计划", "docs/representative_literature_and_validation_plan_20260810.md"),
    ]
    add_table(doc, ["内容", "文件"], evidence, [2600, 6760], font_size=8.7, first_col_bold=True)

    note = doc.add_paragraph(style="Small Note")
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run("说明：本报告依据 2026-08-11 磁盘上的状态文件、统一结果、冻结清单和实验审计生成。历史章节仅用于说明过程，发生冲突时以最新运行报告和冻结证据为准。")
    set_run_font(run, size=9, color=GRAY)

    doc.core_properties.title = "少样本工业异常检测项目全程进展与后续工作说明"
    doc.core_properties.subject = "项目从立项到 DynamicFusion V2 与 BTAD 冻结验证的完整中文说明"
    doc.core_properties.author = "Research Project Team"
    doc.core_properties.keywords = "industrial anomaly detection, few-shot, DynamicFusion, project status"
    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build()

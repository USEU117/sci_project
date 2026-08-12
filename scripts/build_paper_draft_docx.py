from __future__ import annotations

import csv
import json
import math
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "outputs" / "paper_draft_20260810"
ASSET_DIR = OUT_DIR / "assets"
DOCX_PATH = OUT_DIR / "基于不确定性路由的少样本工业异常视觉语言证据融合方法_论文初稿_V0.1.docx"

DOC_SKILL = Path(
    r"C:\Users\lynle\.codex\plugins\cache\openai-primary-runtime\documents"
    r"\26.805.11740\skills\documents"
)
sys.path.insert(0, str(DOC_SKILL / "scripts"))
from table_geometry import apply_table_geometry, column_widths_from_weights, section_content_width_dxa  # noqa: E402


BLUE = "2E74B5"
DARK_BLUE = "17365D"
LIGHT_BLUE = "DCE6F1"
PALE_BLUE = "EFF5FB"
LIGHT_GRAY = "F2F2F2"
MID_GRAY = "D9E2F3"
TEXT = RGBColor(31, 31, 31)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, value: str, *, bold: bool = False, color: str | None = None, size: float = 8.5) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(str(value))
    r.bold = bold
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant = OxmlElement("w:cantSplit")
    tr_pr.append(cant)


def add_page_field(paragraph, *, fmt: str | None = None) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, fld_end])
    if fmt:
        pg_num_type = OxmlElement("w:pgNumType")
        pg_num_type.set(qn("w:fmt"), fmt)
        paragraph._p.getparent().append(pg_num_type)


def set_page_numbering(section, start: int, fmt: str = "decimal") -> None:
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num_type)
    pg_num_type.set(qn("w:start"), str(start))
    pg_num_type.set(qn("w:fmt"), fmt)


def add_toc_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "目录将在 Word 中自动更新"
    separate.append(text)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, end])


def set_east_asia_font(run, name: str) -> None:
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:eastAsia"), name)


def add_body(doc: Document, text: str, *, bold_lead: str | None = None, indent: bool = True) -> None:
    p = doc.add_paragraph(style="Body Text")
    if not indent:
        p.paragraph_format.first_line_indent = Twips(0)
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        r1.bold = True
        set_east_asia_font(r1, "黑体")
        r2 = p.add_run(text[len(bold_lead):])
        set_east_asia_font(r2, "宋体")
    else:
        r = p.add_run(text)
        set_east_asia_font(r, "宋体")


def add_bullet(doc: Document, text: str, *, level: int = 0) -> None:
    style = "List Bullet" if level == 0 else "List Bullet 2"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    set_east_asia_font(r, "宋体")


def add_number(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    set_east_asia_font(r, "宋体")


def add_manual_number(doc: Document, number: int, text: str) -> None:
    p = doc.add_paragraph(style="Body Text")
    p.paragraph_format.left_indent = Pt(24)
    p.paragraph_format.first_line_indent = Pt(-24)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(f"{number}.  {text}")
    set_east_asia_font(r, "宋体")


def add_equation(doc: Document, text: str, number: str) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = ""
    p = table.cell(0, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.name = "Cambria Math"
    r.font.size = Pt(11)
    table.cell(0, 1).text = ""
    p2 = table.cell(0, 1).paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r2 = p2.add_run(number)
    r2.font.name = "Times New Roman"
    r2.font.size = Pt(10)
    widths = column_widths_from_weights([9, 1], section_content_width_dxa(doc.sections[-1]))
    apply_table_geometry(table, widths, indent_dxa=0, cell_margins_dxa={"top": 20, "bottom": 20, "start": 0, "end": 0})
    for cell in table.rows[0].cells:
        tc_pr = cell._tc.get_or_add_tcPr()
        borders = OxmlElement("w:tcBorders")
        for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), "nil")
            borders.append(el)
        tc_pr.append(borders)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_east_asia_font(r, "宋体")


def add_figure(doc: Document, path: Path, caption: str, *, width: float = 6.25) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)


def add_table(
    doc: Document,
    caption: str,
    headers: list[str],
    rows: list[list[str]],
    weights: list[float],
    *,
    footnote: str | None = None,
    font_size: float = 8.2,
    cell_margins_dxa: dict[str, int] | None = None,
) -> None:
    add_caption(doc, caption)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = 0
    for idx, header in enumerate(headers):
        set_cell_text(table.cell(0, idx), header, bold=True, color="FFFFFF", size=font_size)
        set_cell_shading(table.cell(0, idx), BLUE)
    set_repeat_table_header(table.rows[0])
    for row_idx, row in enumerate(rows, start=1):
        cells = table.add_row().cells
        for col_idx, value in enumerate(row):
            set_cell_text(cells[col_idx], value, size=font_size)
            if row_idx % 2 == 0:
                set_cell_shading(cells[col_idx], PALE_BLUE)
        set_cant_split(table.rows[-1])
    widths = column_widths_from_weights(weights, section_content_width_dxa(doc.sections[-1]))
    apply_table_geometry(table, widths, indent_dxa=0, cell_margins_dxa=cell_margins_dxa)
    if footnote:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(footnote)
        r.italic = True
        r.font.size = Pt(8.5)
        set_east_asia_font(r, "宋体")


def add_callout(doc: Document, title: str, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_BLUE)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    set_east_asia_font(r, "黑体")
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(text)
    set_east_asia_font(r2, "宋体")
    widths = [section_content_width_dxa(doc.sections[-1])]
    apply_table_geometry(table, widths, indent_dxa=0, cell_margins_dxa={"top": 120, "bottom": 120, "start": 160, "end": 160})


def create_architecture_figure(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    width, height = 2860, 1410
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    font_path = r"C:\Windows\Fonts\simhei.ttf"
    title_font = ImageFont.truetype(font_path, 54)
    body_font = ImageFont.truetype(font_path, 33)
    small_font = ImageFont.truetype(font_path, 28)

    def box(x, y, w, h, text, fill, edge="#17365D", font=body_font):
        draw.rounded_rectangle((x, y, x + w, y + h), radius=20, fill=fill, outline=edge, width=4)
        bbox = draw.multiline_textbbox((0, 0), text, font=font, spacing=10, align="center")
        tw, th = bbox[2] - bbox[0], bbox[3] - bbox[1]
        draw.multiline_text((x + (w - tw) / 2, y + (h - th) / 2), text, font=font, fill="#1F1F1F", spacing=10, align="center")

    def arrow(x1, y1, x2, y2, color="#4F81BD"):
        draw.line((x1, y1, x2, y2), fill=color, width=7)
        angle = math.atan2(y2 - y1, x2 - x1)
        size = 22
        p1 = (x2 - size * math.cos(angle - math.pi / 6), y2 - size * math.sin(angle - math.pi / 6))
        p2 = (x2 - size * math.cos(angle + math.pi / 6), y2 - size * math.sin(angle + math.pi / 6))
        draw.polygon([(x2, y2), p1, p2], fill=color)

    title = "冻结分支预测 → 正常参考校准 → 不确定性路由 → 融合输出"
    tb = draw.textbbox((0, 0), title, font=title_font)
    draw.text(((width - (tb[2] - tb[0])) / 2, 45), title, font=title_font, fill="#17365D")

    box(40, 355, 330, 170, "待测图像", "#EAF2F8")
    box(40, 980, 330, 190, "K 张正常\n参考图", "#E8F5E9", edge="#548235")
    box(475, 285, 430, 185, "视觉分支\nAnomalyDINO", "#D9EAF7")
    box(475, 555, 430, 185, "文本分支\nAnomalyCLIP", "#FCE4D6", edge="#C65911")
    box(475, 950, 430, 220, "正常参考预测\n确定性多视图", "#E2F0D9", edge="#548235")
    box(1035, 410, 430, 240, "统一缓存与对齐\n样本 ID、标签\n像素图尺寸", "#FFF2CC", edge="#BF9000")
    box(1035, 950, 430, 220, "稳健校准\n中位数 + MAD\n像素 q99", "#E2F0D9", edge="#548235")
    box(1595, 410, 430, 240, "可靠性特征\n二元熵、分支冲突\n空间集中度", "#E4DFEC", edge="#7030A0")
    box(1595, 930, 430, 260, "信息边界检查\n不读取测试标签/真值\n不使用测试集统计量", "#F4CCCC", edge="#A61C00", font=small_font)
    box(2150, 390, 350, 290, "双层路由\n图像温度 0.50\n像素温度 0.20", "#D9EAD3", edge="#38761D", font=small_font)
    box(2580, 350, 240, 170, "图像\n异常分数", "#D9EAF7", font=small_font)
    box(2580, 610, 240, 170, "像素\n异常图", "#FCE4D6", edge="#C65911", font=small_font)

    arrow(370, 440, 475, 378)
    arrow(370, 440, 475, 648)
    arrow(370, 1075, 475, 1060, color="#548235")
    arrow(905, 378, 1035, 480)
    arrow(905, 648, 1035, 580)
    arrow(905, 1060, 1035, 1060, color="#548235")
    arrow(1250, 950, 1250, 650, color="#548235")
    arrow(1465, 530, 1595, 530)
    arrow(2025, 530, 2150, 530)
    arrow(1810, 930, 1810, 650, color="#A61C00")
    arrow(2500, 490, 2580, 435)
    arrow(2500, 580, 2580, 690)

    footer = "核心审计原则：路由器只能看模型预测与正常参考信息，不能看测试真值"
    fb = draw.textbbox((0, 0), footer, font=body_font)
    draw.text(((width - (fb[2] - fb[0])) / 2, 1300), footer, font=body_font, fill="#A61C00")
    image.save(path, optimize=True)


def configure_document(doc: Document) -> None:
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(1.0)
    sec.bottom_margin = Inches(1.0)
    sec.left_margin = Inches(1.0)
    sec.right_margin = Inches(1.0)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)
    sec.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)
    normal.font.color.rgb = TEXT
    normal.paragraph_format.line_spacing = 1.35
    normal.paragraph_format.space_after = Pt(8)

    body = styles["Body Text"]
    body.font.name = "Times New Roman"
    body._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    body.font.size = Pt(11)
    body.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    body.paragraph_format.first_line_indent = Pt(22)
    body.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    body.paragraph_format.line_spacing = 1.35
    body.paragraph_format.space_after = Pt(7)
    body.paragraph_format.widow_control = True

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        st = styles[name]
        st.font.name = "Arial"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True
    styles["Heading 1"].paragraph_format.page_break_before = True

    styles["Title"].font.name = "Arial"
    styles["Title"]._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    styles["Title"].font.size = Pt(28)
    styles["Title"].font.bold = True
    styles["Title"].font.color.rgb = RGBColor.from_string(DARK_BLUE)

    caption = styles["Caption"]
    caption.font.name = "Times New Roman"
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    caption.font.size = Pt(9.5)
    caption.font.color.rgb = RGBColor(64, 64, 64)
    caption.paragraph_format.space_after = Pt(7)

    if "Front Heading" not in styles:
        fh = styles.add_style("Front Heading", WD_STYLE_TYPE.PARAGRAPH)
    else:
        fh = styles["Front Heading"]
    fh.font.name = "Arial"
    fh._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    fh.font.size = Pt(17)
    fh.font.bold = True
    fh.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    fh.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fh.paragraph_format.space_before = Pt(12)
    fh.paragraph_format.space_after = Pt(14)

    for name, size in [("TOC 1", 10.0), ("TOC 2", 9.2), ("TOC 3", 8.8)]:
        st = styles[name] if name in styles else styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        st.font.name = "Times New Roman"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        st.font.size = Pt(size)
        st.paragraph_format.space_before = Pt(0)
        st.paragraph_format.space_after = Pt(1)
        st.paragraph_format.line_spacing = 1.0

    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def configure_running_section(section, title: str, *, page_start: int, page_fmt: str = "decimal") -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    hp = section.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = hp.add_run(title)
    hr.font.name = "Times New Roman"
    set_east_asia_font(hr, "宋体")
    hr.font.size = Pt(8.5)
    hr.font.color.rgb = RGBColor(128, 128, 128)
    fp = section.footer.paragraphs[0]
    add_page_field(fp)
    set_page_numbering(section, page_start, page_fmt)


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def pct(value: str) -> str:
    return f"{float(value) * 100:.2f}"


def build_doc() -> Path:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    architecture = ASSET_DIR / "dynamic_fusion_architecture.png"
    create_architecture_figure(architecture)

    visa_rows = read_csv(ROOT / "experiments" / "summaries" / "visa_baseline_main_table_20260803.csv")
    mvtec_rows = read_csv(ROOT / "experiments" / "summaries" / "mvtec_paper_main_table_template_20260809.csv")
    ablation_rows = read_csv(ROOT / "experiments" / "summaries" / "dynamic_fusion_scientific_analysis_20260809" / "ablation_summary.csv")
    with (ROOT / "outputs" / "dynamic_fusion" / "final_validation" / "summary.json").open("r", encoding="utf-8") as f:
        visa_dynamic = json.load(f)

    fig_dir = ROOT / "outputs" / "dynamic_fusion" / "figures" / "20260809_scientific_analysis"
    doc = Document()
    configure_document(doc)
    props = doc.core_properties
    props.title = "基于不确定性路由的少样本工业异常视觉—语言证据融合方法"
    props.subject = "论文初稿 V0.1；统一少样本基线、动态融合、消融与失效分析"
    props.author = ""
    props.keywords = "工业异常检测, 少样本学习, 视觉语言模型, 动态融合, 不确定性路由"

    # Cover.
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(28)
    r = p.add_run("论文初稿")
    r.font.name = "Arial"
    set_east_asia_font(r, "黑体")
    r.font.size = Pt(18)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(BLUE)
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(58)
    p.paragraph_format.space_after = Pt(16)
    p.add_run("基于不确定性路由的少样本工业异常\n视觉—语言证据融合方法")
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Uncertainty-Routed Vision–Language Evidence Fusion\nfor Few-Shot Industrial Anomaly Detection")
    r2.font.name = "Times New Roman"
    r2.font.size = Pt(15)
    r2.font.italic = True
    r2.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(95)
    for line in ["作者：________________", "学校 / 学院：________________", "专业：________________", "指导教师：________________"]:
        rp = p3.add_run(line + "\n")
        rp.font.size = Pt(12)
        set_east_asia_font(rp, "宋体")
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.paragraph_format.space_before = Pt(35)
    rr = p4.add_run("初稿版本：V0.1\n日期：2026 年 8 月 10 日")
    rr.font.size = Pt(10.5)
    set_east_asia_font(rr, "宋体")

    front = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_running_section(front, "论文初稿 V0.1", page_start=1, page_fmt="lowerRoman")

    # Chinese abstract.
    doc.add_paragraph("摘  要", style="Front Heading")
    for para in [
        "工业视觉异常检测需要在缺陷样本稀少、异常类型未知且类别差异明显的条件下，同时判断整张图像是否异常，并定位异常区域。传统视觉方法善于利用局部纹理和结构信息，视觉—语言方法则能借助文本语义提供跨类别先验。两类方法具有互补可能，但它们的分数范围、空间分辨率和可靠性含义并不一致，直接平均容易破坏强分支原有的排序能力。",
        "本文围绕 1/2/4-shot 工业异常检测，首先在 MVTec AD 和 VisA 上建立统一、可审计的复现协议。协议固定 3 个随机种子，使用嵌套正常参考集，统一输出图像异常分数、像素异常图和样本标识，并分别复现 PatchCore、WinCLIP+、AnomalyDINO 和 PromptAD。随后，本文以 AnomalyDINO 为视觉分支、AnomalyCLIP 为文本分支，设计图像级与像素级分离的不确定性路由框架。路由器只使用冻结预测、正常参考图以及分支不确定性和一致性特征，不读取测试标签、测试掩膜或测试集整体统计量。",
        "实验表明，在 VisA 基线中，AnomalyDINO 在 1/2/4-shot 下的图像 AUROC 分别达到 89.40%、91.40% 和 92.58%，是当前已完成方法中的最强视觉基线。双温度路由在 VisA seed 0 的开发消融中，对 K=2 和 K=4 的 AUPRO 相比最佳固定权重分别提高 4.14 和 8.82 个百分点，说明图像判断与像素定位需要不同的路由强度。然而，冻结后的 V1 在 MVTec 上未超过原始 AnomalyDINO：其图像 AUROC 在 1/2/4-shot 下分别低 16.29、10.49 和 7.94 个百分点。进一步分析发现，少量正常参考图估计出的中位数/MAD 尺度过小，使 sigmoid 校准在 MVTec 上平均有 99.99% 的视觉分数进入不小于 0.999 的饱和区域；二元熵又把饱和值误判为高置信，从而导致路由权重无法可靠反映分支优劣。",
        "本文的主要价值包括：建立统一少样本复现与结果审计流程；实现不使用测试真值的图像—像素双层路由接口；通过跨数据集、跨 shot、逐类别和案例级证据定位动态融合失效机制。结果说明，视觉—语言融合的关键不只是设计权重公式，还要保证分数校准保留排序，并让不确定性能够识别“超出参考分布”的预测。本文据此提出排序保持校准、饱和保护、安全回退和重新划分开发/最终验证集等后续方向。",
    ]:
        add_body(doc, para)
    p = doc.add_paragraph()
    r = p.add_run("关键词：")
    r.bold = True
    set_east_asia_font(r, "黑体")
    r2 = p.add_run("工业异常检测；少样本学习；AnomalyDINO；AnomalyCLIP；视觉—语言融合；不确定性路由；缺陷定位")
    set_east_asia_font(r2, "宋体")

    doc.add_page_break()
    doc.add_paragraph("ABSTRACT", style="Front Heading")
    english_abstract = [
        "Industrial visual anomaly detection must identify defective images and localize anomalous regions when anomalous training samples are unavailable and only a few normal references are provided. Vision-only methods capture fine-grained structural cues, while vision–language methods introduce transferable semantic priors. Their outputs, however, differ in scale, spatial resolution, and reliability semantics; naïve averaging can therefore damage the ranking produced by a strong branch.",
        "This study first establishes an auditable 1/2/4-shot protocol on MVTec AD and VisA with three deterministic seeds, nested normal-reference sets, a unified prediction schema, and common metrics. PatchCore, WinCLIP+, AnomalyDINO, and PromptAD are reproduced under this protocol. A two-level uncertainty router is then developed using AnomalyDINO as the visual branch and AnomalyCLIP as the text-guided branch. The router consumes frozen predictions, normal-reference calibration, uncertainty, and consistency features only; test labels, masks, and aggregate test statistics are explicitly forbidden.",
        "On VisA, AnomalyDINO obtains image AUROC values of 89.40%, 91.40%, and 92.58% for 1-, 2-, and 4-shot settings. A split-temperature router improves AUPRO over the best fixed-weight control by 4.14 and 8.82 percentage points for K=2 and K=4 on the seed-0 development set. Nevertheless, frozen V1 does not outperform raw AnomalyDINO on MVTec: image AUROC decreases by 16.29, 10.49, and 7.94 percentage points. The primary cause is calibration saturation. The median/MAD sigmoid mapping pushes 99.99% of visual test scores above 0.999 on average, while binary entropy incorrectly treats such saturated values as highly confident.",
        "The contribution is therefore both methodological and diagnostic: a reproducible few-shot benchmark, a leakage-safe image/pixel routing interface, and an evidence-based analysis showing that reliable fusion requires rank-preserving calibration, saturation awareness, and safe fallback policies. These findings define a defensible V2 research path without tuning on already inspected final-validation data.",
    ]
    for para in english_abstract:
        p = doc.add_paragraph(style="Body Text")
        p.paragraph_format.first_line_indent = Pt(18)
        r = p.add_run(para)
        r.font.name = "Times New Roman"
        r.font.size = Pt(10.5)
    p = doc.add_paragraph()
    r = p.add_run("Keywords: ")
    r.bold = True
    r.font.name = "Times New Roman"
    r2 = p.add_run("industrial anomaly detection; few-shot learning; vision–language fusion; uncertainty routing; defect localization")
    r2.font.name = "Times New Roman"

    doc.add_page_break()
    doc.add_paragraph("目  录", style="Front Heading")
    toc_p = doc.add_paragraph()
    add_toc_field(toc_p)
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rn = note.add_run("提示：首次在 Microsoft Word 中打开后，选择目录并按 F9 可刷新页码。")
    rn.font.size = Pt(9)
    rn.font.color.rgb = RGBColor(128, 128, 128)
    set_east_asia_font(rn, "宋体")

    body_sec = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_running_section(body_sec, "基于不确定性路由的少样本工业异常视觉—语言证据融合方法", page_start=1, page_fmt="decimal")

    # Chapter 1.
    doc.add_heading("1 绪论", level=1)
    doc.add_heading("1.1 研究背景", level=2)
    add_body(doc, "工业产品表面缺陷通常具有数量少、类型变化快和标注成本高等特点。真实生产线上，大多数样本是正常品，已经出现的缺陷也不能覆盖未来可能发生的异常。因此，监督式分类或分割方法虽然在固定缺陷上效果较好，却很难直接适应新产线、新产品和未知缺陷。工业异常检测把问题改写为“只学习正常状态，再识别偏离正常状态的样本”，由此降低对异常标注的依赖。MVTec AD 提供 15 个物体和纹理类别以及像素级真值[1]；VisA 进一步提供 12 个物体类别和更丰富的工业场景[2]，二者已经成为工业异常检测中常用的公开基准。")
    add_body(doc, "少样本异常检测比常规无监督异常检测更严格：每个目标类别只给出 K 张正常参考图，其中 K 通常为 1、2 或 4。方法既要从极少样本中建立正常外观，又要避免把参考图的偶然姿态、光照或背景当成类别规则。与此同时，工业任务不仅需要图像级判断，还需要像素级定位。前者关注正常图和异常图的全局排序，后者关注异常区域与背景之间的空间对比，两种目标并不完全一致。")
    add_body(doc, "预训练视觉模型和视觉—语言模型为这一问题提供了两条技术路线。PatchCore 使用预训练特征建立正常 patch 记忆库[3]；WinCLIP/WinCLIP+ 利用 CLIP 的文本语义和窗口特征开展零样本或少样本检测[4][5]；AnomalyDINO 使用 DINOv2 的视觉 patch 特征进行近邻匹配[7][8]；PromptAD 则在目标类别的正常样本上学习正常与异常提示[9]。这些方法的信息来源不同，可能在某些缺陷上互补，但也可能因为分数尺度不一致而相互干扰。")

    doc.add_heading("1.2 研究问题与主要挑战", level=2)
    add_body(doc, "本文研究的问题是：在不使用目标测试集真值的前提下，能否根据视觉分支与文本分支的可靠程度，为每张图像、每个像素动态分配融合权重，从而兼顾图像识别和缺陷定位。该问题包含四个直接挑战。")
    for item in [
        "分支输出不可直接比较。AnomalyDINO 的近邻距离与 AnomalyCLIP 的语义相似度具有不同量纲，必须先校准到可比较范围。",
        "少量正常参考难以估计稳定分布。K 很小时，中位数、MAD 或分位数容易被单个视角影响，进而造成校准饱和。",
        "图像级与像素级目标不同。同一权重可能保持图像排序，却把像素热图变得过于平滑；也可能改善局部定位，却破坏整图排序。",
        "实验容易发生信息泄漏。如果根据测试标签选择阈值、融合权重或参数，最终结果将失去独立验证意义。",
    ]:
        add_bullet(doc, item)

    doc.add_heading("1.3 研究目标与工作内容", level=2)
    add_body(doc, "本文不把“提出一个融合公式”作为唯一目标，而是同时解决实验协议、工程接口、结果审计和科学解释四个层面的问题。第一，固定 MVTec AD 和 VisA 的 1/2/4-shot、3-seed 嵌套抽样，建立统一预测格式和统一指标。第二，复现多种视觉与视觉—语言基线，并明确不同方法是否使用目标正常样本调参。第三，设计图像级与像素级分离的不确定性路由器。第四，在冻结参数后进行跨数据集验证，并对不符合预期的结果做机制分析。")

    doc.add_heading("1.4 本文主要贡献", level=2)
    add_number(doc, "建立可审计的少样本工业异常检测复现流程。所有方法读取同一份嵌套 manifest，输出统一 NPZ 预测文件，并通过类别数、样本数、标签方向、像素图尺寸和 SHA256 来源检查。")
    add_number(doc, "实现不依赖测试真值的双层动态融合框架。该框架对视觉与文本分支分别校准，计算图像和像素不确定性，再生成连续融合权重和可解释路由标签。")
    add_number(doc, "完成系统消融与失效分析。实验不仅比较动态权重与固定权重，还量化校准饱和、权重—分支优势相关性、类别差异和排序变化，从而解释 V1 未能超过强视觉分支的原因。")
    add_number(doc, "给出边界清楚的后续研究路线。本文不使用已经查看的最终验证结果继续调参，而是把排序保持校准、饱和保护和安全回退作为 V2 假设，并要求重新建立开发/验证边界。")
    add_callout(doc, "论文表述边界", "本文可以报告已审计结果和局部收益，但不能宣称动态融合全面优于 AnomalyDINO，也不能把尚未通过 Gate A 的 ReMP-AD、AdaptCLIP 写成已完成对比。")

    doc.add_heading("1.5 论文结构", level=2)
    add_body(doc, "第 2 章介绍工业异常检测、视觉—语言异常检测和多模态融合相关工作；第 3 章给出任务定义和统一实验协议；第 4 章介绍动态融合方法；第 5 章说明数据集、实现环境和评价指标；第 6 章汇报基线复现结果；第 7 章汇报动态融合与消融实验；第 8 章分析失败原因、适用边界和有效性威胁；第 9 章总结全文并给出后续工作。")

    # Chapter 2.
    doc.add_heading("2 相关工作", level=1)
    doc.add_heading("2.1 基于视觉特征的工业异常检测", level=2)
    add_body(doc, "视觉特征方法通常使用在大规模自然图像上预训练的网络提取特征，再学习正常样本在特征空间中的分布。PaDiM 对每个空间位置建立多元高斯分布[12]；PatchCore 把正常图像的局部 patch 特征存入记忆库，并用近邻距离作为异常分数[3]。这类方法结构直观，对纹理、边缘和局部结构敏感，但在少样本条件下，正常状态覆盖不充分，参考图的姿态变化会直接影响记忆库。")
    add_body(doc, "DINOv2 通过大规模自监督学习获得通用视觉特征[7]。AnomalyDINO 将其用于少样本异常检测：对查询 patch 与正常参考 patch 做近邻匹配，再形成图像分数和像素热图[8]。其优点是不需要在目标数据集上训练骨干网络，且视觉 patch 对细粒度结构较敏感。本项目的统一复现也表明，AnomalyDINO 是当前已完成矩阵中的最强视觉分支，因此动态融合必须至少保留其排序信息，才有可能获得总体提升。")

    doc.add_heading("2.2 基于视觉—语言模型的异常检测", level=2)
    add_body(doc, "CLIP 通过大规模图文对比学习把图像和文本映射到同一表示空间[5]。在异常检测中，可以把“正常产品”和“损坏产品”等文本描述作为语义锚点。WinCLIP 将不同尺度窗口与文本提示结合，并在 WinCLIP+ 中加入少量正常参考图[4]。AnomalyCLIP 学习与具体物体类别无关的正常/异常提示，以提高跨数据集的零样本检测能力[6]。这类方法不完全依赖目标类别的正常图，在视觉参考不足时可能提供补充；但文本分支也可能对复杂纹理、细小缺陷或类别特有外观不敏感。")
    add_body(doc, "PromptAD 使用目标类别正常样本学习提示，并通过语义拼接构造异常提示[9]。由于它在目标正常样本上发生参数更新，本文将其标记为 target_normal_tuning=true，与完全冻结的推理方法分开解释。ReMP-AD 引入检索增强的多模态提示融合[10]；AdaptCLIP 通过视觉、文本和 prompt-query 适配器实现跨域零/少样本异常检测[11]。二者与本文研究方向密切相关，但当前项目尚未完成其 Gate A，因此只在相关工作中介绍，不进入正式结果排名。")

    doc.add_heading("2.3 不确定性、校准与动态融合", level=2)
    add_body(doc, "多分支融合常见做法包括固定加权、规则路由和学习式门控。固定加权容易复现，但无法随样本变化；学习式门控更灵活，却需要可靠监督。工业异常检测通常没有可用于训练路由器的异常标签，因此本文选择正常参考校准加无监督不确定性路由。关键假设是：校准后的概率越接近 0.5，分支越不确定；越接近 0 或 1，分支越确定。后续实验表明，这个假设只有在校准合理时成立。如果 sigmoid 因参考尺度过小而饱和，极端概率可能只表示“超出校准范围”，并不表示预测正确。")
    add_body(doc, "与简单平均相比，本文还把图像级权重和像素级权重分开。图像级权重控制整图异常分数，直接影响 AUROC 的全局排序；像素级权重控制热图中每个位置的分支贡献，直接影响异常区域与背景的对比。因此，两者需要不同的路由温度，并应分别验证。")

    doc.add_heading("2.4 现有方法比较与研究空缺", level=2)
    add_body(doc, "现有少样本工业异常检测方法大致可以按信息来源分为正常视觉参考、文本语义提示和二者联合三类。视觉方法容易解释“与正常外观相差多少”，但对参考覆盖范围敏感；文本方法能够提供跨类别先验，但对细粒度纹理和具体产品结构的刻画有限；联合方法理论上可以互补，却必须回答“什么时候应该相信哪条分支”。")
    add_table(
        doc,
        "表 2-1 代表性技术路线与本文关注点",
        ["技术路线", "代表方法", "主要证据", "主要风险", "与本文关系"],
        [
            ["视觉记忆/近邻", "PatchCore、AnomalyDINO", "正常 patch 特征", "参考覆盖不足", "提供强视觉分支"],
            ["视觉—语言提示", "WinCLIP、AnomalyCLIP", "文本正常/异常语义", "局部结构不敏感", "提供文本分支"],
            ["目标域提示学习", "PromptAD", "正常图 + 可学习提示", "需要目标域调参", "公平性对照"],
            ["多模态融合", "ReMP-AD 等", "检索与视觉—语言先验", "融合可靠性难验证", "最接近的对比方向"],
        ],
        [1.5, 1.8, 1.9, 1.8, 1.8],
        font_size=7.7,
        cell_margins_dxa={"top": 55, "bottom": 55, "start": 80, "end": 80},
    )
    add_body(doc, "本文关注的研究空缺不是“是否能够把两个分数相加”，而是缺少一种在无异常监督条件下仍可审计的可靠性判断：它既要允许文本证据在视觉分支确实不足时参与，又要保护已经很强的视觉排序；同时还要把整图排序和像素定位作为两个不同目标。后文的 V1 先给出一个可复现实现，再通过失败证据明确其不足。")

    # Chapter 3.
    doc.add_heading("3 任务定义与统一实验协议", level=1)
    doc.add_heading("3.1 少样本异常检测任务", level=2)
    add_body(doc, "设目标类别 c 具有 K 张正常参考图 R_c^K={r_1,…,r_K}，其中 K∈{1,2,4}。测试集合 Q_c 同时包含正常图像和异常图像，但测试标签只允许在最终评价时使用。给定测试图像 x，模型输出图像异常分数 s(x) 和像素异常图 M(x)。s(x) 越大表示整张图越可能异常；M(x) 中较大的位置表示更可能属于缺陷区域。")
    add_table(
        doc,
        "表 3-1 主要符号说明",
        ["符号", "含义", "本文取值或形状"],
        [
            ["K", "每个类别的正常参考图数量", "1、2、4"],
            ["R_c^K", "类别 c 的正常参考集合", "嵌套抽样"],
            ["s_v, s_t", "视觉/文本分支图像分数", "N 维"],
            ["M_v, M_t", "视觉/文本分支像素图", "N×H×W"],
            ["u_v, u_t", "两分支不确定性", "[0,1]"],
            ["w_v", "视觉分支连续权重", "[0.05,0.95]"],
        ],
        [1.0, 4.0, 2.2],
    )

    doc.add_heading("3.2 固定的 1/2/4-shot 嵌套抽样", level=2)
    add_body(doc, "本文使用 seed 0、1、2 三个随机种子生成统一 manifest。对同一数据集、同一类别和同一 seed，1-shot 集合是 2-shot 集合的子集，2-shot 集合又是 4-shot 集合的子集。这样可以把性能变化主要归因于参考图数量增加，而不是抽样完全改变。VisA manifest 共记录 252 个选择路径，MVTec manifest 共记录 315 个选择路径；两份清单都通过路径存在、类别一致、无重复和嵌套关系验证。")
    add_body(doc, "所有方法必须读取同一 manifest。方法可以按自身需要建立记忆库、提取提示或计算校准参数，但不能重新随机选择正常图。若官方代码默认随机抽样，项目通过适配层注入固定路径，并在运行日志中保存实际选中的参考图。")

    doc.add_heading("3.3 统一预测接口与审计", level=2)
    add_body(doc, "不同官方仓库的输出格式并不统一。本文将所有方法的预测转换为共同 NPZ 格式，至少包含 sample_ids、image_scores、pixel_maps、image_labels 和 pixel_masks。融合前先检查样本编号、顺序、标签方向和像素图尺寸；评测前再检查类别数、样本数、有限值和 schema 错误。每次正式运行保留独立 RunId、配置、日志、预测缓存、汇总表和来源哈希，避免覆盖历史证据。")
    add_table(
        doc,
        "表 3-2 路由器信息使用边界",
        ["信息", "是否允许", "原因"],
        [
            ["冻结分支预测", "允许", "路由器的直接输入"],
            ["K 张目标正常参考图", "允许", "少样本协议明确提供"],
            ["正常参考的确定性增强视图", "允许", "只用于稳定性与校准"],
            ["测试图像标签", "禁止", "会直接泄漏检测答案"],
            ["测试像素真值", "禁止", "会泄漏定位答案"],
            ["测试集均值/标准差", "禁止", "会利用整体测试分布调参"],
            ["查看最终结果后继续调 V1", "禁止", "破坏独立验证边界"],
        ],
        [2.4, 1.2, 4.0],
        footnote="所有冻结 V1 正式运行都记录 test_predictions_used=false 和 test_labels_used=false。",
    )

    doc.add_heading("3.4 开发集与最终验证边界", level=2)
    add_body(doc, "动态融合只使用 VisA seed 0 做设计与敏感性分析。图像温度、像素温度、最小权重和路由边界在 seed 0 上确定后冻结。VisA seed 1/2 与 MVTec seed 0/1/2 用于最终验证。虽然项目额外对 VisA seed 0 的 K=2/K=4 做了冻结方案复查，但这些结果属于开发证据，不与 seed 1/2 混合成“独立验证的三 seed 均值”。")

    # Chapter 4.
    doc.add_heading("4 基于不确定性路由的视觉—语言证据融合方法", level=1)
    doc.add_heading("4.1 方法总体结构", level=2)
    add_body(doc, "方法数据流如图 4-1 所示。测试图像分别经过 AnomalyDINO 和 AnomalyCLIP，得到视觉分支与文本分支的图像分数和像素图。正常参考图经过同样的分支，产生校准所需的正常响应。随后，系统对两条分支进行样本对齐和稳健校准，计算不确定性、分支冲突和空间集中度等特征。双层路由器分别生成图像视觉权重和像素视觉权重，最后输出融合图像分数与融合热图。")
    add_figure(doc, architecture, "图 4-1 动态融合 V1 的总体数据流与信息边界", width=6.4)

    doc.add_heading("4.2 视觉分支与文本分支", level=2)
    add_body(doc, "视觉分支使用 AnomalyDINO。它把查询图像切分为 patch，用冻结的 DINOv2 提取局部特征，再与 K 张正常参考图构成的 patch 特征库计算最近邻距离。距离较大表示查询 patch 与正常参考差异较大。像素异常图由 patch 距离插值和聚合得到，图像异常分数由局部异常响应汇总得到。该分支不重新训练 DINOv2 骨干，但会根据类别、shot 和 seed 重建正常参考库。")
    add_body(doc, "文本分支使用冻结的 AnomalyCLIP 预测。该方法使用对象无关的正常/异常文本提示，将图像全局特征和局部特征与文本嵌入比较，得到图像级与像素级异常响应。本文把 AnomalyCLIP 视为零样本文本引导分支；它不冒充 1/2/4-shot 基线，而是在所有 shot 设置中提供相同类型的语义证据。")

    doc.add_heading("4.3 正常参考稳健校准", level=2)
    add_body(doc, "两条分支的原始分数范围差异明显，不能直接融合。对每个数据集、类别、seed、shot 和分支，本文只使用正常参考预测估计中心和尺度。图像分数的中心为中位数 μ，尺度为经过常数修正的中位绝对偏差 MAD；像素图通常包含大量接近零的背景值，如果直接展平，MAD 容易变为零，因此先计算每个参考视图像素响应的 0.99 分位数，再用这些分位数估计中心和尺度。")
    add_equation(doc, "δ = max(1.4826 · median(|r − median(r)|), 10⁻⁶)", "（4-1）")
    add_equation(doc, "p(s) = sigmoid((s − μ) / (δ · τ_c))", "（4-2）")
    add_body(doc, "其中 τ_c 为校准温度，正式 V1 中取 1.0。所有变换结果裁剪到 [10⁻⁶,1−10⁻⁶]，防止对数运算出现无穷值。校准文件还记录数据来源、每个输入缓存的 SHA256，以及 test_predictions_used=false、test_labels_used=false。")

    doc.add_heading("4.4 分支不确定性与连续权重", level=2)
    add_body(doc, "当分支没有提供额外不确定性时，V1 使用校准概率的二元熵作为不确定性。概率接近 0.5 时熵最大，表示分支难以判断；概率接近 0 或 1 时熵较小，表示分支更确定。")
    add_equation(doc, "H(p) = −p·log₂p − (1−p)·log₂(1−p)", "（4-3）")
    add_body(doc, "路由器把负熵作为可靠性 logits，经二分支 softmax 得到视觉权重。较低不确定性的分支获得较大权重；视觉权重被限制在 [0.05,0.95]，避免任一分支被完全关闭。")
    add_equation(doc, "w_v = exp(−u_v / T) / [exp(−u_v / T) + exp(−u_t / T)]", "（4-4）")
    add_equation(doc, "s_f = w_v^I·p_v^I + (1−w_v^I)·p_t^I", "（4-5）")
    add_equation(doc, "M_f = w_v^P⊙p_v^P + (1−w_v^P)⊙p_t^P", "（4-6）")
    add_body(doc, "式中上标 I 和 P 分别表示图像级与像素级。冻结参数为图像温度 T_I=0.50、像素温度 T_P=0.20、最小权重 0.05、decision margin 0.15。温度越小，权重越接近“选择一个分支”；温度越大，权重越平滑。decision margin 只把连续权重转换为“视觉主导、文本主导、加权融合”三种标签，不改变实际融合分数。")

    doc.add_heading("4.5 一致性与冲突特征", level=2)
    add_body(doc, "项目接口还实现了图像/像素分支差异、分支一致度、空间响应集中度、正常参考增强一致性和跨 shot 敏感性。V1 的最终权重主要由熵决定，这些特征用于审计和后续设计，而没有在看过最终结果后重新训练一个门控网络。这样做牺牲了一部分灵活性，但保留了清晰的因果边界：V1 的结果反映冻结规则本身，而不是针对最终测试集重新拟合。")

    doc.add_heading("4.6 方法复杂度与实现特性", level=2)
    add_body(doc, "融合阶段直接读取已经冻结的 NPZ 预测，因此不需要重新运行两个分支模型。其主要计算包括逐元素校准、熵计算、softmax 权重和加权求和，额外计算量相对于分支推理较小。当前文档没有填写正式 GPU 显存峰值和推理耗时，因为这些数值尚未在统一计时协议下测量；论文最终版应在同一硬件、同一输入尺寸、排除首次加载后报告。")

    doc.add_heading("4.7 推理与审计步骤", level=2)
    add_body(doc, "一次正式融合运行按固定顺序执行。任一步失败都停止该运行并保留错误报告，不允许在样本缺失或顺序不一致时继续计算指标。")
    add_table(
        doc,
        "表 4-1 DynamicFusion V1 的执行步骤",
        ["步骤", "输入", "处理", "输出/验收"],
        [
            ["1 分支加载", "视觉/文本 NPZ", "检查类别与 schema", "两分支均为 passed"],
            ["2 样本对齐", "sample_ids、标签、像素图", "核对顺序和尺寸", "N、H、W 完全一致"],
            ["3 校准加载", "正常参考校准文件", "检查状态和 SHA256", "未使用测试信息"],
            ["4 特征计算", "校准概率", "熵、冲突、集中度", "全部有限且形状正确"],
            ["5 双层路由", "图像/像素不确定性", "计算并裁剪权重", "权重位于 [0.05,0.95]"],
            ["6 评测留存", "融合分数和热图", "统一评测与来源记录", "RunId、报告、缓存齐全"],
        ],
        [1.35, 2.0, 2.2, 2.3],
        font_size=7.6,
        cell_margins_dxa={"top": 50, "bottom": 50, "start": 70, "end": 70},
    )

    # Chapter 5.
    doc.add_heading("5 实验设置", level=1)
    doc.add_heading("5.1 数据集", level=2)
    add_table(
        doc,
        "表 5-1 数据集与本项目统一评测范围",
        ["数据集", "官方规模", "类别", "项目测试样本", "少样本设置"],
        [
            ["MVTec AD", "5,354 张高分辨率图像", "15", "1,725", "1/2/4-shot，3 seeds"],
            ["VisA", "10,821 张高分辨率图像", "12", "2,162", "1/2/4-shot，3 seeds"],
        ],
        [1.5, 2.5, 0.8, 1.4, 2.2],
        footnote="官方规模来自数据集论文[1][2]；项目测试样本数来自统一 manifest 与审计报告。",
    )
    add_body(doc, "MVTec AD 包含纹理和物体类别，异常形式包括划痕、凹陷、污染和结构变化等[1]。VisA 包含 12 个物体类别，具有更复杂的对象排列、背景和局部结构[2]。本文不修改官方测试划分，只从正常训练图中选择 K 张参考图。异常测试图及其掩膜仅用于最终计算指标。")

    doc.add_heading("5.2 对比方法与公平性标签", level=2)
    add_table(
        doc,
        "表 5-2 方法、信息来源与当前实验状态",
        ["方法", "类型", "目标正常样本用途", "当前状态", "论文处理"],
        [
            ["PatchCore", "纯视觉记忆库", "建立 patch 库", "VisA/MVTec 9/9", "正式主表"],
            ["WinCLIP+", "视觉—语言 + 参考图", "补充正常视觉证据", "VisA/MVTec 9/9", "正式主表"],
            ["AnomalyDINO", "纯视觉近邻", "建立正常 patch 库", "VisA/MVTec 9/9", "正式主表"],
            ["PromptAD", "目标域提示学习", "训练提示参数", "VisA 9/9；MVTec 4/9", "标记 target_normal_tuning=true"],
            ["AnomalyCLIP", "零样本文本引导", "不使用", "两数据集 zero-shot", "单独报告/作为文本分支"],
            ["ReMP-AD", "检索增强多模态", "参考检索与融合", "Gate A 待完成", "不进入结果排名"],
            ["AdaptCLIP", "CLIP 适配器", "少样本推理", "缺 checkpoint，Gate A 阻塞", "不进入结果排名"],
        ],
        [1.45, 1.8, 2.0, 1.7, 2.1],
        footnote="9/9 表示 3 个 shot × 3 个 seed 全部通过类别数、样本数和 schema 审计。",
        font_size=7.7,
    )

    doc.add_heading("5.3 实现环境", level=2)
    add_body(doc, "实验在 Windows 环境中运行，GPU 为 NVIDIA RTX 3060 Laptop，显存 6 GB。AnomalyCLIP 环境使用 Python 3.10、PyTorch 2.0.0+cu118 和 torchvision 0.15.1+cu118。不同方法保留独立虚拟环境和官方源码版本，避免依赖冲突。PatchCore 和 AnomalyDINO 的 FAISS 首轮使用 CPU，以减少 6 GB 显存压力。所有长任务采用串行队列、类别级 marker 和运行后验收，超时、OOM 或中断记录保留，不以目录存在作为完成证据。")

    doc.add_heading("5.4 评价指标", level=2)
    add_body(doc, "图像级指标包括 AUROC、平均精确率（AP）和最大 F1；像素级指标包括 Pixel AUROC、Pixel AP 和 AUPRO。AUROC 衡量正负样本在所有阈值下的排序能力，对类别不平衡相对稳定。AP 更关注异常样本的精确率—召回率关系。AUPRO 先计算不同连通异常区域的重叠率，再在规定假阳性率范围内积分，能够减少大缺陷区域对像素指标的支配。本文主表优先展示 Image AUROC、Pixel AUROC、Pixel AP 和 AUPRO。")
    add_body(doc, "所有主表按类别先计算指标，再做宏平均；3 个 seed 完整时报告均值±标准差。若不足 3 个 seed，则只标注完成数量，不用不完整均值冒充正式结果。")
    add_table(
        doc,
        "表 5-3 主指标的关注对象与解释限制",
        ["指标", "层级", "主要回答的问题", "解释时的限制"],
        [
            ["Image AUROC", "图像", "异常图能否排在正常图前", "不直接反映阈值后的误报量"],
            ["Image AP", "图像", "异常样本的精确率—召回率", "受异常比例影响"],
            ["Pixel AUROC/AP", "像素", "异常像素能否与背景区分", "大区域可能占据更多像素"],
            ["AUPRO", "区域", "不同缺陷区域是否被覆盖", "依赖规定的假阳性率范围"],
        ],
        [1.4, 1.0, 2.8, 2.4],
        font_size=7.8,
        cell_margins_dxa={"top": 55, "bottom": 55, "start": 80, "end": 80},
    )
    add_body(doc, "均值和标准差的单位均为百分点。正文中的“提高 4.14 个百分点”表示从一个百分比数值增加 4.14，而不是相对增长 4.14%。当图像 AUROC 与 AUPRO给出不同结论时，本文分别解释整图排序和局部定位，不用一个指标替代另一个。")

    # Chapter 6.
    doc.add_heading("6 统一基线复现结果", level=1)
    doc.add_heading("6.1 VisA 基线结果", level=2)
    visa_table = []
    for row in visa_rows:
        visa_table.append([
            row["method"] + ("*" if row["method"] == "PromptAD" else ""),
            row["shot"],
            f'{pct(row["image_auroc_mean"])} ± {float(row["image_auroc_std"])*100:.2f}',
            f'{pct(row["pixel_auroc_mean"])} ± {float(row["pixel_auroc_std"])*100:.2f}',
            f'{pct(row["pixel_ap_mean"])} ± {float(row["pixel_ap_std"])*100:.2f}',
            f'{pct(row["aupro_mean"])} ± {float(row["aupro_std"])*100:.2f}',
        ])
    add_table(
        doc,
        "表 6-1 VisA 统一 1/2/4-shot 基线结果（3 seeds，%）",
        ["方法", "Shot", "Image AUROC", "Pixel AUROC", "Pixel AP", "AUPRO"],
        visa_table,
        [1.6, 0.7, 1.8, 1.8, 1.6, 1.6],
        footnote="* PromptAD 使用目标正常样本进行提示学习，target_normal_tuning=true。其余方法按各自冻结/记忆库协议运行。",
        font_size=7.8,
    )
    add_body(doc, "表 6-1 显示，AnomalyDINO 在 VisA 的四项主指标上总体最强，并且参考图从 1 张增加到 4 张时持续改善。其 Image AUROC 从 89.40% 提升到 92.58%，AUPRO 从 92.21% 提升到 93.69%。这说明 DINOv2 patch 特征与少样本近邻匹配能够有效利用新增正常参考。")
    add_body(doc, "PatchCore 随 shot 增加也有明显提升，Image AUROC 从 68.03% 增至 78.68%；WinCLIP+ 的提升较平缓，但 AUPRO 稳定高于 PatchCore。PromptAD 的 Pixel AUROC 较高，但 Image AUROC 约为 81%，而且不同 shot 并非单调上升。由于 PromptAD 会在目标正常样本上训练提示，它与完全冻结方法的计算成本和信息使用方式不同，最终论文需要保留单独标记。")

    doc.add_heading("6.2 MVTec AD 基线与当前完整性", level=2)
    mvtec_table = []
    for row in mvtec_rows:
        method = row["method"]
        if method == "PromptAD":
            method += "*"
        mvtec_table.append([
            method,
            row["shot"],
            row["image_auroc_mean_std"],
            row["pixel_auroc_mean_std"],
            row["pixel_ap_mean_std"],
            row["aupro_mean_std"],
        ])
    add_table(
        doc,
        "表 6-2 MVTec AD 当前可用结果（%，2026-08-09 审计状态）",
        ["方法", "Shot", "Image AUROC", "Pixel AUROC", "Pixel AP", "AUPRO"],
        mvtec_table,
        [1.65, 0.65, 1.8, 1.8, 1.55, 1.55],
        footnote="PromptAD 只有 4/9 组合完成，因此单元格显示 n=x/3，不生成正式均值±标准差；DynamicFusion 在本表只作为冻结 V1 结果，不用于宣称优于基线。",
        font_size=7.4,
    )
    add_body(doc, "MVTec 上，AnomalyDINO 的 Image AUROC 为 95.71%/96.86%/97.46%，明显高于 PatchCore 和 WinCLIP+。Pixel AUROC 与 AUPRO 同样保持领先。这个结果确立了一个重要约束：动态融合若要成为性能贡献，必须在不破坏 AnomalyDINO 全局排序的情况下，利用文本分支改善少数视觉失败样本。")
    add_body(doc, "PromptAD 的 MVTec 正式矩阵尚未完成。当前 1-shot 有 2 个 seed，2-shot 和 4-shot 各有 1 个 seed，剩余 5 个组合已进入可断点恢复队列，但按项目计划保持暂停。因此，初稿不能给出 PromptAD 的 MVTec 三 seed 排名。ReMP-AD 和 AdaptCLIP 尚未通过 Gate A，也不应出现在数值主表中。")

    doc.add_heading("6.3 基线复现的主要结论", level=2)
    add_body(doc, "统一复现说明，不同方法对 shot 数量的响应差异较大。纯视觉近邻方法通常能直接从更多正常参考中获益；视觉—语言方法的表现还受提示、窗口和文本语义影响；目标域提示学习方法则额外包含训练过程。若只比较最终指标而忽略这些信息来源，会把“训练方式差异”误当成“算法本身差异”。因此，本文后续融合实验固定使用已审计的 AnomalyDINO 与 AnomalyCLIP 预测，并保留固定权重和单分支对照。")

    # Chapter 7.
    doc.add_heading("7 动态融合实验与消融", level=1)
    doc.add_heading("7.1 冻结参数与验证范围", level=2)
    add_body(doc, "V1 在 VisA seed 0 上完成温度和 decision margin 的小范围敏感性实验，最终冻结图像温度 0.50、像素温度 0.20、decision margin 0.15 和 min_weight 0.05。最终审计覆盖 17 个运行：VisA 8 个、MVTec 9 个，共 231 个“类别—运行”组合。全部运行通过来源、样本数、schema 和禁止信息标记检查；来源清单包含 693 行记录、285 个唯一输入文件及其 SHA256。")
    add_callout(doc, "独立验证说明", "VisA seed 1/2 的 6 个组合是独立验证；VisA seed 0 只用于开发和冻结方案复查。MVTec seed 0/1/2 是冻结后运行，但其结果已经被用于失效分析，今后不能再拿同一结果调 V1 或 V2。")

    doc.add_heading("7.2 VisA 独立验证结果", level=2)
    dyn_rows = []
    for key in ["s1_k1", "s1_k2", "s1_k4", "s2_k1", "s2_k2", "s2_k4"]:
        row = visa_dynamic[key]
        seed = key[1]
        shot = key[4]
        dyn_rows.append([
            seed,
            shot,
            f'{float(row["image_auroc"])*100:.2f}',
            f'{float(row["pixel_auroc"])*100:.2f}',
            f'{float(row["pixel_ap"])*100:.2f}',
            f'{float(row["aupro"])*100:.2f}',
        ])
    add_table(
        doc,
        "表 7-1 DynamicFusion V1 在 VisA seed 1/2 的独立验证结果（%）",
        ["Seed", "Shot", "Image AUROC", "Pixel AUROC", "Pixel AP", "AUPRO"],
        dyn_rows,
        [0.8, 0.8, 1.8, 1.8, 1.6, 1.6],
        footnote="seed 0 开发结果不并入本表，以避免把开发集当作独立验证。",
        font_size=8.0,
    )
    add_body(doc, "VisA 独立验证中，Image AUROC 大致保持在 79.77%—82.50%。K=2 和 K=4 的像素定位明显好于 K=1，AUPRO 在 seed 1 分别为 83.31% 和 82.47%，在 seed 2 分别为 79.08% 和 84.81%。这说明多张正常参考有助于像素校准和局部权重，但图像级结果仍明显低于原始 AnomalyDINO。")

    doc.add_heading("7.3 双温度消融", level=2)
    selected_variants = {
        "calibrated_visual_only": "校准后视觉单分支",
        "calibrated_text_only": "校准后文本单分支",
        "fixed_visual_0.50": "固定视觉权重 0.50",
        "fixed_visual_0.75": "固定视觉权重 0.75",
        "single_temperature_0.20": "单温度 0.20",
        "split_image_0.50_pixel_0.20": "双温度 0.50/0.20",
    }
    abl_table = []
    for row in ablation_rows:
        if row["variant"] in selected_variants:
            abl_table.append([
                selected_variants[row["variant"]],
                row["shot"],
                f'{float(row["image_auroc"])*100:.2f}',
                f'{float(row["pixel_auroc"])*100:.2f}',
                f'{float(row["pixel_ap"])*100:.2f}',
                f'{float(row["aupro"])*100:.2f}',
            ])
    add_table(
        doc,
        "表 7-2 VisA seed 0 主要消融结果（%）",
        ["变体", "Shot", "Image AUROC", "Pixel AUROC", "Pixel AP", "AUPRO"],
        abl_table,
        [2.3, 0.65, 1.55, 1.55, 1.35, 1.35],
        footnote="完整消融还包括固定权重 0/0.25/1.00 和单温度 0.50；本表保留最有解释力的对照。",
        font_size=7.3,
    )
    add_figure(doc, fig_dir / "visa_ablation_split_temperature.png", "图 7-1 VisA seed 0 固定权重、单温度与双温度消融", width=6.35)
    add_body(doc, "双温度的作用主要体现在定位目标。K=2 时，双温度方案相对最佳固定权重的 Image AUROC 提高 0.20 个百分点，AUPRO 提高 4.14 个百分点；K=4 时分别提高 0.04 和 8.82 个百分点。K=1 的 Image AUROC 比最佳固定权重低 0.07 个百分点，但 AUPRO 高 2.78 个百分点。该结果支持“图像路由与像素路由需要不同强度”，但不能证明动态融合已经解决了上游校准问题。")

    doc.add_heading("7.4 MVTec 冻结验证", level=2)
    mvtec_compare = [
        ["1", "95.71 ± 0.02", "79.43 ± 2.61", "−16.29", "91.06 ± 0.40", "82.33 ± 3.12"],
        ["2", "96.86 ± 0.51", "86.37 ± 1.98", "−10.49", "93.79 ± 0.43", "90.58 ± 1.36"],
        ["4", "97.46 ± 0.47", "89.52 ± 2.69", "−7.94", "94.18 ± 0.03", "91.81 ± 0.80"],
    ]
    add_table(
        doc,
        "表 7-3 MVTec 上原始视觉分支与 DynamicFusion V1（3 seeds，%）",
        ["Shot", "AnomalyDINO\nImage AUROC", "Dynamic\nImage AUROC", "差值", "Dynamic\nPixel AUROC", "Dynamic\nAUPRO"],
        mvtec_compare,
        [0.7, 2.0, 1.8, 0.9, 1.8, 1.6],
        font_size=7.8,
    )
    add_figure(doc, fig_dir / "mvtec_visual_vs_dynamic_by_shot.png", "图 7-2 MVTec 三个 shot 下 AnomalyDINO 与 DynamicFusion 的 Image AUROC", width=5.65)
    add_body(doc, "冻结 V1 在 MVTec 三个 shot 下都未超过原始 AnomalyDINO，且 1-shot 的下降最大。shot 增加后差距从 16.29 个百分点缩小到 7.94 个百分点，说明更多正常参考能让校准稍稳定，但没有消除根本问题。Pixel AUROC 和 AUPRO 的绝对值相对较高，却仍不足以把方法描述为总体优于视觉分支。")

    # Chapter 8.
    doc.add_heading("8 失效分析、讨论与有效性边界", level=1)
    doc.add_heading("8.1 主要失效链：校准饱和", level=2)
    add_body(doc, "动态融合不符合预期的首要原因发生在路由之前。正常参考图很少时，视觉分支正常分数的 MAD 可能极小。测试图分数只要略高于参考中心，标准化值就会非常大，经过 sigmoid 后全部接近 1。理论上，严格单调变换不会改变排序；但实际实现包含有限精度、裁剪和大量数值并列，强视觉分支原本细致的排序被压扁。")
    add_figure(doc, fig_dir / "calibration_saturation_diagnostic.png", "图 8-1 视觉分支校准饱和比例与校准前后 Image AUROC", width=6.35)
    add_body(doc, "MVTec 中视觉校准分数不小于 0.999 的平均比例达到 99.99%，VisA 为 91.54%。MVTec 原始 AnomalyDINO 的平均类别级 Image AUROC 为 95.71%/96.86%/97.46%，校准后分别降到 51.53%/54.46%/56.74%。这说明 V1 的主要损失不是文本分支简单“拉低”结果，而是视觉分支在进入融合前已经丢失大量排序信息。")
    add_table(
        doc,
        "表 8-1 V1 的主要失效链及其可观测证据",
        ["环节", "发生的问题", "证据", "后果"],
        [
            ["参考校准", "MAD 尺度过小", "视觉分数大面积 ≥0.999", "形成大量并列值"],
            ["不确定性", "熵只看概率极端程度", "饱和 0.999 熵很低", "错误地视为高置信"],
            ["路由权重", "权重与真实分支优势弱相关", "类别 Spearman 较低", "不能稳定选择更强分支"],
            ["图像融合", "每张图权重不同", "融合/视觉排序相关约 0.54", "正常/异常全局排序变化"],
            ["像素融合", "文本响应可能扩散", "失败案例背景升温", "局部对比度下降"],
        ],
        [1.25, 2.1, 2.3, 2.0],
        font_size=7.8,
    )

    doc.add_heading("8.2 熵置信度为什么不能识别饱和错误", level=2)
    add_body(doc, "二元熵在 p=0.5 时最大，在 p→0 或 p→1 时接近零。这个性质适合描述一个已经校准良好的二分类概率，却无法区分“模型真的确定”和“数值因为校准范围过窄而被推到边界”。因此，当前路由器会把由尺度问题产生的 0.999999 当作可靠异常证据。正确的 V2 不确定性需要额外判断：测试分数距离正常参考分布有多远、是否已经超出可校准范围、不同增强视图是否稳定，以及两条分支是否在异常区域上形成一致证据。")

    doc.add_heading("8.3 权重、路由标签与真实分支优势", level=2)
    add_figure(doc, fig_dir / "route_weight_summary.png", "图 8-2 图像权重、像素权重及正常/异常样本权重统计", width=6.35)
    add_body(doc, "如果路由有效，在 AnomalyDINO 明显优于 AnomalyCLIP 的类别中，视觉权重也应明显更高。实际类别级 Spearman 相关性较弱：MVTec 三个 shot 平均约为 0.10、0.16、0.20，VisA 约为 0.36、0.37、0.31。MVTec 的 5,175 个路由判断中没有文本主导；VisA 也只有极少文本主导。这说明离散标签的数量不能证明路由器学会了分支选择，真正需要检查的是连续权重是否与正确性相关。")

    doc.add_heading("8.4 类别差异与案例分析", level=2)
    add_figure(doc, fig_dir / "mvtec_category_image_auroc_delta_heatmap.png", "图 8-3 MVTec 各类别 DynamicFusion 相对 AnomalyDINO 的 Image AUROC 差值", width=6.4)
    add_body(doc, "MVTec 中下降最明显的类别包括 carpet、cable、grid、zipper 和 bottle。例如，carpet 的原始视觉 AUROC 为 100%，动态融合约为 61.31%；cable 从约 92.36% 降至 54.02%。少数类别仍存在互补：capsule 平均提高约 1.42 个百分点，wood 提高约 0.54 个百分点，leather 基本持平。这些局部改善证明文本证据并非完全无效，但其频率和幅度不足以支持总体优越性。")
    add_figure(doc, fig_dir / "mvtec_success_failure_cases.png", "图 8-4 MVTec 自动选择的定位改善与退化案例（原图、真值、两分支、融合图和像素视觉权重）", width=6.15)
    add_body(doc, "成功案例中，文本热图补充了视觉分支未覆盖的细长结构，像素权重在结构边缘产生选择，融合后异常区域与背景的对比提高。失败案例中，视觉分支本已集中在真值区域，而文本响应或融合权重把高响应扩散到纹理、主体轮廓和背景。图 8-4 的案例按照“融合后异常区域对比度减去视觉分支对比度”自动选择，不是人工只挑好看的结果；但它们来自已查看的最终验证，只能用于解释机制，不能再用于调参。")

    doc.add_heading("8.5 为什么图像级和像素级表现不同", level=2)
    add_body(doc, "Image AUROC 只依赖所有样本的相对排序。若不同样本采用不同权重，即使两条分支各自有合理排序，融合后也可能把某些正常图抬高、把异常图压低。MVTec 中动态分数与原始视觉分数的类别级平均 Spearman 相关性约为 0.54，说明排序发生明显改变。像素任务则更关注同一张图内部的空间对比，局部权重有机会利用文本边缘或区域提示，因此双温度更容易在 AUPRO 上表现出局部收益。")

    doc.add_heading("8.6 有效性威胁与论文边界", level=2)
    for item in [
        "内部有效性：分支来自不同官方仓库，输入尺寸、预处理和特征层可能不同。统一接口减少格式差异，但不能消除算法原生配置差异。",
        "协议公平性：PromptAD 会使用目标正常样本学习提示，必须标记 target_normal_tuning=true；AnomalyCLIP 是零样本文本分支，不能冒充少样本矩阵。",
        "结果完整性：MVTec PromptAD 仍缺 5 个组合；ReMP-AD 与 AdaptCLIP 未通过 Gate A。初稿中的空缺不能用于排名或统计显著性结论。",
        "外部有效性：当前数据集只有 MVTec AD 和 VisA，且 V1 已在这两个数据集上完成分析。V2 必须重新建立开发集或使用额外数据，不能反复在同一最终集上优化。",
        "效率有效性：尚未在统一计时协议下测量峰值显存、吞吐和单图推理时间，因此本文不填写估计值。",
        "统计有效性：多数基线有 3 个 seed，但 VisA 动态融合独立验证只有 seed 1/2；论文应报告逐 seed 结果，并避免把 seed 0 开发结果包装为独立三 seed 验证。",
    ]:
        add_bullet(doc, item)

    doc.add_heading("8.7 V2 改进方向", level=2)
    add_body(doc, "V2 的首要任务不是更换温度，而是修复上游校准。建议按照以下顺序提出并验证新假设。")
    add_manual_number(doc, 1, "排序保持校准：使用经验分布、分位数映射或仅做单调且不产生大量并列的尺度变换，并显式检查校准前后 Spearman 相关性。")
    add_manual_number(doc, 2, "饱和保护：统计测试分数超出正常参考支持范围的比例；当饱和比例超过阈值时，不把低熵直接解释为高置信。")
    add_manual_number(doc, 3, "安全回退：以强视觉分支为默认输出，只有当文本分支在正常增强一致性、分支冲突和空间集中度上同时满足条件时才允许明显改权重。")
    add_manual_number(doc, 4, "任务分离：图像级路由优先保持全局排序，像素级路由关注局部对比和连通区域；两者使用独立特征和损失。")
    add_manual_number(doc, 5, "新验证边界：从未用于 V1 分析的数据中建立 V2 开发集，锁定方法后再对保留集测试。若没有新数据，应把 V2 结果明确写成探索性结果。")

    doc.add_heading("8.8 V2 的最小验收条件", level=2)
    add_body(doc, "V2 不以“某一个指标比 V1 高”作为唯一成功标准。为了避免再次出现强分支被静默破坏的问题，开发阶段至少需要同时通过下列保护性检查。")
    add_table(
        doc,
        "表 8-2 DynamicFusion V2 建议验收条件",
        ["检查", "最低要求", "失败时处理"],
        [
            ["排序保持", "校准后视觉分数与原始分数保持高相关", "停用该校准方案"],
            ["饱和比例", "极端并列值不再大面积出现", "增大尺度下限或改用分位数映射"],
            ["安全回退", "文本证据不足时恢复视觉输出", "限制最大改权幅度"],
            ["任务分离", "图像和像素指标分别不劣于控制线", "拆分特征、温度或损失"],
            ["独立验证", "冻结后只在保留数据上测试一次", "结果降级为探索性证据"],
        ],
        [1.6, 3.6, 2.4],
        font_size=7.7,
        cell_margins_dxa={"top": 50, "bottom": 50, "start": 75, "end": 75},
    )

    # Chapter 9.
    doc.add_heading("9 结论与后续工作", level=1)
    doc.add_heading("9.1 研究结论", level=2)
    add_body(doc, "本文围绕少样本工业异常检测，完成了从数据许可与校验、统一抽样、四种主要基线复现，到视觉—语言动态融合设计、冻结验证、消融和失效分析的完整流程。VisA 四方法基线 36/36 运行已经通过统一审计；MVTec 的 PatchCore、WinCLIP+ 和 AnomalyDINO 各完成 9/9，DynamicFusion 完成 9/9。结果表明，AnomalyDINO 是当前最强视觉基线，更多正常参考总体能够提高其图像检测和像素定位能力。")
    add_body(doc, "本文设计的 V1 使用正常参考校准、二元熵不确定性和图像/像素双温度路由。在 VisA seed 0 消融中，双温度对 AUPRO 出现局部收益，说明图像判断与像素定位应采用不同路由强度；但 MVTec 最终验证表明，V1 没有超过原始 AnomalyDINO。科学分析把原因定位为参考尺度过小导致的 sigmoid 饱和、熵对饱和错误的误判，以及样本相关权重对全局排序的破坏。")
    add_body(doc, "因此，本文当前最稳妥的结论不是“动态融合一定优于单分支”，而是：视觉和文本证据确实存在局部互补，但可靠融合的前提是校准必须保留强分支排序，不确定性必须识别超出参考分布的预测，并且路由需要具有安全回退机制。这个结论由运行级、类别级、权重级和案例级证据共同支持。")

    doc.add_heading("9.2 后续工作安排", level=2)
    add_body(doc, "论文初稿完成后，项目后续工作分为实验补全和方法升级两条轨道。实验补全优先恢复 PromptAD MVTec 剩余 5 个组合；随后对 ReMP-AD 与 AdaptCLIP 依次执行 Gate A，只有单类别输入、显存、预测导出和统一评测全部通过后，才进入完整矩阵。最终再补齐公平主表、效率表和统计分析。方法升级则按第 8.7 节开发 V2，并建立新的开发/验证边界。")
    add_table(
        doc,
        "表 9-1 初稿之后的项目工作包与验收条件",
        ["工作包", "主要内容", "资源", "验收条件", "论文影响"],
        [
            ["P1 初稿修订", "按导师意见修改结构与表述", "CPU", "章节逻辑、图表编号、引用一致", "形成 V0.2"],
            ["P2 PromptAD 补全", "MVTec 剩余 5 个 seed/shot", "长时 GPU", "每组 15 类、1,725 样本、0 schema 错", "补完整三 seed 主表"],
            ["P3 ReMP-AD Gate A", "manifest、NPZ、bottle 单类", "GPU", "显存可运行且统一指标通过", "决定是否进入矩阵"],
            ["P4 AdaptCLIP Gate A", "checkpoint、batch 1、bottle", "GPU", "权重校验、6 GB 显存通过", "决定是否进入矩阵"],
            ["P5 效率测量", "显存、加载后单图耗时、吞吐", "GPU+CPU", "统一硬件/尺寸/预热协议", "补效率表"],
            ["P6 DynamicFusion V2", "排序保持校准、饱和保护、回退", "先 CPU 后 GPU", "新开发/验证边界，强分支不退化", "决定最终方法主张"],
        ],
        [1.4, 2.2, 1.1, 2.3, 1.7],
        font_size=7.2,
        cell_margins_dxa={"top": 35, "bottom": 35, "start": 55, "end": 55},
    )

    # References.
    doc.add_heading("参考文献", level=1)
    references = [
        "[1] BERGMANN P, FAUSER M, SATTLEGGER D, et al. MVTec AD—A comprehensive real-world dataset for unsupervised anomaly detection[C]//Proceedings of the IEEE/CVF Conference on Computer Vision and Pattern Recognition. 2019: 9592-9600.",
        "[2] ZOU Y, JEONG J, PEMULA L, et al. SPot-the-Difference self-supervised pre-training for anomaly detection and segmentation[C]//European Conference on Computer Vision. 2022.",
        "[3] ROTH K, PEMULA L, ZEPEDA J, et al. Towards total recall in industrial anomaly detection[C]//Proceedings of the IEEE/CVF Conference on Computer Vision and Pattern Recognition. 2022: 14318-14328.",
        "[4] JEONG J, ZOU Y, KIM T, et al. WinCLIP: Zero-/few-shot anomaly classification and segmentation[C]//Proceedings of the IEEE/CVF Conference on Computer Vision and Pattern Recognition. 2023: 19606-19616.",
        "[5] RADFORD A, KIM J W, HALLACY C, et al. Learning transferable visual models from natural language supervision[C]//Proceedings of the 38th International Conference on Machine Learning. PMLR, 2021, 139: 8748-8763.",
        "[6] ZHOU Q, PANG G, TIAN Y, et al. AnomalyCLIP: Object-agnostic prompt learning for zero-shot anomaly detection[C]//International Conference on Learning Representations. 2024.",
        "[7] OQUAB M, DARCET T, MOUTAKANNI T, et al. DINOv2: Learning robust visual features without supervision[J]. Transactions on Machine Learning Research, 2024.",
        "[8] DAMM S, LASZKIEWICZ M, LEDERER J, et al. AnomalyDINO: Boosting patch-based few-shot anomaly detection with DINOv2[C]//Proceedings of the Winter Conference on Applications of Computer Vision. 2025: 1319-1329.",
        "[9] LI X, ZHANG Z, TAN X, et al. PromptAD: Learning prompts with only normal samples for few-shot anomaly detection[C]//Proceedings of the IEEE/CVF Conference on Computer Vision and Pattern Recognition. 2024.",
        "[10] MA H, YANG G, ZHAO D, et al. ReMP-AD: Retrieval-enhanced multi-modal prompt fusion for few-shot industrial visual anomaly detection[C]//Proceedings of the IEEE/CVF International Conference on Computer Vision. 2025: 20425-20434.",
        "[11] GAO B B, ZHOU Y, YAN J, et al. AdaptCLIP: Adapting CLIP for universal visual anomaly detection[C]//Proceedings of the AAAI Conference on Artificial Intelligence. 2026, 40(6): 4095-4103. DOI: 10.1609/aaai.v40i6.42404.",
        "[12] DEFARD T, SETKOV A, LOESCH A, et al. PaDiM: A patch distribution modeling framework for anomaly detection and localization[C]//25th International Conference on Pattern Recognition. 2021: 475-489.",
        "[13] DENG H, LI X. Anomaly detection via reverse distillation from one-class embedding[C]//Proceedings of the IEEE/CVF Conference on Computer Vision and Pattern Recognition. 2022: 9737-9746.",
    ]
    for ref in references:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(22)
        p.paragraph_format.first_line_indent = Pt(-22)
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(ref)
        r.font.name = "Times New Roman"
        r.font.size = Pt(9.5)
        set_east_asia_font(r, "宋体")

    # Appendices.
    doc.add_heading("附录 A 当前实验完整性与待补内容", level=1)
    add_body(doc, "本附录用于区分“已经可以写入论文的证据”和“仍需补做的实验”。状态以 2026 年 8 月 9 日的项目同步快照为准；后续完成新实验后，应重新运行完整性审计并更新正文表格。")
    add_table(
        doc,
        "表 A-1 当前实验矩阵完整性",
        ["数据集/方法", "计划组合", "已通过", "状态", "当前处理"],
        [
            ["VisA/PatchCore", "9", "9", "完整", "正文正式表"],
            ["VisA/WinCLIP+", "9", "9", "完整", "正文正式表"],
            ["VisA/AnomalyDINO", "9", "9", "完整", "正文正式表"],
            ["VisA/PromptAD", "9", "9", "完整", "正文正式表，标记 tuning"],
            ["MVTec/PatchCore", "9", "9", "完整", "正文正式表"],
            ["MVTec/WinCLIP+", "9", "9", "完整", "正文正式表"],
            ["MVTec/AnomalyDINO", "9", "9", "完整", "正文正式表"],
            ["MVTec/PromptAD", "9", "4", "部分", "只标 n=x/3"],
            ["MVTec/DynamicFusion", "9", "9", "完整", "冻结 V1 分析"],
            ["ReMP-AD", "Gate A→矩阵", "0", "待门控", "不进排名"],
            ["AdaptCLIP", "Gate A→矩阵", "0", "checkpoint 阻塞", "不进排名"],
        ],
        [2.1, 1.15, 1.15, 1.55, 2.6],
        font_size=7.8,
    )
    add_body(doc, "PromptAD MVTec 剩余组合为 seed 1 的 K=2/K=4，以及 seed 2 的 K=1/K=2/K=4。正式队列状态为 paused_by_schedule，已有 seed 1、K=2 的类别级断点，应从该断点恢复而不是覆盖重跑。")

    doc.add_heading("附录 B 可复现性检查清单", level=1)
    add_table(
        doc,
        "表 B-1 论文结果复现与发布前检查",
        ["检查项", "要求", "V0.1 状态"],
        [
            ["数据合法来源", "MVTec 许可流程；VisA 官方公开源", "已确认"],
            ["压缩包与 manifest 哈希", "记录 SHA256", "已完成"],
            ["嵌套抽样", "K1⊂K2⊂K4，3 seeds", "已验证"],
            ["统一预测 schema", "sample_id、分数、像素图、标签", "已实现"],
            ["正式结果完整性", "类别数、样本数、0 schema error", "已审计；PromptAD MVTec 部分"],
            ["动态融合信息边界", "禁止测试标签、掩膜、整体统计", "17/17 通过"],
            ["开发/验证分离", "VisA s0 开发，s1/s2 独立验证", "已执行"],
            ["参数冻结", "最终验证后不继续调 V1", "已冻结"],
            ["效率测量", "统一硬件、尺寸、预热和重复次数", "待补"],
            ["外部方法 Gate A", "ReMP-AD、AdaptCLIP 先门控", "待补"],
        ],
        [2.1, 3.9, 1.7],
        font_size=8.0,
    )

    doc.add_heading("附录 C 初稿使用说明", level=1)
    add_body(doc, "本文件是基于当前真实实验状态形成的 V0.1 初稿，适合用于确定论文主线、与导师讨论章节结构，并作为后续补实验的占位框架。作者、学校、专业和指导教师信息需要在封面补充。参考文献已按当前可核对信息整理，提交学校前仍需按照学校指定的 GB/T 7714 或模板格式统一。")
    add_body(doc, "初稿中已经可以稳定保留的内容包括统一协议、方法接口、VisA 完整基线、MVTec 三种完整基线、DynamicFusion V1 的冻结结果、消融和失效分析。需要在最终稿补充的内容包括：PromptAD MVTec 完整三 seed 结果、ReMP-AD/AdaptCLIP Gate A 结论、正式效率测量，以及如决定开发 V2 后得到的新开发/验证结果。")
    add_callout(doc, "最重要的写作原则", "后续补结果时，只替换明确标为待补的表格和讨论，不删除 V1 的失效证据。若 V2 获得改进，也应把 V1 作为必要消融，说明排序保持校准和安全回退为什么重要。")

    doc.save(DOCX_PATH)
    return DOCX_PATH


if __name__ == "__main__":
    print(build_doc())

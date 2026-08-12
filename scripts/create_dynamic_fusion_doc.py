from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(r"C:\Users\lynle\Desktop\动态融合算法设计与第二阶段项目规划.docx")

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
PALE_YELLOW = "FFF2CC"
GRAY = "666666"


def set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths_dxa: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table_properties = table._tbl.tblPr
    table_width = table_properties.find(qn("w:tblW"))
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        table_properties.append(table_width)
    table_width.set(qn("w:w"), str(sum(widths_dxa)))
    table_width.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[index] / 1440)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_width = tc_pr.find(qn("w:tcW"))
            if tc_width is None:
                tc_width = OxmlElement("w:tcW")
                tc_pr.append(tc_width)
            tc_width.set(qn("w:w"), str(widths_dxa[index]))
            tc_width.set(qn("w:type"), "dxa")


def set_run_font(run, name="Microsoft YaHei", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 11.5, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ["List Bullet", "List Number"]:
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.32)
        style.paragraph_format.first_line_indent = Inches(-0.18)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def add_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("少样本工业异常检测项目｜第二阶段设计规划")
    set_run_font(run, size=8.5, color=GRAY)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("项目内部工作文档　|　2026-07-30　|　")
    set_run_font(run, size=8.5, color=GRAY)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    footer._p.append(field)


def add_title_block(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("项目第二阶段设计规划")
    set_run_font(run, size=23, color="000000", bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run("动态融合算法设计、验证边界与后续执行路线")
    set_run_font(run, size=14, color=GRAY)

    metadata = [
        ("项目主题", "基于不确定性路由的少样本工业异常检测"),
        ("当前策略", "基线轨继续训练；设计轨提前进行"),
        ("设计开发集", "VisA seed 0 的冻结预测缓存"),
        ("最终验证集", "VisA seed 1/2 与 MVTec seed 0/1/2"),
        ("文档目的", "解释第二阶段要做什么，以及如何保证结果公平、可复现"),
    ]
    table = doc.add_table(rows=0, cols=2)
    for label, value in metadata:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        set_cell_shading(cells[0], LIGHT_BLUE)
        for paragraph in cells[0].paragraphs:
            for run in paragraph.runs:
                set_run_font(run, size=10, color=DARK_BLUE, bold=True)
        for paragraph in cells[1].paragraphs:
            for run in paragraph.runs:
                set_run_font(run, size=10)
    set_table_widths(table, [2700, 6660])

    callout = doc.add_paragraph()
    callout.paragraph_format.space_before = Pt(14)
    callout.paragraph_format.space_after = Pt(12)
    callout.paragraph_format.left_indent = Inches(0.15)
    callout.paragraph_format.right_indent = Inches(0.15)
    callout.paragraph_format.line_spacing = 1.2
    ppr = callout._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), PALE_YELLOW)
    ppr.append(shd)
    run = callout.add_run("一句话理解：")
    set_run_font(run, size=10.5, color=DARK_BLUE, bold=True)
    run = callout.add_run(
        "第二阶段不是重新训练所有基线，而是让系统根据每张图、每个区域的情况，"
        "自动判断应该更相信视觉分支、文本分支，还是两者的组合。"
    )
    set_run_font(run, size=10.5)


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_para(doc, text, bold_prefix: str | None = None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        run = p.add_run(bold_prefix)
        set_run_font(run, bold=True)
        run = p.add_run(text[len(bold_prefix) :])
        set_run_font(run)
    else:
        run = p.add_run(text)
        set_run_font(run)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_run_font(run)


def add_numbers(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(item)
        set_run_font(run)


def add_status_table(doc):
    table = doc.add_table(rows=1, cols=4)
    headers = ["部分", "当前状态", "已完成内容", "下一步目标"]
    for i, value in enumerate(headers):
        table.rows[0].cells[i].text = value
        set_cell_shading(table.rows[0].cells[i], LIGHT_BLUE)
        for run in table.rows[0].cells[i].paragraphs[0].runs:
            set_run_font(run, size=9.5, color=DARK_BLUE, bold=True)
    rows = [
        ("数据与协议", "已完成", "MVTec、VisA；1/2/4-shot；3 seeds；固定 manifest", "保持不变，作为所有方法共同输入"),
        ("统一评测", "已完成", "图像 AUROC/AP/F1；像素 AUROC/AP/AUPRO；NPZ schema", "继续复用，不为融合单独改指标"),
        ("VisA基线", "大部分已完成", "PatchCore、WinCLIP+、AnomalyDINO完整矩阵", "继续完成PromptAD及近期方法"),
        ("PromptAD", "正在推进", "VisA seed 0 的1/2/4-shot已完成", "继续seed 1/2，之后做MVTec Gate A"),
        ("动态融合", "设计中", "接口、配置、规则路由、单元测试已建立", "先做校准，再做固定融合和动态路由"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            for paragraph in cells[i].paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, size=9.2)
    set_table_widths(table, [1500, 1500, 3300, 3060])


def add_term_table(doc):
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "术语"
    table.rows[0].cells[1].text = "简单解释"
    for cell in table.rows[0].cells:
        set_cell_shading(cell, LIGHT_BLUE)
        for run in cell.paragraphs[0].runs:
            set_run_font(run, size=10, color=DARK_BLUE, bold=True)
    terms = [
        ("文本分支", "利用“正常产品”“有缺陷产品”等文字概念判断异常，擅长语义层面的判断。"),
        ("视觉分支", "利用少量正常参考图片比较外观差异，擅长细小纹理和局部结构变化。"),
        ("图像异常分数", "用一个数字表示整张图片有多异常，分数越高通常表示越异常。"),
        ("像素异常图", "一张与原图对应的热力图，用来表示异常可能出现在什么位置。"),
        ("不确定性", "模型对自己判断的犹豫程度；接近中间状态通常表示更不确定。"),
        ("一致性", "两条分支是否在同一个样本或同一个区域上给出相近判断。"),
        ("路由器", "读取两条分支的结果，决定应该相信视觉、文本，还是进行加权融合的小模块。"),
        ("分数校准", "把不同方法的分数转换到可以公平比较的尺度，避免数值范围不同造成偏向。"),
        ("测试数据泄漏", "开发时偷看最终测试集的标签、掩码或整体统计量，会导致结果不真实。"),
        ("消融实验", "每次去掉一个模块，观察性能变化，用来证明真正有效的设计。"),
    ]
    for term, explanation in terms:
        cells = table.add_row().cells
        cells[0].text = term
        cells[1].text = explanation
        set_cell_shading(cells[0], "F6F8FB")
        for paragraph in cells[0].paragraphs:
            for run in paragraph.runs:
                set_run_font(run, size=9.5, color=DARK_BLUE, bold=True)
        for paragraph in cells[1].paragraphs:
            for run in paragraph.runs:
                set_run_font(run, size=9.5)
    set_table_widths(table, [1900, 7460])


def build() -> None:
    doc = Document()
    style_document(doc)
    add_header_footer(doc)
    add_title_block(doc)

    add_heading(doc, "一、第二阶段到底要解决什么问题", 1)
    add_para(
        doc,
        "第一阶段的任务是把不同异常检测方法公平地复现出来。第二阶段的任务是在这些已经得到的预测结果之上，设计一个新的动态融合方法。"
    )
    add_para(
        doc,
        "固定融合的做法是每次都用同样的比例，例如视觉分支占50%、文本分支占50%。动态融合则会根据当前图片和当前区域的情况，改变两条分支的权重。"
    )
    add_para(
        doc,
        "例如，视觉分支发现一条很细的裂纹，而文本分支没有明显反应，这时应该提高视觉分支权重；如果视觉分支受到纹理或光照干扰，而文本分支判断很稳定，就可以提高文本分支权重。"
    )

    add_heading(doc, "二、为什么需要两条分支", 1)
    add_heading(doc, "2.1 文本分支", 2)
    add_para(doc, "文本分支把图片和“正常”“异常”“损坏表面”等文字概念进行比较。它的优点是对新类别有一定通用能力，缺点是对很小的局部缺陷可能不够敏感。")
    add_heading(doc, "2.2 视觉分支", 2)
    add_para(doc, "视觉分支把测试图和少量正常参考图片比较。它通常更擅长发现裂纹、划痕、缺口和纹理变化，但也可能把正常的角度、光照或外观变化误认为异常。")
    add_para(doc, "因此，动态融合不是简单把两个结果相加，而是先判断哪条分支在当前情况下更可靠。")

    add_heading(doc, "三、重要术语解释", 1)
    add_term_table(doc)

    add_heading(doc, "四、动态融合的输入和输出", 1)
    add_heading(doc, "4.1 输入", 2)
    add_bullets(
        doc,
        [
            "视觉分支的图像异常分数。",
            "视觉分支的像素异常图。",
            "文本分支的图像异常分数。",
            "文本分支的像素异常图。",
            "两条分支各自的不确定性。",
            "两条分支的分数差异和异常图差异。",
            "用于对齐样本的 sample_id。",
        ],
    )
    add_heading(doc, "4.2 输出", 2)
    add_bullets(
        doc,
        [
            "融合后的图像异常分数。",
            "融合后的像素异常图。",
            "图像级视觉权重和文本权重。",
            "像素级视觉权重和文本权重。",
            "路由决定：视觉主导、文本主导或加权融合。",
            "用于分析的置信度和一致性记录。",
        ],
    )
    add_para(
        doc,
        "当前项目已经建立了第一版 prediction-only 接口。prediction-only 的意思是：路由器只接收模型预测，不接收测试标签或测试掩码。"
    )

    add_heading(doc, "五、第二阶段的核心执行步骤", 1)
    add_heading(doc, "第1步：统一两条分支的格式", 2)
    add_numbers(
        doc,
        [
            "确认两条分支使用完全相同的测试样本。",
            "确认 sample_id 能够一一对应。",
            "确认分数方向一致：分数越高都表示越异常。",
            "确认像素异常图的尺寸和坐标含义一致。",
            "保存每条分支的原始结果，不覆盖第一阶段结果。",
        ],
    )
    add_heading(doc, "第2步：建立固定融合基线", 2)
    add_para(doc, "先建立最简单的参照结果：")
    add_bullets(doc, ["视觉50% + 文本50%。", "视觉70% + 文本30%。", "视觉30% + 文本70%。"])
    add_para(doc, "固定融合的作用是提供一个清楚的比较基准。如果动态融合连固定融合都不能超过，就说明路由器没有产生实际价值。")
    add_heading(doc, "第3步：做分数校准", 2)
    add_para(doc, "不同方法的分数范围可能完全不同。比如一个方法的分数在0.1到0.6之间，另一个方法的分数在0.01到0.05之间，不能直接比较大小。")
    add_para(doc, "校准只能使用允许的数据：源域验证结果、目标域正常参考样本或训练阶段允许使用的正常图片。不能使用目标测试图片的标签、掩码或整个测试集的均值和标准差。")
    add_heading(doc, "第4步：计算不确定性和一致性", 2)
    add_bullets(
        doc,
        [
            "比较两条分支的图像分数差异。",
            "比较两条分支异常图的高分区域是否重合。",
            "观察异常图是否集中在少数区域，还是全图都很亮。",
            "计算每条分支的图像级和像素级不确定性。",
            "记录两条分支意见冲突的样本，作为后续分析重点。",
        ],
    )
    add_heading(doc, "第5步：先做图像级动态路由", 2)
    add_para(doc, "第一版不要直接训练复杂的神经网络。先用容易解释的规则：视觉更确定时视觉主导，文本更确定时文本主导，两者接近时进行加权融合。")
    add_heading(doc, "第6步：再做像素级动态权重", 2)
    add_para(doc, "当图像级路由确认有效后，再让每个像素使用不同的视觉/文本权重。像素级权重必须限制在0到1之间，并避免相邻像素出现剧烈跳变。")
    add_heading(doc, "第7步：做消融实验", 2)
    add_numbers(
        doc,
        [
            "只用视觉分支。",
            "只用文本分支。",
            "固定50%融合。",
            "固定比例的最好结果。",
            "只使用不确定性的动态融合。",
            "不确定性加一致性的动态融合。",
            "图像级动态融合。",
            "像素级动态融合。",
            "去掉分数校准。",
            "去掉像素平滑约束。",
        ],
    )

    add_heading(doc, "六、数据使用和防止泄漏的规则", 1)
    add_heading(doc, "6.1 开发阶段", 2)
    add_para(doc, "开发阶段只使用 VisA seed 0 的冻结预测缓存和合成数据。此阶段可以调整接口、特征、校准方法和少量超参数。")
    add_heading(doc, "6.2 设计锁定", 2)
    add_para(doc, "当输入特征、校准规则、路由结构、权重范围和阈值确定后，就要锁定设计。锁定后不能再根据最终验证结果修改。")
    add_heading(doc, "6.3 最终验证", 2)
    add_para(doc, "设计锁定后才使用 VisA seed 1/2 和 MVTec seed 0/1/2。这样才能判断方法是否能换随机样本、换数据集后仍然有效。")
    add_heading(doc, "6.4 路由器明确禁止的信息", 2)
    add_bullets(
        doc,
        [
            "测试图像标签。",
            "测试图像掩码。",
            "类别测试标签。",
            "测试集整体均值、标准差或分位数。",
            "根据测试集结果反复选择参数。",
        ],
    )

    add_heading(doc, "七、当前项目进度", 1)
    add_status_table(doc)
    add_para(doc, "当前基线训练和动态融合设计可以同时进行。动态融合设计不改变统一划分、指标和已有结果。")

    add_heading(doc, "八、目前已经实现的设计轨内容", 1)
    add_bullets(
        doc,
        [
            "统一输入输出接口。",
            "图像级和像素级不确定性特征。",
            "图像分数和异常图一致性特征。",
            "确定性置信度路由器。",
            "图像级和像素级动态权重。",
            "动态融合配置文件。",
            "合成数据单元测试。",
            "NaN、Inf、形状不匹配和极端数值检查。",
            "冻结 VisA seed 0 缓存的无训练烟雾测试。",
        ],
    )
    add_para(doc, "烟雾测试已经证明接口可以运行，但也发现两个方法的原始分数尺度不同，直接计算不确定性会造成路由偏置。因此下一步首先要做分数校准，暂时不能把烟雾测试结果当作性能提升。")

    add_heading(doc, "九、后续任务和每一步的目的", 1)
    add_status_table(doc)
    add_numbers(
        doc,
        [
            "继续完成基线轨的 PromptAD seed 1/2 训练，得到稳定的基准结果。",
            "在 VisA seed 0 上完成分数校准和固定融合比较。",
            "完成图像级规则路由，并记录每类样本更依赖哪条分支。",
            "如果图像级路由有效，再实现像素级动态权重。",
            "完成消融实验和失败案例分析。",
            "锁定动态融合设计，不再修改结构和主要超参数。",
            "使用 VisA seed 1/2 和 MVTec 进行最终验证。",
            "最后整理主结果表、逐类别结果、效率表和论文图。",
        ],
    )

    add_heading(doc, "十、最终需要交付的结果", 1)
    add_bullets(
        doc,
        [
            "一份可重复运行的动态融合代码。",
            "一份不使用测试真值的训练/验证协议。",
            "文本分支、视觉分支、固定融合和动态融合的主结果表。",
            "包含分数校准、不确定性、一致性和像素权重的消融表。",
            "路由权重和典型成功/失败案例可视化。",
            "不同类别、不同shot和不同seed下的稳定性分析。",
            "对测试数据泄漏、分数尺度差异和失败模式的说明。",
        ],
    )

    add_heading(doc, "十一、当前最重要的判断", 1)
    add_para(
        doc,
        "现在最重要的不是马上增加更复杂的网络，而是先把两个分支的分数校准做好。只有在同一尺度上比较“不确定性”，路由器才有可能做出合理的选择。基线轨可以继续耗时训练，设计轨则可以在不占用GPU的情况下完成接口、校准规则、固定融合和单元测试。"
    )

    doc.core_properties.title = "项目第二阶段设计规划：动态融合算法"
    doc.core_properties.subject = "少样本工业异常检测项目"
    doc.core_properties.author = "Codex"
    doc.core_properties.keywords = "少样本, 工业异常检测, 动态融合, PromptAD, VisA, MVTec"
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()

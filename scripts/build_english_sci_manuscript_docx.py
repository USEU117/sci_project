from __future__ import annotations

import csv
import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips
import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt  # noqa: E402
from matplotlib.patches import FancyArrowPatch, FancyBboxPatch, Rectangle  # noqa: E402


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "outputs" / "paper_draft_20260810"
ASSET_DIR = OUT_DIR / "assets_english"
DOCX_PATH = OUT_DIR / "Leakage-Safe_Uncertainty_Routing_English_SCI_Draft_V0.3.docx"

DOC_SKILL = Path(
    r"C:\Users\lynle\.codex\plugins\cache\openai-primary-runtime\documents"
    r"\26.805.11740\skills\documents"
)
sys.path.insert(0, str(DOC_SKILL / "scripts"))
from table_geometry import (  # noqa: E402
    apply_table_geometry,
    column_widths_from_weights,
    section_content_width_dxa,
)


# Journal-neutral manuscript override built on the narrative_proposal preset.
# It deliberately uses restrained black/gray typography, a single-column review
# layout, numbered sections, and compact data displays that can later be moved
# into a target journal's Word or LaTeX template.
BLACK = RGBColor(24, 24, 24)
DARK_GRAY = "404040"
MID_GRAY = "D9D9D9"
LIGHT_GRAY = "F2F2F2"
PALE_GRAY = "F8F8F8"
NOTE_FILL = "FFF2CC"
ACCENT = "1F4E79"


def set_run_font(run, name: str = "Times New Roman", size: float | None = None,
                 bold: bool | None = None, italic: bool | None = None,
                 color: str | None = None) -> None:
    run.font.name = name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), name)
    r_fonts.set(qn("w:hAnsi"), name)
    r_fonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, value: str, *, bold: bool = False,
                  size: float = 8.0, align=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(str(value))
    set_run_font(run, size=size, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    separate.append(text)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, end])
    set_run_font(run, size=8.5, color="666666")


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.4)
    section.footer_distance = Inches(0.4)
    section.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.18
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.widow_control = True

    body = styles["Body Text"]
    body.font.name = "Times New Roman"
    body._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    body._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    body._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    body.font.size = Pt(10.5)
    body.font.color.rgb = BLACK
    body.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    body.paragraph_format.first_line_indent = Twips(0)
    body.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    body.paragraph_format.line_spacing = 1.18
    body.paragraph_format.space_after = Pt(5)
    body.paragraph_format.widow_control = True

    for name, size, before, after in [
        ("Heading 1", 13.5, 14, 6),
        ("Heading 2", 11.5, 10, 4),
        ("Heading 3", 10.5, 8, 3),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string("000000")
        style.paragraph_format.left_indent = Twips(0)
        style.paragraph_format.right_indent = Twips(0)
        style.paragraph_format.first_line_indent = Twips(0)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.page_break_before = False

    title = styles["Title"]
    title.font.name = "Arial"
    title._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    title._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    title.font.size = Pt(18)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string("000000")
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(4)
    title.paragraph_format.space_after = Pt(8)
    title_ppr = title._element.get_or_add_pPr()
    title_border = title_ppr.find(qn("w:pBdr"))
    if title_border is not None:
        title_ppr.remove(title_border)

    caption = styles["Caption"]
    caption.font.name = "Times New Roman"
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    caption.font.size = Pt(8.7)
    caption.font.color.rgb = RGBColor.from_string(DARK_GRAY)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(5)

    for style_name in ["List Bullet", "List Number"]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.32)
        style.paragraph_format.first_line_indent = Inches(-0.18)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.line_spacing = 1.15

    if "Reference" not in styles:
        ref = styles.add_style("Reference", WD_STYLE_TYPE.PARAGRAPH)
    else:
        ref = styles["Reference"]
    ref.font.name = "Times New Roman"
    ref._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    ref._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    ref.font.size = Pt(8.8)
    ref.paragraph_format.left_indent = Inches(0.25)
    ref.paragraph_format.first_line_indent = Inches(-0.25)
    ref.paragraph_format.space_after = Pt(2.5)
    ref.paragraph_format.line_spacing = 1.05

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = header.add_run("Leakage-Safe Uncertainty Routing for Few-Shot Industrial Anomaly Detection")
    set_run_font(hr, size=8.0, color="777777")
    footer = section.footer.paragraphs[0]
    add_page_field(footer)
    first_footer = section.first_page_footer.paragraphs[0]
    add_page_field(first_footer)

    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def add_body(doc: Document, text: str, *, bold_lead: str | None = None) -> None:
    paragraph = doc.add_paragraph(style="Body Text")
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        set_run_font(lead, bold=True)
        rest = paragraph.add_run(text[len(bold_lead):])
        set_run_font(rest)
    else:
        run = paragraph.add_run(text)
        set_run_font(run)


def add_numbered_item(doc: Document, text: str, *, bold_lead: str | None = None) -> None:
    paragraph = doc.add_paragraph(style="List Number")
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        set_run_font(lead, bold=True)
        rest = paragraph.add_run(text[len(bold_lead):])
        set_run_font(rest)
    else:
        run = paragraph.add_run(text)
        set_run_font(run)


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    run = paragraph.add_run(text)
    set_run_font(run)


def add_equation(doc: Document, expression: str, number: str) -> None:
    table = doc.add_table(rows=1, cols=2)
    left = table.cell(0, 0)
    right = table.cell(0, 1)
    left.text = ""
    lp = left.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lr = lp.add_run(expression)
    set_run_font(lr, name="Cambria Math", size=10.5)
    right.text = ""
    rp = right.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rr = rp.add_run(number)
    set_run_font(rr, size=9.0)
    widths = column_widths_from_weights([9.2, 0.8], section_content_width_dxa(doc.sections[-1]))
    apply_table_geometry(
        table,
        widths,
        indent_dxa=0,
        cell_margins_dxa={"top": 10, "bottom": 10, "start": 0, "end": 0},
    )
    for cell in table.rows[0].cells:
        tc_pr = cell._tc.get_or_add_tcPr()
        borders = OxmlElement("w:tcBorders")
        for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), "nil")
            borders.append(el)
        tc_pr.append(borders)


def add_table_caption(doc: Document, number: int, title: str) -> None:
    paragraph = doc.add_paragraph(style="Caption")
    paragraph.paragraph_format.keep_with_next = True
    r1 = paragraph.add_run(f"Table {number}. ")
    set_run_font(r1, size=8.7, bold=True)
    r2 = paragraph.add_run(title)
    set_run_font(r2, size=8.7)


def add_figure_caption(doc: Document, number: int, title: str) -> None:
    paragraph = doc.add_paragraph(style="Caption")
    r1 = paragraph.add_run(f"Fig. {number}. ")
    set_run_font(r1, size=8.7, bold=True)
    r2 = paragraph.add_run(title)
    set_run_font(r2, size=8.7)


def add_table(doc: Document, number: int, title: str, headers: list[str],
              rows: list[list[str]], weights: list[float], *,
              footnote: str | None = None, font_size: float = 7.8,
              left_columns: set[int] | None = None) -> None:
    add_table_caption(doc, number, title)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = 0
    left_columns = left_columns or set()
    for index, header in enumerate(headers):
        set_cell_text(table.cell(0, index), header, bold=True, size=font_size)
        set_cell_shading(table.cell(0, index), MID_GRAY)
    set_repeat_table_header(table.rows[0])
    for row_index, values in enumerate(rows, start=1):
        cells = table.add_row().cells
        for col_index, value in enumerate(values):
            align = WD_ALIGN_PARAGRAPH.LEFT if col_index in left_columns else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_text(cells[col_index], value, size=font_size, align=align)
            if row_index % 2 == 0:
                set_cell_shading(cells[col_index], PALE_GRAY)
        set_cant_split(table.rows[-1])
    widths = column_widths_from_weights(weights, section_content_width_dxa(doc.sections[-1]))
    apply_table_geometry(
        table,
        widths,
        indent_dxa=120,
        cell_margins_dxa={"top": 65, "bottom": 65, "start": 120, "end": 120},
    )
    if footnote:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(2)
        paragraph.paragraph_format.space_after = Pt(5)
        run = paragraph.add_run(footnote)
        set_run_font(run, size=8.2, italic=True, color="555555")


def add_figure(doc: Document, path: Path, number: int, title: str,
               alt_text: str, *, width: float = 6.25) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(width))
    inline = doc.inline_shapes[-1]._inline
    inline.docPr.set("descr", alt_text)
    add_figure_caption(doc, number, title)


def add_draft_note(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, NOTE_FILL)
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("Draft note - remove or revise before submission. ")
    set_run_font(run, size=8.8, bold=True)
    body = paragraph.add_run(text)
    set_run_font(body, size=8.8)
    apply_table_geometry(
        table,
        [section_content_width_dxa(doc.sections[-1])],
        indent_dxa=120,
        cell_margins_dxa={"top": 90, "bottom": 90, "start": 120, "end": 120},
    )


def _diagram_box(ax, x: float, y: float, w: float, h: float, text: str,
                 face: str, edge: str, *, fontsize: float = 8.2,
                 linewidth: float = 1.4, linestyle: str = "-") -> None:
    patch = FancyBboxPatch(
        (x, y), w, h,
        boxstyle="round,pad=0.018,rounding_size=0.025",
        facecolor=face, edgecolor=edge, linewidth=linewidth, linestyle=linestyle,
    )
    ax.add_patch(patch)
    ax.text(x + w / 2, y + h / 2, text, ha="center", va="center",
            fontsize=fontsize, color="#202020", linespacing=1.15)


def _diagram_arrow(ax, start: tuple[float, float], end: tuple[float, float],
                   *, color: str = "#4F81BD", linestyle: str = "-",
                   connectionstyle: str = "arc3") -> None:
    ax.add_patch(FancyArrowPatch(
        start, end, arrowstyle="-|>", mutation_scale=11, linewidth=1.35,
        color=color, linestyle=linestyle, connectionstyle=connectionstyle,
    ))


def _save_figure(fig, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(path, dpi=300, bbox_inches="tight", facecolor="white")
    fig.savefig(path.with_suffix(".svg"), bbox_inches="tight", facecolor="white")
    plt.close(fig)


def create_english_architecture_figure(path: Path) -> None:
    """Create a leakage-explicit V0.3 architecture figure.

    Normal-only preparation and per-query inference are shown as separate lanes.
    Ground truth is isolated inside the evaluator and has no return edge.
    """
    fig, ax = plt.subplots(figsize=(12.6, 7.45))
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")
    ax.text(
        0.5, 0.965,
        "Leakage-safe two-branch evidence routing: preparation, query inference, and evaluation",
        ha="center", va="center", fontsize=13.4, fontweight="bold", color="#17365D",
    )

    # Lane backgrounds make the source of every statistic visible.
    ax.add_patch(Rectangle((0.015, 0.63), 0.97, 0.27, facecolor="#F4FAF4",
                           edgecolor="#70AD47", linewidth=1.2))
    ax.add_patch(Rectangle((0.015, 0.23), 0.97, 0.34, facecolor="#F7FAFD",
                           edgecolor="#5B9BD5", linewidth=1.2))
    ax.add_patch(Rectangle((0.66, 0.025), 0.325, 0.14, facecolor="#FFF4F2",
                           edgecolor="#C00000", linewidth=1.2, linestyle="--"))
    ax.text(0.025, 0.875, "A  Normal-only preparation (offline)", fontsize=10.2,
            fontweight="bold", color="#385723", va="top")
    ax.text(0.025, 0.545, "B  Query inference (one image at a time)", fontsize=10.2,
            fontweight="bold", color="#1F4E79", va="top")
    ax.text(0.67, 0.145, "C  Evaluation only", fontsize=9.5,
            fontweight="bold", color="#9C0006", va="top")

    # Offline preparation lane.
    _diagram_box(ax, 0.035, 0.69, 0.12, 0.11, "K normal\nreferences", "#E2F0D9", "#548235")
    _diagram_box(ax, 0.205, 0.675, 0.16, 0.14,
                 "Frozen branches\nreference outputs\n+ deterministic views", "#EAF2F8", "#4472C4", fontsize=7.8)
    _diagram_box(ax, 0.415, 0.675, 0.18, 0.14,
                 "Branch-wise normal\ncalibrators and\nsupport ranges", "#FFF2CC", "#BF9000", fontsize=7.9)
    _diagram_box(ax, 0.645, 0.69, 0.15, 0.11,
                 "Frozen reference\npackage", "#E2F0D9", "#548235")
    _diagram_box(ax, 0.84, 0.675, 0.125, 0.14,
                 "Audit fields\nhashes; test labels\nused = false", "#FCE4D6", "#C65911", fontsize=7.5)
    _diagram_arrow(ax, (0.155, 0.745), (0.205, 0.745), color="#548235")
    _diagram_arrow(ax, (0.365, 0.745), (0.415, 0.745), color="#548235")
    _diagram_arrow(ax, (0.595, 0.745), (0.645, 0.745), color="#548235")
    _diagram_arrow(ax, (0.795, 0.745), (0.84, 0.745), color="#548235")

    # Query lane.
    _diagram_box(ax, 0.035, 0.335, 0.09, 0.10, "Query\nimage x", "#D9EAF7", "#4472C4")
    _diagram_box(ax, 0.16, 0.39, 0.145, 0.105,
                 "AnomalyDINO\n(s_v, M_v)", "#D9EAF7", "#4472C4")
    _diagram_box(ax, 0.16, 0.265, 0.145, 0.105,
                 "AnomalyCLIP\n(s_t, M_t)", "#FCE4D6", "#C65911")
    _diagram_box(ax, 0.35, 0.315, 0.15, 0.13,
                 "Sample/geometry\nalignment + frozen\ncalibration", "#FFF2CC", "#BF9000", fontsize=7.8)
    _diagram_box(ax, 0.545, 0.315, 0.145, 0.13,
                 "V1 reliability\nentropy, conflict,\nconcentration", "#E4DFEC", "#7030A0", fontsize=7.8)
    _diagram_box(ax, 0.735, 0.39, 0.10, 0.105,
                 "Image router\nweight w_I", "#E2F0D9", "#548235", fontsize=7.8)
    _diagram_box(ax, 0.735, 0.265, 0.10, 0.105,
                 "Pixel router\nmap W_P(x)", "#E2F0D9", "#548235", fontsize=7.8)
    _diagram_box(ax, 0.88, 0.39, 0.09, 0.105,
                 "Fused score\ns_f", "#D9EAF7", "#4472C4", fontsize=7.8)
    _diagram_box(ax, 0.88, 0.265, 0.09, 0.105,
                 "Fused map\nM_f", "#FCE4D6", "#C65911", fontsize=7.8)
    _diagram_arrow(ax, (0.125, 0.385), (0.16, 0.442))
    _diagram_arrow(ax, (0.125, 0.385), (0.16, 0.317))
    _diagram_arrow(ax, (0.305, 0.442), (0.35, 0.405))
    _diagram_arrow(ax, (0.305, 0.317), (0.35, 0.35))
    _diagram_arrow(ax, (0.50, 0.38), (0.545, 0.38))
    _diagram_arrow(ax, (0.69, 0.405), (0.735, 0.442))
    _diagram_arrow(ax, (0.69, 0.35), (0.735, 0.317))
    _diagram_arrow(ax, (0.835, 0.442), (0.88, 0.442))
    _diagram_arrow(ax, (0.835, 0.317), (0.88, 0.317))
    _diagram_arrow(ax, (0.72, 0.69), (0.425, 0.445), color="#548235",
                   connectionstyle="arc3,rad=0.20")

    # Evaluation lane: predictions and labels meet only here.
    _diagram_box(ax, 0.675, 0.055, 0.115, 0.065,
                 "Predictions\n(s_f, M_f)", "#EAF2F8", "#4472C4", fontsize=7.4)
    _diagram_box(ax, 0.82, 0.055, 0.145, 0.065,
                 "Labels + masks\nmetrics only", "#F4CCCC", "#C00000", fontsize=7.4)
    _diagram_arrow(ax, (0.925, 0.265), (0.735, 0.12), color="#4472C4",
                   connectionstyle="arc3,rad=-0.15")
    _diagram_arrow(ax, (0.79, 0.087), (0.82, 0.087), color="#C00000")
    ax.text(
        0.33, 0.105,
        "Leakage boundary: no test label, mask, cross-query memory, or aggregate test statistic\n"
        "is available to reference preparation, calibration, reliability estimation, or routing.",
        ha="center", va="center", fontsize=8.5, color="#9C0006", fontweight="bold",
    )
    _save_figure(fig, path)


def create_v2_roadmap_figure(path: Path) -> None:
    """Create a prospective V2 roadmap that is explicitly separated from V1 evidence."""
    fig, ax = plt.subplots(figsize=(12.4, 4.3))
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")
    ax.text(0.5, 0.93, "DynamicFusion V2: prospective safety-oriented development path",
            ha="center", fontsize=13.2, fontweight="bold", color="#17365D")
    ax.text(0.5, 0.845, "PLANNED DESIGN - not implemented or evaluated in the V1 results",
            ha="center", fontsize=9.5, fontweight="bold", color="#9C0006")
    boxes = [
        (0.025, "V1 evidence\nsaturation; ties;\nweak gate-error link", "#F4CCCC", "#C00000"),
        (0.225, "Rank-preserving\nnormal-only\ncalibration", "#FFF2CC", "#BF9000"),
        (0.425, "Out-of-support\nand saturation\ndiagnostics", "#E4DFEC", "#7030A0"),
        (0.625, "Safe visual default\nand guarded text\nintervention", "#E2F0D9", "#548235"),
        (0.825, "Separate image/pixel\nrouters + new\ndev/holdout test", "#D9EAF7", "#4472C4"),
    ]
    for x, label, face, edge in boxes:
        _diagram_box(ax, x, 0.38, 0.15, 0.30, label, face, edge, fontsize=8.4)
    for i in range(len(boxes) - 1):
        _diagram_arrow(ax, (boxes[i][0] + 0.15, 0.53), (boxes[i + 1][0], 0.53))
    ax.text(
        0.5, 0.16,
        "Acceptance sequence: ranking retained -> saturation bounded -> unsafe cases fall back -> "
        "weights track branch advantage -> frozen independent validation",
        ha="center", va="center", fontsize=8.9, color="#404040",
    )
    _save_figure(fig, path)


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        return list(csv.DictReader(handle))


def pct(value: str) -> str:
    return f"{float(value) * 100:.2f}"


def build_doc() -> Path:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    architecture = ASSET_DIR / "leakage_safe_uncertainty_routing_architecture_v03.png"
    v2_roadmap = ASSET_DIR / "dynamic_fusion_v2_roadmap_v03.png"
    create_english_architecture_figure(architecture)
    create_v2_roadmap_figure(v2_roadmap)

    visa_rows = read_csv(ROOT / "experiments" / "summaries" / "visa_baseline_main_table_20260803.csv")
    mvtec_rows = read_csv(ROOT / "experiments" / "summaries" / "mvtec_paper_main_table_template_20260809.csv")
    ablation_rows = read_csv(
        ROOT / "experiments" / "summaries" / "dynamic_fusion_scientific_analysis_20260809" / "ablation_summary.csv"
    )
    with (ROOT / "outputs" / "dynamic_fusion" / "final_validation" / "summary.json").open(
        "r", encoding="utf-8"
    ) as handle:
        visa_dynamic = json.load(handle)

    fig_dir = ROOT / "outputs" / "dynamic_fusion" / "figures" / "20260809_scientific_analysis"

    doc = Document()
    configure_document(doc)
    props = doc.core_properties
    props.title = "When Should Visual and Language Evidence Be Fused? A Leakage-Safe Study of Uncertainty Routing for Few-Shot Industrial Anomaly Detection"
    props.subject = "English SCI-style manuscript draft V0.3"
    props.author = ""
    props.keywords = (
        "industrial anomaly detection; few-shot learning; vision-language model; "
        "uncertainty routing; score calibration; defect localization"
    )

    # Title page and abstract: compact journal style, not a thesis cover.
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(0)
    meta.paragraph_format.space_after = Pt(5)
    meta_run = meta.add_run("ORIGINAL RESEARCH ARTICLE | ENGLISH MANUSCRIPT DRAFT V0.3")
    set_run_font(meta_run, name="Arial", size=8.5, bold=True, color=ACCENT)

    title = doc.add_paragraph(style="Title")
    title_ppr = title._p.get_or_add_pPr()
    title_border = title_ppr.find(qn("w:pBdr"))
    if title_border is not None:
        title_ppr.remove(title_border)
    title.add_run(
        "When Should Visual and Language Evidence Be Fused? "
        "A Leakage-Safe Study of Uncertainty Routing for Few-Shot Industrial Anomaly Detection"
    )

    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author.paragraph_format.space_after = Pt(2)
    ar = author.add_run("Author information to be completed")
    set_run_font(ar, size=10.5, italic=True, color="555555")
    affiliation = doc.add_paragraph()
    affiliation.alignment = WD_ALIGN_PARAGRAPH.CENTER
    affiliation.paragraph_format.space_after = Pt(9)
    afr = affiliation.add_run("Affiliation, postal address, ORCID, and corresponding-author email to be completed before submission")
    set_run_font(afr, size=8.8, color="666666")

    abstract_heading = doc.add_paragraph()
    abstract_heading.paragraph_format.space_before = Pt(4)
    abstract_heading.paragraph_format.space_after = Pt(2)
    ah = abstract_heading.add_run("Abstract")
    set_run_font(ah, name="Arial", size=10.5, bold=True)
    abstract = (
        "Few-shot industrial anomaly detection must classify defective images and localize defects from only a few normal "
        "references. Visual nearest-neighbor models capture fine structural deviations, whereas vision-language models "
        "supply transferable semantic evidence. Their anomaly scores, however, have different meanings and distributions; "
        "sample-dependent fusion can therefore damage the ranking of an already strong branch. This study establishes an "
        "auditable 1-, 2-, and 4-shot protocol on MVTec AD and VisA with three deterministic seeds, nested reference sets, "
        "a common prediction schema, and explicit leakage controls. A frozen first-generation router combines AnomalyDINO "
        "and AnomalyCLIP through separate image- and pixel-level weights. Calibration uses only normal-reference predictions; "
        "test labels, masks, cross-query memory, and aggregate test statistics are excluded. AnomalyDINO achieves VisA image "
        "AUROC values of 89.40%, 91.40%, and 92.58% at 1, 2, and 4 shots. On the VisA seed-0 development split, separate "
        "routing temperatures improve AUPRO over the best fixed-weight control by 4.14 and 8.82 percentage points at 2 and "
        "4 shots. The frozen router nevertheless reduces MVTec image AUROC relative to raw AnomalyDINO by 16.29, 10.49, "
        "and 7.94 points. The failure is traced to median/MAD calibration saturation: 99.99% of visual scores are at least "
        "0.999 on average, after which binary entropy mistakes numerical extremeness for confidence and the router no longer "
        "tracks branch advantage. The evidence supports a conditional conclusion, not universal fusion superiority: "
        "reliable routing requires rank-preserving calibration, out-of-support detection, a safe strong-branch fallback, and "
        "independent validation on a new development boundary."
    )
    ap = doc.add_paragraph(style="Body Text")
    ap.paragraph_format.space_after = Pt(4)
    ap_run = ap.add_run(abstract)
    set_run_font(ap_run, size=9.4)

    keywords = doc.add_paragraph()
    keywords.paragraph_format.space_after = Pt(8)
    kw1 = keywords.add_run("Keywords: ")
    set_run_font(kw1, size=9.2, bold=True)
    kw2 = keywords.add_run(
        "industrial anomaly detection; few-shot learning; vision-language model; uncertainty routing; "
        "score calibration; defect localization"
    )
    set_run_font(kw2, size=9.2)

    # 1 Introduction
    doc.add_heading("1 Introduction", level=1)
    add_body(
        doc,
        "Automated visual inspection is expected to detect rare defects while production conditions, object pose, "
        "illumination, material appearance, and camera settings continue to change. In many factories, abundant normal "
        "samples are available but representative defect samples are not. A practical detector must therefore learn the "
        "normal appearance of a product and identify deviations without assuming that future defect types have already "
        "been annotated. MVTec AD [1] and VisA [2] formalize this setting at both image level, where a complete image is "
        "classified as normal or anomalous, and pixel level, where the defective region must be localized. Recent surveys "
        "organize this field by supervision, architecture, benchmark, and deployment constraints [3], including the "
        "particularly difficult low-shot setting considered here [4]."
    )
    add_body(
        doc,
        "The few-shot setting is especially demanding. Only K normal reference images, with K commonly equal to 1, 2, "
        "or 4, are available for each target category. A method must cover legitimate normal variation from this small "
        "sample while remaining sensitive to subtle scratches, contamination, missing components, deformation, and "
        "texture changes. PatchCore [5] and related visual-memory methods compare local test features with a bank of normal "
        "features. Foundation encoders such as DINOv2 [6] have strengthened this paradigm, and AnomalyDINO [7] shows that "
        "high-quality frozen patch features can be highly competitive without target-domain backbone training. FastRecon "
        "[8], UniVAD [10], SubspaceAD [11], and FastRef [12] further show that reconstruction, unified frozen features, "
        "compact subspaces, and query-conditioned prototype refinement are strong alternatives to a complex fusion system."
    )
    add_body(
        doc,
        "Vision-language models offer a second source of evidence. CLIP [13] aligns image and text representations. "
        "WinCLIP [9], AnomalyCLIP [14], PromptAD [15], AdaCLIP [16], VCP-CLIP [17], FiLo [18], and AA-CLIP [19] adapt "
        "prompts, image context, or patch-text alignment to industrial anomalies. Recent systems also combine retrieval, "
        "prompt interaction, DINO features, or text-guided residuals, including ReMP-AD [20], AdaptCLIP [21], AnoPLe [22], "
        "DLVP-CLIP [23], PAPL [24], and TGRF-CLIP [25]. This progress makes multimodal fusion timely, but it also sharpens "
        "the question addressed in this paper: when is semantic evidence reliable enough to modify a strong visual result?"
    )
    add_body(
        doc,
        "Two reliability problems are often hidden by a final mean score. First, branch outputs differ in range, density, "
        "spatial resolution, and probabilistic meaning. A calibration transform can make values numerically comparable "
        "while still destroying the rank information needed by AUROC. Second, a router can accidentally use test labels, "
        "masks, or aggregate test-set statistics when selecting thresholds or fusion weights. Such leakage makes a result "
        "look adaptive but removes its independent validation meaning. The problem addressed in this paper is therefore "
        "not merely how to define a weighted sum, but how to build, audit, and diagnose a visual-language router under "
        "normal-only supervision."
    )
    add_body(doc, "The contributions of this study are fourfold.")
    add_bullet(
        doc,
        "A unified and auditable 1/2/4-shot benchmark is established on MVTec AD and VisA. Nested reference manifests, "
        "three deterministic seeds, method-independent prediction files, and category/sample/schema checks make the "
        "comparison reproducible and expose incomplete result cells rather than filling them with incompatible values."
    )
    add_bullet(
        doc,
        "A leakage-safe fusion interface is implemented for frozen AnomalyDINO and AnomalyCLIP predictions. It aligns "
        "sample identifiers and pixel-map geometry, calibrates branches from normal references, and produces separate "
        "image-level and pixel-level weights."
    )
    add_bullet(
        doc,
        "A frozen two-temperature router is evaluated with single-branch and fixed-weight controls. The experiments "
        "distinguish local segmentation gains from global ranking losses and keep the VisA seed-0 development evidence "
        "separate from independent validation."
    )
    add_bullet(
        doc,
        "A systematic failure analysis connects numerical saturation, entropy misinterpretation, weak weight-advantage "
        "correlation, category-level degradation, and qualitative heatmaps. This analysis yields concrete requirements "
        "for a safer second-generation router."
    )
    add_body(
        doc,
        "The principal claim is deliberately bounded. The current router demonstrates a reproducible mechanism and local "
        "pixel-level benefits, but it does not provide universal improvement over the strongest visual branch. Reporting "
        "this limitation is essential: dynamic fusion is useful only when its reliability estimate is related to actual "
        "branch error, its calibration retains the information used by the metric, and unsafe interventions can be rejected."
    )

    # 2 Related work
    doc.add_heading("2 Related Work", level=1)
    doc.add_heading("2.1 Few-shot visual anomaly detection", level=2)
    add_body(
        doc,
        "Normal-only anomaly detection models the distribution of defect-free appearance and assigns large anomaly "
        "scores to deviations. PaDiM [26] estimates a multivariate Gaussian distribution at each spatial position, while "
        "PatchCore [5] stores representative normal patch features and uses nearest-neighbor distance for detection. "
        "Reverse distillation [27] reconstructs one-class embeddings through a student network. SimpleNet [28] introduces "
        "a lightweight feature adapter and synthetic feature anomalies, and EfficientAD [29] emphasizes millisecond-level "
        "latency. These methods illustrate a broad trade-off among memory, training cost, localization detail, and "
        "deployment speed."
    )
    add_body(
        doc,
        "In the extreme few-shot regime, the normal reference bank is small enough that coverage of legitimate variation "
        "becomes a dominant issue. FastRecon [8] reconstructs query features from few normal supports, whereas AnomalyDINO "
        "[7] uses frozen DINOv2 [6] patches and deep nearest-neighbor matching. UniVAD [10] studies a training-free unified "
        "model across anomaly domains. SubspaceAD [11] uses principal subspaces fitted to normal patches, providing a "
        "particularly important 2026 visual-only challenge to the necessity of multimodal complexity. FastRef [12] refines "
        "normal prototypes at inference while constraining anomaly absorption. These approaches motivate two fairness "
        "distinctions used below: fitting a normal reference representation is not target-domain backbone training, and "
        "single-query adaptation is not equivalent to using aggregate statistics from the complete test set."
    )

    doc.add_heading("2.2 Vision-language anomaly adaptation", level=2)
    add_body(
        doc,
        "CLIP [13] maps images and natural-language descriptions into a common representation space. WinCLIP [9] adapts "
        "this representation through compositional prompts and multi-scale windows; WinCLIP+ additionally uses a few "
        "normal images. AnomalyCLIP [14] learns object-agnostic normal and anomalous prompts, while PromptAD [15] learns "
        "prompts from target normal samples. AdaCLIP [16] and VCP-CLIP [17] introduce adaptive prompts and visual-context "
        "prompting. FiLo [18] and AA-CLIP [19] strengthen fine-grained defect language and anomaly-aware patch alignment. "
        "AnomalyGPT [30] extends anomaly analysis to large vision-language models and natural-language interaction, although "
        "its training and output setting differs from the frozen few-normal-shot comparison studied here."
    )
    add_body(
        doc,
        "The closest methods fuse information at different locations and under different supervision. ReMP-AD [20] fuses "
        "retrieved normal tokens with visual-language priors, whereas this study routes already frozen branch outputs. "
        "AdaptCLIP [21] and AnoPLe [22] learn adapters or bidirectional prompts, so their base-data training and target-normal "
        "use must be labeled separately. DLVP-CLIP [23] improves local visual prompting. PAPL [24] performs upstream CLIP-"
        "DINO modeling in a zero-shot framework, and TGRF-CLIP [25] uses trained text-guided residual fusion. These methods "
        "are relevant conceptual neighbors, but they are not interchangeable with a post-hoc, normal-only, frozen-output "
        "router. Citation therefore does not imply protocol-matched reproduction."
    )

    doc.add_heading("2.3 Dynamic and uncertainty-aware fusion", level=2)
    add_body(
        doc,
        "Dynamic fusion aims to reduce the influence of an unreliable modality for each sample. Multimodal Dynamics [31] "
        "models sample-dependent information before fusion. Quality-aware multimodal fusion (QMF) [32] gives a more direct "
        "condition: estimated modality quality must be inversely related to generalization error for dynamic weighting to "
        "reliably beat static fusion. MoECLIP [33] dynamically routes image patches to specialized experts, but it routes "
        "experts inside a zero-shot CLIP model rather than choosing between frozen visual and text anomaly branches."
    )
    add_body(
        doc,
        "Industrial anomaly methods also expose the importance of heterogeneous score distributions. The histogram-based "
        "fusion in PGAD [34] explicitly aligns global and local evidence under one-shot scarcity. Bias correction for "
        "language-image similarity scores [35] likewise shows that numerical comparability cannot be assumed. Together, "
        "these studies support the motivation for routing "
        "but do not guarantee that a particular uncertainty proxy is useful. In the present experiments, the weak "
        "correlation between routing weight and true branch advantage violates the practical condition highlighted by QMF "
        "[32] and explains why changing weights alone is not evidence of successful routing."
    )

    doc.add_heading("2.4 Calibration and the information-use boundary", level=2)
    add_body(
        doc,
        "Temperature scaling and related post-hoc methods show that neural scores should not be treated as calibrated "
        "probabilities by default [36]. Predictive uncertainty can deteriorate under dataset shift [37], and more flexible "
        "Dirichlet mappings are not universally safe when validation evidence is scarce [38]. CADET calibrates anomaly "
        "scores to reduce hardness bias [39], while instance-wise monotonic calibration explicitly preserves "
        "within-instance order [40]. These ideas motivate the rank and saturation audits proposed for the next router, but "
        "they do not retrospectively repair V1."
    )
    add_body(
        doc,
        "The source of adaptation statistics is equally important. MuSc uses unlabeled test-image relations in a training-"
        "free framework [41], and DNPR updates normal prototypes over the test stream [42]. Such transductive or online "
        "protocols can be valid when clearly declared, but they differ from this study's inductive single-query boundary. "
        "Here, the router may access K normal references, deterministic reference views, and the current query's frozen "
        "predictions. It may not access test labels, masks, cross-query memory, or aggregate test-set statistics."
    )

    # 3 Problem formulation and protocol
    doc.add_heading("3 Problem Formulation and Evaluation Protocol", level=1)
    doc.add_heading("3.1 Few-shot anomaly detection task", level=2)
    add_body(
        doc,
        "For category c, let R_c^K = {r_1, ..., r_K} denote K normal reference images, where K is in {1, 2, 4}. The test "
        "set Q_c contains both normal and anomalous images. For a query image x, a method outputs an image-level anomaly "
        "score s(x) and a pixel-level anomaly map M(x). Larger values indicate stronger anomalous evidence. Test image "
        "labels and pixel masks are used only by the evaluator after all predictions and routing weights have been frozen."
    )
    add_body(
        doc,
        "Image-level detection and pixel-level localization are related but non-identical objectives. Image AUROC depends "
        "on the ordering of all normal and anomalous samples. Pixel metrics instead depend on within-image spatial contrast "
        "and defect coverage. Consequently, one routing temperature can preserve global ranking while over-smoothing a "
        "heatmap, or improve local coverage while changing the image ordering. The protocol evaluates both objectives "
        "separately."
    )

    doc.add_heading("3.2 Deterministic nested sampling", level=2)
    add_body(
        doc,
        "Seeds 0, 1, and 2 generate deterministic manifests for every dataset and category. Within one seed, the one-shot "
        "set is a subset of the two-shot set, and the two-shot set is a subset of the four-shot set. This nesting reduces "
        "the confounding effect of replacing all references when K changes. The VisA manifest records 252 selected paths, "
        "and the MVTec AD manifest records 315 selected paths. Both manifests pass path-existence, category, uniqueness, "
        "count, and nesting checks. Every method is required to read these manifests rather than sample independently."
    )

    doc.add_heading("3.3 Unified prediction schema and audit trail", level=2)
    add_body(
        doc,
        "Official repositories expose different output conventions. Each method is therefore adapted to export a common "
        "NPZ schema containing sample_ids, image_scores, pixel_maps, image_labels, and pixel_masks. Before evaluation or "
        "fusion, the pipeline checks category membership, sample order, label direction, map dimensions, finite values, "
        "and expected sample counts. Each run retains a unique run identifier, configuration, log, source hashes, "
        "prediction cache, and evaluation report. A directory or a completion marker alone is not accepted as evidence of "
        "a valid run."
    )
    protocol_rows = [
        ["References", "K normal images from the frozen manifest", "Allowed"],
        ["Reference augmentations", "Deterministic normal-only views", "Allowed"],
        ["Frozen branch predictions", "Direct router inputs", "Allowed"],
        ["Current query", "Single-image branch outputs and geometry", "Allowed"],
        ["Single-query adaptation", "Must reset for every query and be declared", "Conditional"],
        ["Test image labels", "Reveal the image-level answer", "Forbidden"],
        ["Test pixel masks", "Reveal the localization answer", "Forbidden"],
        ["Cross-query memory", "Transfers information between test samples", "Forbidden"],
        ["Aggregate test statistics", "Permit transductive tuning", "Forbidden"],
    ]
    add_table(
        doc, 1, "Information boundary enforced during router development and inference.",
        ["Information source", "Role", "Router access"], protocol_rows,
        [2.0, 3.2, 1.3], font_size=7.7, left_columns={0, 1}
    )

    doc.add_heading("3.4 Development and validation boundary", level=2)
    add_body(
        doc,
        "VisA seed 0 is used for router design and limited sensitivity analysis. After image temperature, pixel temperature, "
        "minimum weight, and decision margin are frozen, VisA seeds 1 and 2 are evaluated independently. MVTec AD seeds "
        "0, 1, and 2 are then processed with the frozen configuration. Because the MVTec results have now been inspected "
        "and used for failure analysis, they cannot serve as a new development set for a revised router. Any second-" 
        "generation method requires a new development/holdout boundary or must be labeled exploratory."
    )

    # 4 Method
    doc.add_heading("4 Leakage-Safe Uncertainty-Routed Fusion", level=1)
    doc.add_heading("4.1 Framework overview", level=2)
    add_body(
        doc,
        "The framework operates on frozen predictions rather than retraining the two branch backbones. A query image is "
        "processed by the visual and text-guided branches to obtain image scores and pixel maps. The same branches process "
        "the K normal references and their deterministic views. The system aligns the cached samples, estimates branch-" 
        "specific normal centers and scales, calculates uncertainty features, and produces separate image-level and pixel-" 
        "level visual weights. Figure 1 summarizes the data flow and the information boundary."
    )
    add_figure(
        doc, architecture, 1,
        "Leakage-safe data flow of the frozen two-branch V1 router.",
        "Normal-only preparation and per-query inference are separated. Ground-truth labels and masks enter only the evaluator; no arrow returns to calibration, reliability estimation, or routing."
    )

    doc.add_heading("4.2 Visual and text-guided branches", level=2)
    add_body(
        doc,
        "The visual branch is AnomalyDINO [7]. A frozen DINOv2 encoder maps query and reference images to patch features. "
        "Reference patches form a category-specific normal bank, and nearest-neighbor distances give a pixel anomaly map. "
        "An image score is obtained by aggregating local responses according to the official implementation. The visual "
        "backbone is not optimized on the target dataset, although the reference bank changes with category, seed, and K."
    )
    add_body(
        doc,
        "The text-guided branch is frozen AnomalyCLIP [14]. Normal and anomalous textual concepts are compared with global "
        "and local image representations to obtain an image score and an anomaly map. This branch supplies the same type "
        "of semantic evidence for every shot setting; it is not presented as a 1/2/4-shot baseline because it does not "
        "consume K in the same way as the visual memory bank. Its role is to provide a complementary, category-transferable "
        "signal to the router."
    )

    doc.add_heading("4.3 Normal-reference robust calibration", level=2)
    add_body(
        doc,
        "Raw distances and vision-language similarities occupy different numerical ranges. For branch b, dataset d, "
        "category c, seed, and shot, the center mu_b is the median of normal-reference responses. The robust scale is a "
        "median absolute deviation (MAD) with a numerical floor epsilon. Pixel maps contain many near-zero background "
        "values, so the pixel calibration first summarizes each deterministic reference view by its 0.99 quantile and fits "
        "the robust statistics to those view-level tails."
    )
    add_equation(doc, "delta_b = max(1.4826 median(|r_b - median(r_b)|), epsilon)", "(1)")
    add_equation(doc, "p_b(s) = sigmoid((s_b - mu_b) / (delta_b tau_c))", "(2)")
    add_body(
        doc,
        "Here tau_c is the calibration temperature and all transformed probabilities are clipped to a finite open interval "
        "to avoid logarithmic singularities. Calibration artifacts store the input source, parameter values, SHA256 hashes, "
        "and the flags test_predictions_used=false and test_labels_used=false. Although this mapping is intended to make "
        "branches comparable, Section 7 shows that an extremely small scale can cause severe saturation and ranking loss."
    )

    doc.add_heading("4.4 Uncertainty and continuous routing weights", level=2)
    add_body(
        doc,
        "The frozen V1 router uses the binary entropy of each calibrated probability as its primary uncertainty. Entropy is "
        "largest at p=0.5 and approaches zero near p=0 or p=1. The negative uncertainty values are converted into branch "
        "reliability logits and normalized by a two-branch softmax. The visual weight is clipped to [0.05, 0.95] so that "
        "neither branch is completely disabled."
    )
    add_equation(doc, "H(p) = -p log2(p) - (1 - p) log2(1 - p)", "(3)")
    add_equation(
        doc,
        "w_v = exp(-u_v / T) / [exp(-u_v / T) + exp(-u_t / T)]",
        "(4)",
    )
    add_body(
        doc,
        "A decision margin of 0.15 converts the continuous weight into visual-dominant, text-dominant, or blended labels "
        "for interpretation only; these labels do not change the fused value. The interface also records branch conflict, "
        "normal-view consistency, and spatial concentration. In V1 these auxiliary quantities are diagnostics rather than "
        "learned routing targets, which preserves a clear causal interpretation of the final weight."
    )

    doc.add_heading("4.5 Separate image-level and pixel-level routing", level=2)
    add_body(
        doc,
        "A single routing strength is not assumed to fit both tasks. The image weight controls the complete-image anomaly "
        "score and therefore the global ordering used by image AUROC. The pixel weight controls the contribution of each "
        "spatial location and therefore the contrast between defective regions and background. The frozen configuration "
        "uses image temperature T_I=0.50 and pixel temperature T_P=0.20, with minimum weight 0.05."
    )
    add_equation(doc, "s_f(x) = w_v^I p_v^I(x) + (1 - w_v^I) p_t^I(x)", "(5)")
    add_equation(doc, "M_f(x) = w_v^P(x) .* p_v^P(x) + [1 - w_v^P(x)] .* p_t^P(x)", "(6)")
    add_body(
        doc,
        "The element-wise operator in (6) permits spatially varying evidence selection. This separation is central to the "
        "study: a router may improve defect coverage without preserving image ranking, so each output must be compared with "
        "its own single-branch and fixed-weight controls."
    )

    doc.add_heading("4.6 Inference and audit sequence", level=2)
    add_body(
        doc,
        "A formal run proceeds in six ordered steps: load both branch caches; align sample identifiers, labels, masks, and "
        "map dimensions; load the passed normal-reference calibration; calculate finite reliability features; produce "
        "clipped image and pixel weights; and save fused predictions with a unique run identifier and provenance report. "
        "Any missing sample, order mismatch, non-finite value, or failed calibration terminates the run before metrics are "
        "computed. Seventeen frozen validation runs, comprising 231 category-run combinations, passed these checks. The "
        "source manifest contains 693 provenance rows and 285 unique input files with recorded SHA256 hashes."
    )

    # 5 Experimental setup
    doc.add_heading("5 Experimental Setup", level=1)
    doc.add_heading("5.1 Datasets", level=2)
    dataset_rows = [
        ["MVTec AD", "5,354", "15", "1,725", "1/2/4 shots; 3 seeds"],
        ["VisA", "10,821", "12", "2,162", "1/2/4 shots; 3 seeds"],
    ]
    add_table(
        doc, 2, "Datasets and the unified evaluation scope.",
        ["Dataset", "Official size", "Categories", "Test images", "Few-shot protocol"],
        dataset_rows, [1.3, 1.25, 0.8, 1.0, 2.15], font_size=7.7
    )
    add_body(
        doc,
        "MVTec AD contains object and texture categories with pixel-level ground truth [1]. VisA contains 12 industrial "
        "categories and more complex arrangements and backgrounds [2]. The official test splits are unchanged. Only K "
        "images are selected from each category's normal training set; anomalous test images and masks are never used to "
        "construct a reference bank or router calibration."
    )

    doc.add_heading("5.2 Baselines and fairness labels", level=2)
    add_body(
        doc,
        "PatchCore [5], WinCLIP+ [9], AnomalyDINO [7], and PromptAD [15] are evaluated under the common manifest and "
        "prediction schema. PromptAD updates prompt parameters using target-category normal images and is therefore labeled "
        "target_normal_tuning=true. The other three baselines use their respective frozen feature/reference protocols. "
        "AnomalyCLIP [14] is reported separately as a zero-shot text branch. ReMP-AD [20], AdaptCLIP [21], SubspaceAD [11], "
        "and FastRef [12] are prioritized for Gate A verification. AnoPLe [22] remains a conditional later candidate. No "
        "numbers are copied from their papers into the local main table: a method enters only after its official checkpoint, "
        "information-use boundary, prediction export, memory use, and unified metrics pass verification."
    )

    doc.add_heading("5.3 Implementation details", level=2)
    add_body(
        doc,
        "Experiments run on Windows with an NVIDIA RTX 3060 Laptop GPU with 6 GB memory. The implemented pipelines use "
        "PyTorch, not TensorFlow. AnomalyCLIP uses Python 3.10, PyTorch 2.0.0+cu118, and torchvision 0.15.1+cu118. Each "
        "official method is isolated in its own environment to prevent dependency conflicts. FAISS nearest-neighbor search "
        "for PatchCore and AnomalyDINO uses a CPU first pass where required by the memory budget. Exact source revisions, "
        "patches, manifests, command logs, and prediction hashes are retained in the project."
    )
    add_body(
        doc,
        "The fusion stage reads cached NPZ predictions and therefore adds little computation relative to branch inference. "
        "Peak GPU memory, warmed-up latency, and throughput are intentionally not estimated in this draft because they have "
        "not yet been measured under one hardware, input-size, warm-up, and repetition protocol. These values must be added "
        "before submission if the target journal requires an efficiency comparison."
    )

    doc.add_heading("5.4 Evaluation metrics", level=2)
    add_body(
        doc,
        "Image-level evaluation reports AUROC, average precision (AP), and maximum F1. Pixel-level evaluation reports pixel "
        "AUROC, pixel AP, and area under the per-region-overlap curve (AUPRO). AUROC measures ranking over all thresholds, "
        "AP emphasizes the precision-recall trade-off under class imbalance, and AUPRO reduces domination by large defect "
        "regions by aggregating region overlap over a prescribed false-positive-rate interval. Main tables prioritize image "
        "AUROC, pixel AUROC, pixel AP, and AUPRO."
    )
    add_body(
        doc,
        "Metrics are first computed by category and then macro-averaged. When all three seeds are complete, results are "
        "reported as mean ± standard deviation in percentage points. Incomplete seed cells are not averaged. Development "
        "results are reported by seed rather than merged with independent validation."
    )

    # 6 Results
    doc.add_heading("6 Results", level=1)
    doc.add_heading("6.1 Unified VisA baselines", level=2)
    visa_table = []
    for row in visa_rows:
        method = row["method"] + ("*" if row["method"] == "PromptAD" else "")
        visa_table.append([
            method,
            row["shot"],
            f'{pct(row["image_auroc_mean"])} ± {pct(row["image_auroc_std"])}',
            f'{pct(row["pixel_auroc_mean"])} ± {pct(row["pixel_auroc_std"])}',
            f'{pct(row["pixel_ap_mean"])} ± {pct(row["pixel_ap_std"])}',
            f'{pct(row["aupro_mean"])} ± {pct(row["aupro_std"])}',
        ])
    add_table(
        doc, 3, "Unified VisA baseline results over three seeds (%, mean ± standard deviation).",
        ["Method", "Shot", "Image AUROC", "Pixel AUROC", "Pixel AP", "AUPRO"],
        visa_table, [1.55, 0.55, 1.35, 1.35, 1.15, 1.15],
        footnote=(
            "* PromptAD uses target-category normal images to learn prompts (target_normal_tuning=true). "
            "Other methods follow their frozen/reference-bank protocols."
        ),
        font_size=7.25,
    )
    add_body(
        doc,
        "AnomalyDINO is the strongest completed VisA baseline across the four principal metrics. Its image AUROC increases "
        "from 89.40% at one shot to 92.58% at four shots, while AUPRO increases from 92.21% to 93.69%. PatchCore also "
        "benefits substantially from more references, whereas WinCLIP+ changes more gradually. PromptAD obtains high pixel "
        "AUROC but lower image AUROC than AnomalyDINO under this reproduction. Because PromptAD learns target-domain "
        "prompts, its result should not be interpreted as a fully frozen inference comparison."
    )

    doc.add_heading("6.2 Unified MVTec AD baselines", level=2)
    mvtec_baseline_table = []
    for row in mvtec_rows:
        if row["method"] not in {"PatchCore", "WinCLIP+", "AnomalyDINO"}:
            continue
        mvtec_baseline_table.append([
            row["method"], row["shot"], row["image_auroc_mean_std"],
            row["pixel_auroc_mean_std"], row["pixel_ap_mean_std"], row["aupro_mean_std"],
        ])
    add_table(
        doc, 4, "Complete MVTec AD baselines over three seeds (%, mean ± standard deviation).",
        ["Method", "Shot", "Image AUROC", "Pixel AUROC", "Pixel AP", "AUPRO"],
        mvtec_baseline_table, [1.55, 0.55, 1.35, 1.35, 1.15, 1.15],
        font_size=7.4,
    )
    add_body(
        doc,
        "AnomalyDINO reaches image AUROC values of 95.71%, 96.86%, and 97.46% for 1, 2, and 4 shots, respectively, "
        "and also leads the completed pixel metrics. This establishes a strict requirement for fusion: semantic evidence "
        "must correct a limited number of visual failures without disturbing the already strong global visual ranking. "
        "The MVTec PromptAD matrix currently contains only four of nine seed-shot runs, so no three-seed mean or ranking is "
        "reported."
    )

    doc.add_heading("6.3 Independent VisA validation of the frozen router", level=2)
    dynamic_table = []
    for key in ["s1_k1", "s1_k2", "s1_k4", "s2_k1", "s2_k2", "s2_k4"]:
        row = visa_dynamic[key]
        dynamic_table.append([
            key[1], key[4], f'{float(row["image_auroc"]) * 100:.2f}',
            f'{float(row["pixel_auroc"]) * 100:.2f}', f'{float(row["pixel_ap"]) * 100:.2f}',
            f'{float(row["aupro"]) * 100:.2f}',
        ])
    add_table(
        doc, 5, "Frozen DynamicFusion V1 on independent VisA seeds 1 and 2 (%).",
        ["Seed", "Shot", "Image AUROC", "Pixel AUROC", "Pixel AP", "AUPRO"],
        dynamic_table, [0.7, 0.7, 1.55, 1.55, 1.35, 1.35],
        footnote="Seed 0 is excluded because it was used for development and parameter selection.",
        font_size=7.7,
    )
    add_body(
        doc,
        "Independent VisA image AUROC ranges from 79.77% to 82.50%, below raw AnomalyDINO. Pixel localization benefits "
        "more from additional references: AUPRO at two and four shots is 83.31% and 82.47% for seed 1 and 79.08% and "
        "84.81% for seed 2. The difference between image and pixel behavior motivates the split-temperature analysis, but "
        "the independent results do not support a claim that V1 is stronger overall."
    )

    doc.add_heading("6.4 Development ablation of split temperatures", level=2)
    selected_variants = {
        "calibrated_visual_only": "Calibrated visual only",
        "calibrated_text_only": "Calibrated text only",
        "fixed_visual_0.50": "Fixed visual weight 0.50",
        "fixed_visual_0.75": "Fixed visual weight 0.75",
        "single_temperature_0.20": "Single temperature 0.20",
        "split_image_0.50_pixel_0.20": "Split temperature 0.50/0.20",
    }
    ablation_table = []
    for row in ablation_rows:
        if row["variant"] not in selected_variants:
            continue
        ablation_table.append([
            selected_variants[row["variant"]], row["shot"],
            f'{float(row["image_auroc"]) * 100:.2f}', f'{float(row["pixel_auroc"]) * 100:.2f}',
            f'{float(row["pixel_ap"]) * 100:.2f}', f'{float(row["aupro"]) * 100:.2f}',
        ])
    add_table(
        doc, 6, "Selected VisA seed-0 development ablations (%).",
        ["Variant", "Shot", "Image AUROC", "Pixel AUROC", "Pixel AP", "AUPRO"],
        ablation_table, [2.2, 0.5, 1.25, 1.25, 1.1, 1.1],
        footnote=(
            "The complete ablation also includes fixed visual weights 0, 0.25, and 1.00 and a single temperature of 0.50. "
            "This table retains the most diagnostic controls."
        ),
        font_size=6.95,
        left_columns={0},
    )
    add_figure(
        doc, fig_dir / "visa_ablation_split_temperature.png", 2,
        "Image AUROC and AUPRO for selected fixed, single-temperature, and split-temperature variants on VisA seed 0.",
        "Two line charts comparing image AUROC and AUPRO over one, two, and four shots for fixed-weight and temperature-based fusion variants."
    )
    add_body(
        doc,
        "At two shots, the split-temperature setting improves image AUROC by 0.20 percentage points and AUPRO by 4.14 "
        "points over the best fixed-weight control. At four shots, the corresponding gains are 0.04 and 8.82 points. At "
        "one shot, image AUROC is 0.07 points lower while AUPRO is 2.78 points higher. The result supports different routing "
        "strengths for image ranking and pixel localization; it does not show that temperature separation repairs the "
        "upstream calibration failure."
    )

    doc.add_heading("6.5 Frozen MVTec AD validation", level=2)
    mvtec_dynamic = [
        ["1", "95.71 ± 0.02", "79.43 ± 2.61", "-16.29", "91.06 ± 0.40", "82.33 ± 3.12"],
        ["2", "96.86 ± 0.51", "86.37 ± 1.98", "-10.49", "93.79 ± 0.43", "90.58 ± 1.36"],
        ["4", "97.46 ± 0.47", "89.52 ± 2.69", "-7.94", "94.18 ± 0.03", "91.81 ± 0.80"],
    ]
    add_table(
        doc, 7, "Raw AnomalyDINO and frozen DynamicFusion V1 on MVTec AD (three seeds, %).",
        ["Shot", "Raw visual\nImage AUROC", "Dynamic\nImage AUROC", "Difference", "Dynamic\nPixel AUROC", "Dynamic\nAUPRO"],
        mvtec_dynamic, [0.55, 1.65, 1.55, 0.85, 1.55, 1.35], font_size=7.35
    )
    add_figure(
        doc, fig_dir / "mvtec_visual_vs_dynamic_by_shot.png", 3,
        "Image AUROC of raw AnomalyDINO and frozen DynamicFusion V1 on MVTec AD.",
        "Grouped bar chart showing that DynamicFusion has lower MVTec image AUROC than raw AnomalyDINO at one, two, and four shots.",
        width=5.7,
    )
    add_body(
        doc,
        "V1 remains below raw AnomalyDINO at every shot. The image AUROC gap decreases from 16.29 points at one shot to "
        "7.94 points at four shots, indicating that additional normal references stabilize calibration but do not remove "
        "the underlying defect. Pixel AUROC and AUPRO remain numerically high, yet those absolute values cannot justify an "
        "overall superiority claim when the image-level ranking is materially worse."
    )

    # 7 Analysis and discussion
    doc.add_heading("7 Failure Analysis and Discussion", level=1)
    doc.add_heading("7.1 Calibration saturation precedes routing failure", level=2)
    add_body(
        doc,
        "The dominant failure occurs before branch mixing. With only a few normal references, the MAD of the visual normal "
        "scores can be extremely small. A query score only slightly above the reference median then receives a very large "
        "standardized value, and the sigmoid maps a broad range of distinct scores to values numerically indistinguishable "
        "from one. A strictly monotonic transform would preserve ranking in exact arithmetic, but clipping, finite precision, "
        "and massive ties flatten the fine ordering that makes the visual branch strong."
    )
    add_figure(
        doc, fig_dir / "calibration_saturation_diagnostic.png", 4,
        "Visual-score saturation and the image-AUROC change caused by calibration.",
        "Two charts showing the percentage of calibrated visual scores at or above 0.999 and the decrease in image AUROC after visual calibration on VisA and MVTec AD."
    )
    add_body(
        doc,
        "On MVTec AD, an average of 99.99% of calibrated visual scores are at least 0.999; the VisA average is 91.54%. "
        "MVTec raw AnomalyDINO image AUROC values of 95.71%, 96.86%, and 97.46% become 51.53%, 54.46%, and 56.74% after "
        "visual calibration alone. Thus, the primary loss cannot be explained as the text branch merely pulling down a "
        "strong score. Much of the visual ranking has already disappeared when the score reaches the router."
    )
    failure_rows = [
        ["Reference calibration", "MAD scale becomes too small", "Most visual scores >=0.999", "Large groups of tied scores"],
        ["Uncertainty", "Entropy sees only extremeness", "Saturated 0.999 has low entropy", "False high confidence"],
        ["Routing weight", "Weak relation to true branch advantage", "Low category Spearman correlation", "Unreliable branch selection"],
        ["Image fusion", "Weight varies across samples", "Fused/raw visual rank correlation about 0.54", "Normal/anomalous ordering changes"],
        ["Pixel fusion", "Text response can be spatially diffuse", "Background activation in failure cases", "Reduced local contrast"],
    ]
    add_table(
        doc, 8, "Observed failure chain of DynamicFusion V1.",
        ["Stage", "Observed problem", "Diagnostic evidence", "Consequence"],
        failure_rows, [1.35, 2.0, 2.1, 1.85], font_size=7.2, left_columns={0, 1, 2, 3}
    )

    doc.add_heading("7.2 Why binary entropy cannot detect saturation error", level=2)
    add_body(
        doc,
        "Binary entropy is appropriate only after p has a defensible probabilistic interpretation. It becomes small both "
        "when a well-calibrated model is genuinely certain and when an ill-scaled transform pushes an ordinary score to a "
        "boundary. V1 cannot distinguish these cases and therefore treats calibration-induced 0.999999 values as reliable "
        "anomalous evidence. A safer uncertainty model must also measure distance beyond the support of the normal-reference "
        "distribution, saturation frequency, stability across deterministic views, branch disagreement, and spatial "
        "concentration."
    )

    doc.add_heading("7.3 Routing weights versus true branch advantage", level=2)
    add_figure(
        doc, fig_dir / "route_weight_summary.png", 5,
        "Mean image and pixel visual weights, including normal/anomalous image statistics.",
        "Line charts comparing image-level and pixel-level visual weights across one, two, and four shots on VisA and MVTec AD."
    )
    add_body(
        doc,
        "If routing is effective, categories in which AnomalyDINO substantially outperforms AnomalyCLIP should receive "
        "higher visual weights. The observed category-level Spearman correlations are weak: approximately 0.10, 0.16, and "
        "0.20 on MVTec AD and 0.36, 0.37, and 0.31 on VisA for 1, 2, and 4 shots. None of the 5,175 MVTec routing decisions "
        "is labeled text-dominant, and VisA has only a small number of text-dominant cases. Counts of discrete route labels "
        "therefore do not demonstrate that the router selected the better branch; the relevant test is whether continuous "
        "weights correlate with branch correctness and preserve the strong branch ranking."
    )

    doc.add_heading("7.4 Category-level heterogeneity", level=2)
    add_figure(
        doc, fig_dir / "mvtec_category_image_auroc_delta_heatmap.png", 6,
        "Category-level image-AUROC difference between DynamicFusion and raw AnomalyDINO on MVTec AD.",
        "Heatmap of DynamicFusion minus raw AnomalyDINO image AUROC for 15 MVTec categories at one, two, and four shots."
    )
    add_body(
        doc,
        "The largest decreases occur for carpet, cable, grid, zipper, and bottle. For example, carpet falls from 100% raw "
        "visual AUROC to approximately 61.31% after fusion, and cable falls from approximately 92.36% to 54.02%. A small "
        "number of categories retain complementary gains: capsule improves by about 1.42 points on average, wood by about "
        "0.54 points, and leather is approximately unchanged. These cases demonstrate that text evidence is not uniformly "
        "uninformative; rather, its gains are too infrequent and too small to offset the damage produced by unstable "
        "calibration."
    )

    doc.add_heading("7.5 Localization successes and failures", level=2)
    add_figure(
        doc, fig_dir / "mvtec_success_failure_cases.png", 7,
        "Automatically selected MVTec localization improvements and degradations at four shots.",
        "Rows of MVTec examples showing the input, ground truth, AnomalyDINO map, AnomalyCLIP map, fused map, and pixel visual weight for automatically selected improvement and degradation cases.",
        width=6.05,
    )
    add_body(
        doc,
        "In successful cases, the text-guided heatmap supplements thin structures missed by the visual branch, while the "
        "pixel weight selects evidence near structural boundaries and increases defect-to-background contrast. In failure "
        "cases, the visual response is already concentrated on the ground-truth region, whereas the text response or fusion "
        "weight spreads high activation across textures, object contours, or background. The displayed cases were selected "
        "automatically by the change in anomalous-region contrast relative to the visual branch, rather than by manual "
        "preference. Because they come from inspected final validation, they are explanatory evidence and cannot be reused "
        "for parameter selection."
    )

    doc.add_heading("7.6 Why image-level and pixel-level outcomes differ", level=2)
    add_body(
        doc,
        "Image AUROC depends only on the relative order of complete images. When each image receives a different weight, "
        "fusion can raise a normal image and suppress an anomalous one even when each branch is individually reasonable. "
        "The category-level mean Spearman correlation between fused and raw visual MVTec scores is approximately 0.54, "
        "showing a material ordering change. Pixel metrics instead emphasize spatial contrast within an image. Local "
        "weights can exploit semantic edges or region cues even when the image score is not improved, explaining why the "
        "split temperature yields larger AUPRO gains than image-AUROC gains in the development ablation."
    )

    doc.add_heading("7.7 Implications for a safer second-generation router", level=2)
    add_body(
        doc,
        "The next design should begin with a strong visual-default policy rather than another temperature sweep. Four "
        "requirements follow directly from the evidence. First, calibration must preserve ranking, for example through an "
        "strict monotonic mapping whose ties and pre/post Spearman correlation are audited. Instance-wise monotonic "
        "calibration [40] is conceptually relevant, but the mapping must be redesigned for normal-only anomaly scores. "
        "Second, the router must flag values beyond the supported normal-reference range instead of interpreting saturation "
        "as confidence; hardness-aware calibration such as CADET [39] motivates additional evidence but cannot use labels "
        "under this protocol. Third, a safe fallback should retain the visual prediction unless the text branch satisfies "
        "joint conditions on reference-view consistency, branch conflict, and spatial concentration. Fourth, image and "
        "pixel routing should use distinct features and acceptance tests because global ranking and local contrast are "
        "different optimization targets."
    )
    add_body(
        doc,
        "These changes must be developed on data that were not used to analyze V1. Minimum acceptance criteria should "
        "include high correlation between calibrated and raw strong-branch rankings, a bounded saturation rate, recovery "
        "to the visual branch when reliability checks fail, independent validation after parameter freezing, and the absence "
        "of material regression in categories where the visual branch is already near perfect. Because QMF [32] links "
        "useful quality scores to branch error, V2 must report weight-versus-advantage correlation as a primary diagnostic, "
        "not merely the number of samples assigned to each route. PGAD [34] further motivates explicit distribution-alignment "
        "controls. Figure 8 summarizes this prospective sequence; none of its safeguards is counted as a completed V1 result."
    )
    add_figure(
        doc, v2_roadmap, 8,
        "Prospective development path for DynamicFusion V2; the safeguards shown here were not implemented or evaluated in V1.",
        "Roadmap from the observed V1 failure chain through rank-preserving calibration, support diagnostics, visual-default fallback, separated image/pixel routing, and a new development/holdout boundary.",
        width=6.15,
    )
    add_body(
        doc,
        "Concurrent 2026 work also limits the scope of any future performance claim. SubspaceAD [11] can test whether a "
        "simple visual subspace already captures the needed few-shot structure; FastRef [12] tests whether safe prototype "
        "refinement is more effective than output fusion; AnoPLe [22] and MoECLIP [33] test learned prompt interaction and "
        "expert routing; and PAPL [24] or TGRF-CLIP [25] test upstream visual-language fusion. V2 should therefore be judged "
        "not only by its average metric, but also by whether it adds stable, interpretable benefit beyond these simpler or "
        "more directly trained alternatives."
    )

    doc.add_heading("7.8 Limitations and threats to validity", level=2)
    add_bullet(
        doc,
        "Repository heterogeneity: branch implementations use different native image sizes, preprocessing operations, and feature layers. The unified schema controls output consistency but cannot remove all algorithm-specific configuration differences."
    )
    add_bullet(
        doc,
        "Protocol heterogeneity: PromptAD learns prompts from target normal data and must remain separated from fully frozen inference methods through the target_normal_tuning=true label."
    )
    add_bullet(
        doc,
        "Matrix completeness: five MVTec PromptAD seed-shot runs remain incomplete. ReMP-AD, AdaptCLIP, SubspaceAD, and FastRef have not yet passed the project's unified Gate A verification. Their absence cannot be used to infer a ranking."
    )
    add_bullet(
        doc,
        "External validity: only MVTec AD and VisA are currently included. V1 has already been inspected on both, so V2 requires a new development boundary or an explicit exploratory label."
    )
    add_bullet(
        doc,
        "Efficiency validity: peak memory, warmed-up latency, and throughput have not yet been measured under one controlled protocol and are therefore omitted rather than estimated."
    )
    add_bullet(
        doc,
        "Statistical validity: the primary baselines use three seeds, whereas independent VisA router validation uses seeds 1 and 2 because seed 0 was reserved for development. The individual validation results are therefore reported without presenting a misleading three-seed mean."
    )

    # 8 Conclusion
    doc.add_heading("8 Conclusion", level=1)
    add_body(
        doc,
        "This study establishes a reproducible 1/2/4-shot industrial anomaly-detection workflow and evaluates a leakage-" 
        "safe visual-language router at image and pixel levels. The unified baseline results identify AnomalyDINO as a "
        "strong visual branch on both VisA and MVTec AD. Split temperatures provide meaningful local AUPRO gains in VisA "
        "development experiments, confirming that image ranking and pixel localization require different routing strengths."
    )
    add_body(
        doc,
        "Frozen validation also reveals a central limitation: robust-looking median/MAD calibration can become unstable "
        "when estimated from extremely few normal references. Severe sigmoid saturation destroys visual ranking, binary "
        "entropy misclassifies the saturation as confidence, and sample-dependent weights then propagate the error. The "
        "result is a substantial MVTec image-AUROC decrease despite isolated category and localization gains."
    )
    add_body(
        doc,
        "The defensible conclusion is therefore conditional. Visual and textual evidence can be locally complementary, but "
        "reliable fusion requires rank-preserving calibration, explicit out-of-support detection, and a safe fallback to the "
        "strong visual branch. These requirements, together with a new development/holdout boundary, define the next "
        "experimental stage and prevent repeated tuning on already inspected final results."
    )

    # Declarations expected by many current journal workflows.
    doc.add_heading("Declarations", level=1)
    add_body(doc, "Funding: To be completed by the authors before submission.", bold_lead="Funding:")
    add_body(doc, "Conflicts of interest: To be confirmed by all authors before submission.", bold_lead="Conflicts of interest:")
    add_body(doc, "Author contributions: To be completed using the target journal's contribution taxonomy, preferably CRediT.", bold_lead="Author contributions:")
    add_body(
        doc,
        "Data availability: MVTec AD and VisA are available from their original providers under the respective access and licensing conditions. The derived manifests, audit reports, and result tables are retained in the project repository and can be made available by the corresponding author subject to dataset licenses.",
        bold_lead="Data availability:",
    )
    add_body(
        doc,
        "Code availability: Reproduction adapters, unified evaluation scripts, routing code, and configuration files are retained in the project repository. The release location and archival identifier should be inserted before submission.",
        bold_lead="Code availability:",
    )
    add_body(doc, "Ethics approval: Not applicable; the study uses public industrial image datasets and does not involve human participants or animals.", bold_lead="Ethics approval:")
    add_body(
        doc,
        "Declaration of generative AI and AI-assisted technologies in manuscript preparation: During preparation of this draft, the authors used OpenAI Codex to assist with English drafting, organization, and document formatting. The authors must review and edit the complete manuscript, verify every technical claim and reference, and take full responsibility for the submitted content.",
        bold_lead="Declaration of generative AI and AI-assisted technologies in manuscript preparation:",
    )
    add_body(doc, "Acknowledgments: To be completed before submission.", bold_lead="Acknowledgments:")

    # References, ordered by first appearance.
    doc.add_heading("References", level=1)
    references = [
        "[1] P. Bergmann, M. Fauser, D. Sattlegger, and C. Steger, \"MVTec AD - A comprehensive real-world dataset for unsupervised anomaly detection,\" in Proc. IEEE/CVF Conf. Comput. Vis. Pattern Recognit., 2019, pp. 9592-9600.",
        "[2] Y. Zou, J. Jeong, L. Pemula, D. Zhang, and O. Dabeer, \"SPot-the-Difference self-supervised pre-training for anomaly detection and segmentation,\" in Proc. Eur. Conf. Comput. Vis., 2022, pp. 392-408.",
        "[3] J. Liu et al., \"Deep industrial image anomaly detection: A survey,\" Mach. Intell. Res., vol. 21, pp. 104-135, 2024, doi: 10.1007/s11633-023-1459-z.",
        "[4] X. Guo, L. Song, W. Zhu, F. Du, and Z. Ma, \"Review of low-shot industrial image anomaly detection,\" Comput. Eng. Appl., vol. 61, no. 13, pp. 26-45, 2025, doi: 10.3778/j.issn.1002-8331.2408-0230 (in Chinese).",
        "[5] K. Roth, L. Pemula, J. Zepeda, B. Schölkopf, T. Brox, and P. Gehler, \"Towards total recall in industrial anomaly detection,\" in Proc. IEEE/CVF Conf. Comput. Vis. Pattern Recognit., 2022, pp. 14318-14328.",
        "[6] M. Oquab et al., \"DINOv2: Learning robust visual features without supervision,\" Trans. Mach. Learn. Res., 2024.",
        "[7] S. Damm, M. Laszkiewicz, J. Lederer, and A. Fischer, \"AnomalyDINO: Boosting patch-based few-shot anomaly detection with DINOv2,\" in Proc. IEEE/CVF Winter Conf. Appl. Comput. Vis., 2025, pp. 1319-1329.",
        "[8] Z. Fang et al., \"FastRecon: Few-shot industrial anomaly detection via fast feature reconstruction,\" in Proc. IEEE/CVF Int. Conf. Comput. Vis., 2023.",
        "[9] J. Jeong, Y. Zou, T. Kim, D. Zhang, A. Ravichandran, and O. Dabeer, \"WinCLIP: Zero-/few-shot anomaly classification and segmentation,\" in Proc. IEEE/CVF Conf. Comput. Vis. Pattern Recognit., 2023, pp. 19606-19616.",
        "[10] Z. Gu, B. Zhu, G. Zhu, Y. Chen, M. Tang, and J. Wang, \"UniVAD: A training-free unified model for few-shot visual anomaly detection,\" in Proc. IEEE/CVF Conf. Comput. Vis. Pattern Recognit., 2025, pp. 15194-15203.",
        "[11] C. Lendering, E. Akdag, and E. Bondarau, \"SubspaceAD: Training-free few-shot anomaly detection via subspace modeling,\" in Proc. IEEE/CVF Conf. Comput. Vis. Pattern Recognit., 2026, pp. 28557-28566.",
        "[12] Y. Li, L. Tian, Y. Dai, W. Chen, L. Bao, and X. Liu, \"FastRef: Fast prototype refinement for few-shot industrial anomaly detection,\" in Proc. IEEE/CVF Conf. Comput. Vis. Pattern Recognit., 2026, pp. 43040-43049.",
        "[13] A. Radford et al., \"Learning transferable visual models from natural language supervision,\" in Proc. 38th Int. Conf. Mach. Learn., vol. 139, 2021, pp. 8748-8763.",
        "[14] Q. Zhou, G. Pang, Y. Tian, S. He, and J. Chen, \"AnomalyCLIP: Object-agnostic prompt learning for zero-shot anomaly detection,\" in Proc. Int. Conf. Learn. Representations, 2024.",
        "[15] X. Li, Z. Zhang, X. Tan, C. Chen, Y. Qu, Y. Xie, and L. Ma, \"PromptAD: Learning prompts with only normal samples for few-shot anomaly detection,\" in Proc. IEEE/CVF Conf. Comput. Vis. Pattern Recognit., 2024, pp. 16838-16848.",
        "[16] Y. Cao, J. Zhang, L. Frittoli, Y. Cheng, W. Shen, and G. Boracchi, \"AdaCLIP: Adapting CLIP with hybrid learnable prompts for zero-shot anomaly detection,\" in Proc. Eur. Conf. Comput. Vis., 2024.",
        "[17] Z. Qu, X. Tao, M. Prasad, F. Shen, Z. Zhang, X. Gong, and G. Ding, \"VCP-CLIP: A visual context prompting model for zero-shot anomaly segmentation,\" in Proc. Eur. Conf. Comput. Vis., 2024.",
        "[18] Z. Gu et al., \"FiLo: Zero-shot anomaly detection by fine-grained description and high-quality localization,\" in Proc. ACM Multimedia, 2024, doi: 10.1145/3664647.3680685.",
        "[19] W. Ma et al., \"AA-CLIP: Enhancing zero-shot anomaly detection via anomaly-aware CLIP,\" in Proc. IEEE/CVF Conf. Comput. Vis. Pattern Recognit., 2025, pp. 4744-4754.",
        "[20] H. Ma, G. Yang, D. Zhao, Y. Ji, and W. Zuo, \"ReMP-AD: Retrieval-enhanced multi-modal prompt fusion for few-shot industrial visual anomaly detection,\" in Proc. IEEE/CVF Int. Conf. Comput. Vis., 2025, pp. 20425-20434.",
        "[21] B.-B. Gao et al., \"AdaptCLIP: Adapting CLIP for universal visual anomaly detection,\" Proc. AAAI Conf. Artif. Intell., vol. 40, no. 6, pp. 4095-4103, 2026, doi: 10.1609/aaai.v40i6.42404.",
        "[22] Y. Lee, S. Kim, D. Moon, S. Jang, and H. Yoon, \"Bidirectional multimodal prompt learning with scale-aware training for few-shot multi-class anomaly detection,\" in Proc. IEEE/CVF Conf. Comput. Vis. Pattern Recognit., 2026, pp. 35577-35586.",
        "[23] G. Zhang and L. Zhang, \"DLVP-CLIP: Enhancing fine-grained zero-shot anomaly detection via dynamic local visual prompting,\" in Proc. IEEE/CVF Conf. Comput. Vis. Pattern Recognit., 2026, pp. 35524-35533.",
        "[24] R. Ma et al., \"PAPL: Particle-based adaptive prompt learning for zero-shot industrial anomaly detection,\" Pattern Recognit., vol. 178, art. 113489, 2026, doi: 10.1016/j.patcog.2026.113489.",
        "[25] H.-L. Yan and X.-S. Xu, \"TGRF-CLIP: CLIP-based text-guided fusion of visual residuals for few-shot anomaly detection,\" Expert Syst. Appl., art. 132817, 2026, doi: 10.1016/j.eswa.2026.132817.",
        "[26] T. Defard, A. Setkov, A. Loesch, and R. Audigier, \"PaDiM: A patch distribution modeling framework for anomaly detection and localization,\" in Proc. 25th Int. Conf. Pattern Recognit., 2021, pp. 475-489.",
        "[27] H. Deng and X. Li, \"Anomaly detection via reverse distillation from one-class embedding,\" in Proc. IEEE/CVF Conf. Comput. Vis. Pattern Recognit., 2022, pp. 9737-9746.",
        "[28] Z. Liu, Y. Zhou, Y. Xu, and Z. Wang, \"SimpleNet: A simple network for image anomaly detection and localization,\" in Proc. IEEE/CVF Conf. Comput. Vis. Pattern Recognit., 2023, pp. 20402-20411.",
        "[29] K. Batzner, L. Heckler, and R. König, \"EfficientAD: Accurate visual anomaly detection at millisecond-level latencies,\" in Proc. IEEE/CVF Winter Conf. Appl. Comput. Vis., 2024, pp. 128-138.",
        "[30] Z. Gu, B. Zhu, G. Zhu, Y. Chen, M. Tang, and J. Wang, \"AnomalyGPT: Detecting industrial anomalies using large vision-language models,\" Proc. AAAI Conf. Artif. Intell., vol. 38, no. 3, pp. 1932-1940, 2024, doi: 10.1609/aaai.v38i3.27963.",
        "[31] Z. Han et al., \"Multimodal dynamics: Dynamical fusion for trustworthy multimodal classification,\" in Proc. IEEE/CVF Conf. Comput. Vis. Pattern Recognit., 2022, pp. 20707-20717.",
        "[32] Q. Zhang, H. Wu, C. Zhang, Q. Hu, H. Fu, J. T. Zhou, and X. Peng, \"Provable dynamic fusion for low-quality multimodal data,\" in Proc. 40th Int. Conf. Mach. Learn., vol. 202, 2023, pp. 41753-41769.",
        "[33] J. Y. Park, J. Seo, M. Kang, and Y. R. Park, \"MoECLIP: Patch-specialized experts for zero-shot anomaly detection,\" in Proc. IEEE/CVF Conf. Comput. Vis. Pattern Recognit., 2026, pp. 35534-35544.",
        "[34] J. Zhou, W. K. Wong, and F. Liao, \"One-shot unsupervised industrial anomaly detection: Enhanced performance under extreme data scarcity,\" Pattern Recognit., vol. 173, art. 112759, 2026, doi: 10.1016/j.patcog.2025.112759.",
        "[35] A. Goodge, B. Hooi, W. S. Ng, and S. K. Ng, \"When text and images don't mix: Bias-correcting language-image similarity scores for anomaly detection,\" in Proc. Brit. Mach. Vis. Conf., 2024.",
        "[36] C. Guo, G. Pleiss, Y. Sun, and K. Q. Weinberger, \"On calibration of modern neural networks,\" in Proc. 34th Int. Conf. Mach. Learn., vol. 70, 2017, pp. 1321-1330.",
        "[37] Y. Ovadia et al., \"Can you trust your model's uncertainty? Evaluating predictive uncertainty under dataset shift,\" in Adv. Neural Inf. Process. Syst., vol. 32, 2019.",
        "[38] M. Kull, M. Perello-Nieto, M. Kängsepp, T. Silva Filho, H. Song, and P. Flach, \"Beyond temperature scaling: Obtaining well-calibrated multiclass probabilities with Dirichlet calibration,\" in Adv. Neural Inf. Process. Syst., vol. 32, 2019.",
        "[39] A. Deng, A. Goodge, L. Y. Ang, and B. Hooi, \"CADET: Calibrated anomaly detection for mitigating hardness bias,\" in Proc. 31st Int. Joint Conf. Artif. Intell., 2022, pp. 2002-2008, doi: 10.24963/ijcai.2022/278.",
        "[40] Y. Zhang, G. E. Batista, and S. S. Kanhere, \"Instance-wise monotonic calibration by constrained transformation,\" in Proc. 41st Conf. Uncertainty Artif. Intell., vol. 286, 2025, pp. 4920-4932.",
        "[41] X. Li, Z. Huang, F. Xue, and Y. Zhou, \"MuSc: Zero-shot industrial anomaly classification and segmentation with mutual scoring of the unlabeled images,\" in Proc. Int. Conf. Learn. Representations, 2024.",
        "[42] S. Li et al., \"DNPR: Zero-shot industrial anomaly detection via dynamic normal prototype refinement,\" Expert Syst. Appl., vol. 312, art. 131331, 2026, doi: 10.1016/j.eswa.2026.131331.",
    ]
    for reference in references:
        paragraph = doc.add_paragraph(style="Reference")
        run = paragraph.add_run(reference)
        set_run_font(run, size=8.8)

    # Appendix: reproducibility evidence that may be moved to supplementary material.
    doc.add_heading("Appendix A. Reproducibility and Evidence Checklist", level=1)
    checklist_rows = [
        ["Dataset provenance", "Official MVTec license workflow; official VisA source", "Verified"],
        ["Archive and manifest hashes", "SHA256 recorded", "Complete"],
        ["Nested few-shot sampling", "K1 subset K2 subset K4; three seeds", "Verified"],
        ["Unified prediction schema", "IDs, scores, maps, labels, masks", "Implemented"],
        ["Result completeness", "Category count, sample count, zero schema error", "Audited; MVTec PromptAD partial"],
        ["Fusion information boundary", "No test labels, masks, or aggregate statistics", "17/17 runs passed"],
        ["Development/validation split", "VisA seed 0 development; seeds 1/2 validation", "Enforced"],
        ["Frozen V1", "No post-validation retuning", "Enforced"],
        ["Efficiency measurement", "Common hardware, size, warm-up, repetitions", "Pending"],
        ["Recent external methods", "ReMP-AD, AdaptCLIP, SubspaceAD, FastRef Gate A", "Pending"],
    ]
    add_table(
        doc, 9, "Reproducibility evidence and remaining pre-submission items.",
        ["Check", "Requirement", "Draft status"], checklist_rows,
        [1.75, 3.35, 1.4], font_size=7.3, left_columns={0, 1, 2}
    )
    # Word needs a body paragraph after a terminal table. Keep it at 1 pt so it
    # remains on the appendix page instead of creating a blank trailing page.
    terminal = doc.add_paragraph()
    terminal.paragraph_format.space_before = Pt(0)
    terminal.paragraph_format.space_after = Pt(0)
    terminal.paragraph_format.line_spacing = Pt(1)
    terminal.paragraph_format.keep_with_next = False
    terminal_run = terminal.add_run(" ")
    set_run_font(terminal_run, size=1)
    doc.save(DOCX_PATH)
    return DOCX_PATH


if __name__ == "__main__":
    print(build_doc())

"""Create a concise Chinese project overview Word document."""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "few_shot_industrial_ad_project_overview_expanded.docx"

BLUE = "2E74B5"
NAVY = "1F4D78"
LIGHT = "E8EEF5"
GRAY = "F2F4F7"


def set_font(run, size=11, bold=False, color="000000"):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shade = OxmlElement("w:shd")
    shade.set(qn("w:fill"), fill)
    tc_pr.append(shade)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def fix_table(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    ind = OxmlElement("w:tblInd")
    ind.set(qn("w:w"), "120")
    ind.set(qn("w:type"), "dxa")
    tbl_pr.append(ind)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Inches(widths[i])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            tc_w = cell._tc.tcPr.tcW
            tc_w.set(qn("w:w"), str(round(widths[i] * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def add_para(doc, text="", style=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    if bold_prefix and text.startswith(bold_prefix):
        set_font(p.add_run(bold_prefix), bold=True)
        set_font(p.add_run(text[len(bold_prefix):]))
    else:
        set_font(p.add_run(text))
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.10
    set_font(p.add_run(text))
    return p


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = Inches(1)
    section.left_margin = section.right_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    for name, size, color, before, after in (("Heading 1", 16, BLUE, 16, 8), ("Heading 2", 13, BLUE, 12, 6)):
        style = doc.styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(header.add_run("少样本工业异常检测项目 | 精简总览"), size=9, color="666666")
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(footer.add_run("内部项目简报"), size=9, color="666666")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(18)
    title.paragraph_format.space_after = Pt(5)
    set_font(title.add_run("少样本工业异常检测项目总览"), size=22, bold=True, color=NAVY)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(18)
    set_font(sub.add_run("V1–V3 方案演进、当前结论与论文输出定位"), size=11, color="666666")

    lead = doc.add_table(rows=1, cols=1)
    fix_table(lead, [6.5])
    set_cell_shading(lead.cell(0, 0), LIGHT)
    p = lead.cell(0, 0).paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    set_font(p.add_run("核心结论："), bold=True, color=NAVY)
    set_font(p.add_run(" 项目已建立可复核的少样本基线、统一预测缓存与泄漏审计流程。动态融合的 V1、V2、V3 均完成了相应验证，但尚未得到可稳定超越强视觉分支、可正式冻结的动态融合算法。当前应如实保留负结果，并把 GPU 资源优先用于既定基线矩阵和论文证据补齐。"))

    doc.add_heading("1. 从项目启动到动态融合研究的整体过程", level=1)
    add_para(doc, "项目最初目标是建立一套少样本工业异常检测的可复现实验体系，并在公平的基线比较之上探索视觉—文本动态融合。工作不是直接从融合算法开始，而是先完成数据、训练协议、预测缓存和统一评测的基础建设。")
    add_bullet(doc, "数据与协议：整理 VisA、MVTec AD，并为第二阶段准备 MPDD 开发集和 BTAD 保持集；建立 1/2/4-shot、3 个随机种子、嵌套抽样和类别/样本数量核验机制。")
    add_bullet(doc, "基线复现与训练：接入 PatchCore、WinCLIP+、AnomalyDINO 和 PromptAD。每个方法的输出被转为统一 NPZ 缓存，随后由同一评价器计算图像级和像素级指标，避免直接比较不同仓库各自的评测口径。")
    add_bullet(doc, "结果与审计：VisA 上四个主要方法的 3 seeds × 1/2/4-shot 矩阵已完成；MVTec 上 PatchCore、WinCLIP+、AnomalyDINO 已完整，PromptAD 仍有部分组合待恢复。动态融合输出另有校准文件、哈希、日志和独立审计。")
    add_bullet(doc, "第二阶段动机：强视觉分支擅长定位，但文本分支可能在类别语义或部分缺陷区域提供补充。因此，融合研究的核心问题不是简单平均两个分数，而是能否在不看测试标签的条件下判断“什么时候应相信文本”。")

    doc.add_heading("2. 基线训练与当前实验格局", level=1)
    baseline = doc.add_table(rows=1, cols=3)
    baseline.style = "Table Grid"
    fix_table(baseline, [1.35, 3.10, 2.05])
    for i, text in enumerate(["部分", "已完成内容", "当前状态"]):
        set_cell_shading(baseline.cell(0, i), GRAY)
        p = baseline.cell(0, i).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(p.add_run(text), bold=True, color=NAVY)
    for row in [
        ("VisA 基线", "PatchCore、WinCLIP+、AnomalyDINO、PromptAD 均完成 3 seeds × 1/2/4-shot 的统一预测与评测。", "可作为主要基线证据；PromptAD 需注明 target-normal prompt tuning。"),
        ("MVTec 基线", "PatchCore、WinCLIP+、AnomalyDINO 的完整矩阵与统一审计已具备。", "PromptAD 尚缺 5 个组合；正式跨方法主表不能把不完整矩阵当作最终排名。"),
        ("动态融合", "完成冻结缓存、校准、路由、消融、独立验证与 V3 AdaptCLIP Gate。", "工程链路完整，但尚未得到可冻结为成功方法的性能结论。"),
    ]:
        cells = baseline.add_row().cells
        for i, text in enumerate(row):
            set_font(cells[i].paragraphs[0].add_run(text), bold=(i == 0), color=(NAVY if i == 0 else "000000"))

    doc.add_heading("3. V1–V3 的方案演进", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    fix_table(table, [0.75, 1.55, 2.85, 1.35])
    headers = ["版本", "核心改动", "验证发现", "当前定位"]
    for i, text in enumerate(headers):
        set_cell_shading(table.cell(0, i), GRAY)
        p = table.cell(0, i).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(p.add_run(text), bold=True, color=NAVY)
    rows = [
        ("V1", "以 AnomalyDINO 为视觉分支、AnomalyCLIP 为文本分支，分别融合图像分数和像素异常图；比较固定权重、双温度与熵驱动的动态权重。", "在完整接口、校准和消融后发现：校准把大量不同视觉分数压到边界，排序信息丢失；熵随之误把“数值贴边”当成高置信度。", "第一版设计和失败链已冻结；不再通过继续搜索温度来追求表面提升。"),
        ("V2", "保留双分支，但改用排序保持校准、正常参考支持范围、异常输入保护、视觉默认回退，以及受限的文本像素辅助。", "V1 的数值饱和已修复，泄漏审计、缓存消融和 3 seeds × 1/2/4-shot 开发矩阵均完成；但文本路由很少被可靠激活，实际平均收益接近零。", "证明“安全融合框架”可运行，但没有证明“动态文本救援”具有稳定性能收益。"),
        ("V3", "以更强的 AdaptCLIP 代替旧文本分支；先做缓存验收，再用 Gate A1 检查文本可提供的理论上限，用 Gate A2 检查无标签路由是否可预测这些收益。", "A1 通过：文本在多数类别和约一半异常区域具有额外信息；A2 失败：仅用 K=1 正常参考校准时，路由无法可靠选中这些区域，最保守配置仍整体退化。", "作为可审计负结果保留；不扩展至更多 seed、shot、数据集或 GPU 矩阵。"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            p = cells[i].paragraphs[0]
            set_font(p.add_run(text), bold=(i == 0), color=(NAVY if i == 0 else "000000"))

    doc.add_heading("4. 为什么 V1、V2、V3 都未达到预期", level=1)
    doc.add_heading("V1：校准损害了动态权重的依据", level=2)
    add_para(doc, "V1 的问题发生在分数进入路由器之前。视觉分支的强弱排序在校准后被大量压缩到接近同一边界，动态权重所依赖的熵、分歧和置信度不再能反映真实的分支优势。因此即使固定温度或权重偶尔有效，也不能证明路由器真正学会了选择。")
    doc.add_heading("V2：安全性提高，但文本帮助空间没有被稳定利用", level=2)
    add_para(doc, "V2 的工作重点是修复 V1：让校准保持排序，用正常参考图定义支持范围，把不可信的文本输出限制在较小残差内，并默认回退到视觉分支。结果显示这些约束避免了明显错误，但也意味着文本只有在严格条件下才允许介入。在 3 seeds × 1/2/4-shot 的开发验证中，文本对像素定位的平均增益几乎为零，说明当前路由特征没有稳定抓住真正有帮助的文本信号。")
    doc.add_heading("V3：更强文本分支有上限，但无标签路由仍不能识别它", level=2)
    add_para(doc, "V3 排除了“旧文本模型太弱”这一种解释。AdaptCLIP 缓存已验收：6 类、458 张图、样本 ID 与视觉分支完全一致，数值有限且无标签泄漏。Gate A1 的 Oracle 结果表明文本可带来潜在帮助：图像级上限在 5/6 类别为正、像素级上限在 6/6 类别为正，204/407 个异常区域中文本更强。")
    add_para(doc, "但 Gate A2 只允许使用 K=1 正常参考图拟合文本校准，并让路由器仅观察两个分支的无标签证据。结果不通过：最保守方案的平均像素 AP 下降 0.07369、像素 AUROC 下降 0.00769，6 个类别没有正向 AP 增益。这表明瓶颈不是“文本有无信息”，而是“如何无标签地判断该信息在当前图像或区域是否可信”。")
    add_para(doc, "因此，V2 和 V3 都应被表述为严谨的负结果：前者验证了安全机制与审计框架，后者进一步证明仅更换为更强文本模型仍不足以解决无标签路由问题。它们不能被写成已经优于强视觉分支的成功算法。")

    doc.add_heading("5. 目前项目状态与下一步", level=1)
    add_bullet(doc, "已完成：从数据协议、基线训练、统一预测缓存到方法无关评估的主链路；动态融合 V1/V2/V3 的实现、审计和关键负结果。")
    add_bullet(doc, "不再做：对 V2/V3 进行无边界参数搜索；在未出现新的、预先定义的无标签可靠性信号前，启动 V3 的大规模 GPU 矩阵。")
    add_bullet(doc, "优先做：补齐并审计基线方法矩阵，保持 MVTec 主表只使用完整配置；整理方法比较、失败案例和动态融合消融证据。")
    add_bullet(doc, "若未来重启动态融合：必须提出新的无标签可靠性特征，先在开发集重做 Gate A1/A2，通过后才允许进入多 seed、多 shot 和保持集验证。")

    doc.add_heading("6. 论文初稿输出（简要）", level=1)
    add_para(doc, "论文初稿可以先完成问题背景、少样本实验协议、基线训练与统一评测框架、已完成的基线结果和动态融合设计。动态融合部分应将 V1–V3 写成严格的消融与失败分析：说明哪些改动修复了数值或工程问题，哪些门控结果未通过，以及为何不扩大未通过方案。")
    add_para(doc, "在基线矩阵和公平比较完整前，不给出跨方法最终排名；在没有稳定提升前，不把 V2/V3 写为论文的核心性能贡献。这样的初稿仍可作为可审计的研究记录，并为后续新路由器设计保留清晰的对照基线。")

    doc.add_heading("7. 结语", level=1)
    add_para(doc, "当前项目最有价值的产出是：以统一协议把“分支有潜力”和“无标签动态融合能稳定发挥该潜力”区分开来。V1–V3 已经说明，后者不能仅靠温度、校准或更强文本模型自动获得。后续工作应先保证基线结果完整可靠，再决定是否投入下一轮动态路由研究。")

    doc.core_properties.title = "少样本工业异常检测项目总览"
    doc.core_properties.subject = "V1-V3 方案演进与当前结论"
    doc.core_properties.author = "SCI Project"
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()

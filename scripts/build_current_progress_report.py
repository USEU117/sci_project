from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "few_shot_industrial_ad_progress_report_2026-07-30.docx"
BLUE = "2E74B5"
DARK = "1F4D78"
INK = "0B2545"
HEADER = "E8EEF5"
CALLOUT = "F4F6F9"
CAUTION = "FFF8E8"


def font(run, size=11, color="000000", bold=False, italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def shade(cell, fill):
    props = cell._tc.get_or_add_tcPr()
    elem = props.find(qn("w:shd"))
    if elem is None:
        elem = OxmlElement("w:shd")
        props.append(elem)
    elem.set(qn("w:fill"), fill)


def cell_width(cell, width):
    props = cell._tc.get_or_add_tcPr()
    elem = props.find(qn("w:tcW"))
    if elem is None:
        elem = OxmlElement("w:tcW")
        props.append(elem)
    elem.set(qn("w:w"), str(width))
    elem.set(qn("w:type"), "dxa")


def cell_margin(cell, top=80, bottom=80, start=120, end=120):
    props = cell._tc.get_or_add_tcPr()
    margins = props.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        props.append(margins)
    for side, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        elem = margins.find(qn(f"w:{side}"))
        if elem is None:
            elem = OxmlElement(f"w:{side}")
            margins.append(elem)
        elem.set(qn("w:w"), str(value))
        elem.set(qn("w:type"), "dxa")


def fix_table(table, widths, header=True):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    props = table._tbl.tblPr
    width = props.find(qn("w:tblW"))
    if width is None:
        width = OxmlElement("w:tblW")
        props.append(width)
    width.set(qn("w:w"), str(sum(widths)))
    width.set(qn("w:type"), "dxa")
    indent = props.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        props.append(indent)
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for value in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(value))
        grid.append(col)
    for row_index, row in enumerate(table.rows):
        for index, cell in enumerate(row.cells):
            cell_width(cell, widths[index])
            cell_margin(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(3)
                paragraph.paragraph_format.line_spacing = 1.05
                for run in paragraph.runs:
                    font(run, size=9.2, bold=(header and row_index == 0))
            if header and row_index == 0:
                shade(cell, HEADER)


def table(doc, headers, rows, widths):
    item = doc.add_table(rows=1, cols=len(headers))
    for index, value in enumerate(headers):
        item.rows[0].cells[index].text = str(value)
    for values in rows:
        cells = item.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = str(value)
    fix_table(item, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def para(doc, text="", size=11, color="000000", bold=False, after=6, align=None):
    item = doc.add_paragraph()
    item.paragraph_format.space_after = Pt(after)
    item.paragraph_format.line_spacing = 1.10
    if align is not None:
        item.alignment = align
    run = item.add_run(text)
    font(run, size=size, color=color, bold=bold)
    return item


def bullet(doc, text):
    item = doc.add_paragraph(style="List Bullet")
    item.paragraph_format.left_indent = Inches(0.5)
    item.paragraph_format.first_line_indent = Inches(-0.25)
    item.paragraph_format.space_after = Pt(4)
    item.paragraph_format.line_spacing = 1.167
    font(item.add_run(text))


def callout(doc, label, text, fill=CALLOUT):
    item = doc.add_table(rows=1, cols=1)
    cell = item.cell(0, 0)
    shade(cell, fill)
    cell_margin(cell, top=120, bottom=120, start=160, end=160)
    fix_table(item, [9360], header=False)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.10
    font(paragraph.add_run(label + "："), color=DARK, bold=True)
    font(paragraph.add_run(text))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def heading(doc, text, level=1):
    item = doc.add_heading(text, level=level)
    for run in item.runs:
        font(run, size={1: 16, 2: 13, 3: 12}[level], color=BLUE if level < 3 else DARK, bold=True)
    item.paragraph_format.space_before = Pt({1: 16, 2: 12, 3: 8}[level])
    item.paragraph_format.space_after = Pt({1: 8, 2: 6, 3: 4}[level])
    item.paragraph_format.keep_with_next = True


def setup(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = section.bottom_margin = Inches(1)
    section.left_margin = section.right_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in (("Heading 1", 16, BLUE, 16, 8), ("Heading 2", 13, BLUE, 12, 6), ("Heading 3", 12, DARK, 8, 4)):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
    header = section.header.paragraphs[0]
    font(header.add_run("少样本工业异常检测项目 · 进度报告"), size=9, color="666666")
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    font(footer.add_run("项目进度报告 | 2026-07-30"), size=9, color="666666")


def build():
    doc = Document()
    setup(doc)
    title = para(doc, "少样本工业异常检测项目进度报告", size=25, color=INK, bold=True, after=6, align=WD_ALIGN_PARAGRAPH.CENTER)
    title.paragraph_format.space_before = Pt(24)
    para(doc, "基准复现进展、当前训练状态与下一阶段计划", size=14, color="555555", after=18, align=WD_ALIGN_PARAGRAPH.CENTER)
    callout(doc, "报告结论", "第一阶段已完成数据、统一划分和统一评估基础设施；PatchCore、WinCLIP+、AnomalyDINO 已完成 VisA 完整矩阵。PromptAD 已完成 seed 0 的 1/2/4-shot，并正在继续 seed 1 的 1-shot。")
    table(doc, ["项目项", "当前状态"], [
        ("数据集", "VisA（12 类）与 MVTec AD（15 类）均已验证"),
        ("统一协议", "1/2/4-shot，seed 0/1/2，嵌套固定 manifest"),
        ("统一指标", "图像 AUROC/AP/F1-max；像素 AUROC/AP/AUPRO"),
        ("当前 GPU 任务", "PromptAD · VisA · seed 1 · 1-shot · candle 分类训练"),
        ("阶段二状态", "仅开展动态融合设计准备，尚未进行正式融合实验"),
    ], [2300, 7060])

    heading(doc, "一、已完成的项目基础", 1)
    para(doc, "项目已形成可追溯的实验框架：官方数据来源、环境、第三方代码提交、固定少样本划分、统一预测 NPZ 格式和方法无关的评估器均已建立。每次运行保留日志、完成标记、预测缓存和逐类结果。")
    for text in [
        "VisA 与 MVTec AD 已完成目录、样本和异常掩码校验；MVTec 的 15 类数据及 315 条 manifest 引用均通过检查。",
        "统一评估层已通过 5 项单元测试，能够严格检查标签、像素图形状和分数方向。",
        "所有方法使用相同的 1/2/4-shot 嵌套样本清单；因此结果差异可主要归因于方法，而不是随机抽样。",
        "长任务采用完成标记与预测缓存，训练中断后可从未完成类别恢复，不覆盖已有成功结果。",
    ]:
        bullet(doc, text)

    heading(doc, "二、已完成的核心基准结果", 1)
    table(doc, ["方法", "VisA 完成范围", "代表性结论", "MVTec 当前状态"], [
        ("AnomalyCLIP", "全类官方协议推理", "VisA 图像 AUROC 81.97%，AUPRO 83.60%", "15 类官方零样本推理完成"),
        ("PatchCore", "1/2/4-shot × 3 seeds", "4-shot：图像 AUROC 78.68%，AUPRO 62.21%", "bottle 1-shot Gate A 完成"),
        ("WinCLIP+", "1/2/4-shot × 3 seeds", "4-shot：图像 AUROC 72.58%，AUPRO 69.05%", "15 类 1-shot Gate B 完成"),
        ("AnomalyDINO", "1/2/4-shot × 3 seeds", "4-shot：图像 AUROC 92.58%，AUPRO 93.69%", "bottle 1-shot Gate A 完成"),
    ], [1700, 2700, 2850, 2110])
    para(doc, "说明：PatchCore、WinCLIP+ 和 AnomalyDINO 的 VisA 表格均为 3 个 seed 的 macro mean；PromptAD 单列报告，因为它会利用目标域正常样本学习 prompt。", size=9.5, color="555555", after=8)

    heading(doc, "三、PromptAD 当前成果", 1)
    para(doc, "PromptAD 的分类和分割分支均从固定的目标域正常样本学习 prompt，因此在所有结果中明确标注 target_normal_tuning=true。每个配置均先分别训练分类与分割分支，再导出、合并为统一预测并重新评估。")
    table(doc, ["VisA 配置", "样本数", "图像 AUROC", "像素 AUROC", "像素 AP", "AUPRO", "状态"], [
        ("1-shot, seed 0", "2,162", "80.25%", "96.20%", "28.54%", "81.73%", "完成"),
        ("2-shot, seed 0", "2,162", "81.15%", "96.79%", "29.63%", "82.25%", "完成"),
        ("4-shot, seed 0", "2,162", "80.46%", "97.01%", "31.87%", "83.71%", "完成"),
        ("1-shot, seed 1", "2,162", "—", "—", "—", "—", "运行中"),
    ], [1650, 900, 1250, 1250, 1050, 1050, 1210])
    callout(doc, "完整性检查", "seed 0 的三个 shot 均生成 24 个训练完成标记、12 个合并预测 NPZ 文件和统一评估报告；每个报告的 validation_errors 均为 0。", "E8EEF5")

    heading(doc, "四、当前工作与后续顺序", 1)
    table(doc, ["优先级", "工作项", "完成条件"], [
        ("P0", "完成 PromptAD VisA 的 seed 1/2、1/2/4-shot", "9 个配置齐全，报告 mean ± std，所有配置零 schema 错误"),
        ("P1", "PromptAD MVTec Gate A", "bottle、1-shot、seed 0 完成训练、预测导出和统一评估"),
        ("P2", "扩展 MVTec 基线矩阵", "PatchCore/AnomalyDINO 通过 Gate B 后扩展至完整 shot × seed"),
        ("P3", "ReMP-AD 与 AdaptCLIP Gate A", "权重、依赖、数据入口和单类可复现实验均通过"),
        ("P4", "汇总对比与论文材料", "逐类表、mean ± std 主表、效率表、失败案例与差异报告齐全"),
    ], [800, 4200, 4360])
    para(doc, "实验执行遵循 Gate A（单类 1-shot seed 0）→ Gate B（全类 1-shot seed 0）→ Gate C（完整 1/2/4-shot × 3 seeds）。未通过上一关的运行不进入下一关。")

    heading(doc, "五、第二阶段：动态融合的可行推进方式", 1)
    callout(doc, "建议", "可以现在开展算法设计、接口定义、配置和单元测试，但不应在基线尚未完整时报告最终融合性能或反复调参。", CAUTION)
    para(doc, "建议把后续工作分为两条轨道：基线轨继续完成耗时实验；设计轨利用已冻结的预测缓存开展不占 GPU 的准备工作。")
    for text in [
        "定义融合输入：视觉分支分数/像素图、文本分支分数/像素图、不确定性和跨分支一致性特征。",
        "定义路由输出：视觉主导、文本主导、固定加权或像素级动态权重；明确不能使用测试真值或测试集整体统计量。",
        "先在合成数据和已冻结的 VisA seed 0 预测上测试接口、数值稳定性和消融代码；不以此作为最终成绩。",
        "锁定结构后，使用 VisA seed 1/2 与 MVTec 作为独立验证，报告完整 mean ± std。",
    ]:
        bullet(doc, text)

    heading(doc, "六、当前判断", 1)
    para(doc, "项目的主要风险已从“数据、环境和评估能否稳定运行”转移为“完整矩阵的计算时间与新方法的可复现性”。当前基准框架已经足以支撑动态融合的工程设计；正式训练与结论应等待预注册的验证协议和剩余基线结果完成后再进行。")
    para(doc, "数据来源：PLAN.md、PROJECT_STATUS.md、experiments/registry.csv、PromptAD 统一评估 summary.csv、日志与预测缓存。报告中的运行中状态以 2026-07-30 17:30 的终端检查为准。", size=9.5, color="555555", after=0)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()

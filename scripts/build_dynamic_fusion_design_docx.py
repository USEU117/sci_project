from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "outputs" / "第二阶段动态融合完整设计说明.docx"
ASSET_DIR = ROOT / "outputs" / "_docx_assets_dynamic_fusion"
DIAGRAM = ASSET_DIR / "dynamic_fusion_overview.png"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "203748"
TEXT = "262626"
MUTED = "666666"
LIGHT_BLUE = "E8EEF5"
PALE_BLUE = "F4F7FA"
PALE_GREEN = "EAF4EC"
PALE_YELLOW = "FFF5D6"
PALE_RED = "FCE8E6"
WHITE = "FFFFFF"
GRID = "AAB7C4"

FONT_CN = "Microsoft YaHei"
FONT_FALLBACK = "SimHei"
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    if sum(widths_dxa) != CONTENT_WIDTH_DXA:
        raise ValueError(f"table widths must sum to {CONTENT_WIDTH_DXA}: {widths_dxa}")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        prevent_row_split(row)
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, color=GRID, size="4") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:color"), color)


def set_run_font(run, size=None, bold=None, italic=None, color=TEXT, name=FONT_CN) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_border_bottom(paragraph, color="B8C5D1", size="6", space="5") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, 9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    tail = paragraph.add_run(" 页")
    set_run_font(tail, 9, color=MUTED)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = FONT_CN
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_CN)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_CN)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[style_name]
        style.font.name = FONT_CN
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_CN)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_CN)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = FONT_CN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def set_running_furniture(section) -> None:
    header_p = section.header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_p.paragraph_format.space_after = Pt(2)
    run = header_p.add_run("少样本工业异常检测｜第二阶段设计参考")
    set_run_font(run, 9, bold=True, color=MUTED)
    set_paragraph_border_bottom(header_p, color="D7E0E8", size="4", space="4")
    footer_p = section.footer.paragraphs[0]
    add_page_number(footer_p)


def add_body(doc, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.keep_together = False
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2)
    else:
        run = p.add_run(text)
        set_run_font(run)


def add_bullet(doc, text: str, level=0) -> None:
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.left_indent = Inches(0.375 + 0.25 * level)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    set_run_font(p.add_run(text))


def create_numbering_instance(doc: Document) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
    ]
    abstract_id = max(abstract_ids, default=-1) + 100
    num_id = max(num_ids, default=0) + 100

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), "decimal")
    level.append(fmt)
    text = OxmlElement("w:lvlText")
    text.set(qn("w:val"), "%1.")
    level.append(text)
    justify = OxmlElement("w:lvlJc")
    justify.set(qn("w:val"), "left")
    level.append(justify)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "271")
    p_pr.append(indent)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    level.append(p_pr)
    abstract.append(level)
    first_num_index = next(
        (
            index
            for index, child in enumerate(numbering)
            if child.tag == qn("w:num")
        ),
        len(numbering),
    )
    numbering.insert(first_num_index, abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_number(doc, text: str, num_id: int) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_node])
    p_pr.insert(0, num_pr)
    set_run_font(p.add_run(text))


def add_callout(doc, title: str, body: str, fill=PALE_BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    set_table_borders(table, color=fill, size="2")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    set_run_font(p.add_run(title), 11, bold=True, color=DARK_BLUE)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    set_run_font(p2.add_run(body), 10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc, headers: list[str], rows: list[list[str]], widths_dxa: list[int]):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths_dxa)
    set_table_borders(table)
    header = table.rows[0]
    set_repeat_table_header(header)
    for i, text in enumerate(headers):
        set_cell_shading(header.cells[i], LIGHT_BLUE)
        p = header.cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(text), 9.5, bold=True, color=NAVY)
    for values in rows:
        cells = table.add_row().cells
        prevent_row_split(table.rows[-1])
        for i, value in enumerate(values):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            set_run_font(p.add_run(value), 9.2)
    set_table_geometry(table, widths_dxa)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_equation(doc, formula: str, explanation: str) -> None:
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [4320, 5040])
    set_table_borders(table, color="C9D5E0", size="4")
    set_cell_shading(table.cell(0, 0), PALE_BLUE)
    p1 = table.cell(0, 0).paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p1.add_run(formula), 9.5, bold=True, color=DARK_BLUE, name="Consolas")
    p2 = table.cell(0, 1).paragraphs[0]
    set_run_font(p2.add_run(explanation), 9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def make_diagram() -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    width, height = 1600, 980
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    font_path = Path("C:/Windows/Fonts/simhei.ttf")
    font = ImageFont.truetype(str(font_path), 38)
    small = ImageFont.truetype(str(font_path), 29)
    tiny = ImageFont.truetype(str(font_path), 24)
    bold = ImageFont.truetype(str(font_path), 42)

    def box(x1, y1, x2, y2, fill, title, lines):
        draw.rounded_rectangle((x1, y1, x2, y2), radius=24, fill=fill, outline="#65839A", width=4)
        draw.text(((x1 + x2) / 2, y1 + 42), title, font=font, fill="#203748", anchor="mm")
        y = y1 + 95
        for line in lines:
            draw.text(((x1 + x2) / 2, y), line, font=tiny, fill="#333333", anchor="mm")
            y += 38

    def arrow(x1, y1, x2, y2, color="#527A95"):
        draw.line((x1, y1, x2, y2), fill=color, width=7)
        import math
        angle = math.atan2(y2 - y1, x2 - x1)
        length = 22
        for offset in (2.55, -2.55):
            draw.line(
                (
                    x2,
                    y2,
                    x2 + length * math.cos(angle + offset),
                    y2 + length * math.sin(angle + offset),
                ),
                fill=color,
                width=7,
            )

    draw.text((800, 55), "第二阶段动态融合：一张测试图的完整数据流", font=bold, fill="#203748", anchor="mm")
    box(70, 150, 455, 395, "#E8EEF5", "视觉分支", ["看少量正常参考图", "输出图像分数 + 像素热图", "擅长产品具体外观"])
    box(70, 570, 455, 815, "#FFF5D6", "文本分支", ["看“正常/异常”文字概念", "输出图像分数 + 像素热图", "擅长通用异常语义"])
    box(610, 220, 990, 470, "#F4F7FA", "校准与可靠性", ["先把两种分数放到同一尺度", "再计算不确定性、一致性", "以及两分支的冲突程度"])
    box(610, 560, 990, 810, "#EAF4EC", "动态路由器", ["不确定性低的一支权重更大", "可整图分配权重", "也可逐像素分配权重"])
    box(1145, 330, 1530, 700, "#FCE8E6", "统一输出", ["融合后的图像异常分数", "融合后的像素异常热图", "视觉/文本权重", "路由决定和诊断特征"])
    arrow(455, 270, 610, 315)
    arrow(455, 690, 610, 380)
    arrow(800, 470, 800, 560)
    arrow(990, 685, 1145, 545)
    draw.text((800, 910), "红线规则：路由器不能读取测试真值、掩码、测试类别标签或整套测试集统计量", font=small, fill="#A1322A", anchor="mm")
    image.save(DIAGRAM, quality=95)


def add_cover(doc: Document) -> None:
    for _ in range(4):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(12)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(18)
    set_run_font(kicker.add_run("少样本工业异常检测 · 第二阶段参考指南"), 11, bold=True, color=BLUE)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    set_run_font(title.add_run("动态融合完整设计说明"), 29, bold=True, color=NAVY)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(8)
    set_run_font(subtitle.add_run("从“两位质检员”到“不确定性路由器”"), 15, color=DARK_BLUE)

    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    desc.paragraph_format.space_after = Pt(48)
    set_run_font(desc.add_run("用简单语言解释目标、输入、校准、可靠性、路由、实验与风险控制"), 10.5, italic=True, color=MUTED)

    meta = doc.add_table(rows=4, cols=2)
    set_table_geometry(meta, [2700, 6660])
    set_table_borders(meta, color="D5DEE6", size="3")
    metadata = [
        ("项目阶段", "第二阶段：动态融合算法设计与验证"),
        ("核心创新", "基于不确定性的图像级 / 像素级动态路由"),
        ("技术栈", "PyTorch 基线 + NumPy 动态融合与评测工具"),
        ("文档快照", "2026-07-31；以当前项目代码、配置和实验规范为准"),
    ]
    for row, (label, value) in zip(meta.rows, metadata):
        set_cell_shading(row.cells[0], LIGHT_BLUE)
        set_run_font(row.cells[0].paragraphs[0].add_run(label), 10, bold=True, color=NAVY)
        set_run_font(row.cells[1].paragraphs[0].add_run(value), 10)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(48)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("阅读方式：先看第 1—4 章理解“在做什么”，再看第 5—10 章理解“怎么做和怎么证明”。"), 9.5, color=MUTED)
    doc.add_page_break()


def build_document() -> None:
    make_diagram()
    doc = Document()
    style_document(doc)
    set_running_furniture(doc.sections[0])
    add_cover(doc)

    doc.add_heading("阅读摘要：一句话先讲明白", level=1)
    add_callout(
        doc,
        "一句话定义",
        "第二阶段不是重新训练一个很大的模型，而是在两个已有异常检测分支之间增加一个“调度员”：先估计每个分支此刻有多可靠，再决定整张图或每个像素更应该相信谁。",
        fill=PALE_GREEN,
    )
    add_body(doc, "可以把系统想成工厂里的两位质检员。第一位拿着少量合格品照片，擅长比较“这件产品和正常样品哪里不一样”；第二位理解“划痕、裂纹、缺口”等文字概念，擅长从通用语义判断异常。两个人各有强项，也都会犯错。动态融合模块不固定各占一半，而是根据当前证据临时分配权重。")

    doc.add_heading("这阶段最终要交付什么", level=2)
    add_bullet(doc, "一个统一的融合接口：不同视觉/文本基线可以替换，核心路由代码不用重写。")
    add_bullet(doc, "一个不使用测试真值的可靠性估计方法。")
    add_bullet(doc, "一个先做图像级、再做像素级的动态路由器。")
    add_bullet(doc, "完整对照实验、消融实验、可视化和失败案例分析。")
    add_bullet(doc, "冻结后的 VisA 与 MVTec AD 1/2/4-shot、3-seed 最终结果。")

    doc.add_heading("这阶段不做什么", level=2)
    add_bullet(doc, "不改变第一阶段已经冻结的数据划分、指标和基线预测。")
    add_bullet(doc, "不使用测试图标签、掩码或整套测试集统计量帮助路由器做决定。")
    add_bullet(doc, "不同时堆叠很多创新点；论文核心只保留“不确定性路由”。")
    add_bullet(doc, "不因为一次单类别结果好看就宣称方法有效。")

    doc.add_heading("1. 为什么需要动态融合", level=1)
    doc.add_heading("1.1 固定融合的问题", level=2)
    add_body(doc, "最简单的融合是把两个分支的分数相加，或者永远各占 50%。问题是：不同图片的难点不同。轻微划痕可能更适合文本语义分支；产品纹理或摆放方式的细微变化，可能更适合正常参考视觉分支。固定权重无法随样本变化。")
    add_body(doc, "另一个更隐蔽的问题是分数尺度不同。某分支输出 0.9，并不一定比另一分支的 0.6 更“确信”。如果不先校准，数值大的分支会长期压过另一分支，得到的只是分数偏置，不是真正的动态判断。")

    doc.add_heading("1.2 论文要解决的核心问题", level=2)
    add_callout(
        doc,
        "核心研究问题",
        "在只有 1/2/4 张正常参考图、没有目标异常训练样本的条件下，怎样估计视觉证据和文本证据在当前图像、当前局部区域上的可靠程度，并据此进行动态融合？",
    )
    add_body(doc, "因此，“少样本”是应用条件，“视觉 + 文本”是信息来源，“不确定性路由”才是论文要重点证明的新方法。")

    doc.add_heading("2. 两个分支分别在看什么", level=1)
    add_table(
        doc,
        ["分支", "它主要依据什么", "通常擅长什么", "可能失败的情况"],
        [
            ["视觉正常参考分支", "1/2/4 张同类别正常产品图；比较局部视觉特征", "产品特有纹理、结构和正常外观", "参考图太少、参考不典型、光照/姿态变化大"],
            ["文本语义分支", "“正常/异常”提示词及 CLIP 类视觉-文本知识", "通用异常语义、跨类别知识", "细小缺陷、行业专有外观、提示词不贴合"],
        ],
        [1900, 2500, 2380, 2580],
    )
    add_body(doc, "当前设计中，视觉分支以 AnomalyDINO 的冻结预测为主要输入；文本分支优先使用 AnomalyCLIP，早期也用 WinCLIP+ 做接口与对齐验证。分支本身仍由 PyTorch 运行，动态融合层目前主要用 NumPy 实现，便于低成本、可重复地处理冻结预测。")

    doc.add_heading("2.1 图像分数和像素热图", level=2)
    add_bullet(doc, "图像异常分数：回答“整张图是否异常”。分数越高，越可能异常。")
    add_bullet(doc, "像素异常热图：回答“异常可能在哪里”。热图中越亮的位置越可疑。")
    add_bullet(doc, "图像级融合负责分类；像素级融合负责定位。两者可以使用不同权重。")

    doc.add_heading("3. 整体系统结构", level=1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run()
    run.add_picture(str(DIAGRAM), width=Inches(6.35))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    set_run_font(cap.add_run("图 1  动态融合从两分支输入到统一输出的完整流程"), 9, italic=True, color=MUTED)

    doc.add_heading("3.1 一张测试图经过的 8 个步骤", level=2)
    flow_num_id = create_numbering_instance(doc)
    for text in [
        "读取同一张图在视觉分支和文本分支中的冻结预测。",
        "按 sample_id 对齐，确认两边确实是同一张图、同一顺序。",
        "把像素热图调整到统一大小，并检查数值是否有限。",
        "用正常参考图拟合的参数分别校准图像分数和像素分数。",
        "计算每个分支的不确定性、稳定性和响应集中程度。",
        "计算两分支之间的一致、分歧与冲突。",
        "路由器生成图像权重和可选的逐像素权重。",
        "输出融合分数、融合热图、权重、路由决定和诊断特征。",
    ]:
        add_number(doc, text, flow_num_id)

    doc.add_heading("4. 第一道关键：分数校准", level=1)
    add_body(doc, "校准可以理解为“先把两把刻度不同的尺子换算成同一种单位”。视觉分支和文本分支来自不同模型，原始分数不能直接比较。项目曾出现 200/200 个样本都被路由到文本分支，这正说明原始尺度存在偏置；它是故障诊断，不是性能结论。")

    doc.add_heading("4.1 校准数据从哪里来", level=2)
    add_bullet(doc, "只读取 few-shot 清单中的正常训练参考图。")
    add_bullet(doc, "每张参考图使用预先固定的 5 个视图：原图、亮度 0.90/1.10、对比度 0.90/1.10。")
    add_bullet(doc, "视觉与文本分支使用同一组 source_id 和 augmentation_id。")
    add_bullet(doc, "绝不读取目标测试图、测试标签或测试掩码来拟合参数。")

    doc.add_heading("4.2 当前稳健校准方法", level=2)
    add_equation(
        doc,
        "z = (s - median) / (1.4826 × MAD + ε)",
        "s 是原始分数；median 是正常参考分数的中位数；MAD 是到中位数距离的中位数；ε 防止除以 0。它比均值/标准差更不容易被单个异常值带偏。",
    )
    add_equation(
        doc,
        "p = sigmoid(z / T)",
        "把换算后的 z 压到 0～1；T 是温度，控制曲线陡峭程度。最终得到可比较的“异常概率式分数”。",
    )
    add_body(doc, "像素热图不能简单把所有背景像素摊平后求中位数，因为正常图中大量像素接近 0，会让尺度退化。当前实现先取每个参考视图热图的 q99 尾部分数，再拟合像素校准器。这样仍只使用正常参考数据，同时更关注正常图中最高的一小部分响应。")

    doc.add_heading("4.3 校准要比较的三种候选", level=2)
    add_table(
        doc,
        ["方案", "含义", "作用"],
        [
            ["不校准", "直接使用两分支原始分数", "作为反例，观察尺度偏置有多严重"],
            ["固定温度缩放", "使用预先设定的统一温度", "判断简单缩放是否足够"],
            ["正常参考稳健校准", "每类别、每分支、图像/像素分别拟合", "主方案；不使用测试数据"],
        ],
        [1900, 4000, 3460],
    )

    doc.add_heading("5. 第二道关键：怎样判断“谁更可靠”", level=1)
    add_body(doc, "可靠性不是只看分数高低，而是看这个分支的判断是否稳定、是否犹豫、是否与另一分支明显冲突。项目把这些信息整理成“可靠性特征”，再交给路由器。")
    add_table(
        doc,
        ["特征", "简单解释", "数值高通常说明"],
        [
            ["概率熵", "分数靠近 0.5 时更犹豫；靠近 0 或 1 时更明确", "不确定性更大"],
            ["增强一致性", "同一正常参考图稍微改变亮度/对比度后，结果是否变化", "变化大则更不稳定"],
            ["跨 shot 敏感度", "使用 1/2/4 张参考图时，结果是否波动", "波动大则更依赖参考选择"],
            ["响应集中程度", "热图是否集中在少数区域，而不是整张图到处发亮", "需要结合分支和场景解释"],
            ["分支分歧", "视觉和文本的分数/热图相差多大", "两位“质检员”意见冲突"],
            ["层间一致性（后续）", "同一模型不同特征层是否给出相近结论", "变化大则内部证据不稳定"],
        ],
        [2050, 4700, 2610],
    )

    doc.add_heading("5.1 为什么不能只用熵", level=2)
    add_body(doc, "熵要求输入已经像概率。如果原始分数没有校准，0.9 和 0.6 并不能公平比较；而且一个模型可能“自信地犯错”。因此主设计需要把熵、增强稳定性、跨 shot 稳定性、分支冲突等多种证据结合起来。")

    doc.add_heading("6. 第三道关键：动态路由器怎样分配权重", level=1)
    add_body(doc, "第一版路由器采用确定性规则：不确定性越低，权重越大。它不训练大型网络，便于排错、复现和证明数据没有泄漏。后续只有在规则路由不足时，才比较逻辑回归或很小的全连接网络。")
    add_equation(
        doc,
        "wᵥ = exp(-uᵥ/T) / [exp(-uᵥ/T)+exp(-uₜ/T)]",
        "uᵥ、uₜ 分别是视觉与文本不确定性；T 控制权重变化速度；wᵥ 是视觉权重，文本权重为 1-wᵥ。不确定性低的一支得到更大权重。",
    )
    add_equation(
        doc,
        "s_fused = wᵥ·sᵥ + (1-wᵥ)·sₜ",
        "融合分数等于两个校准后分数的加权和。像素热图使用同样思想，但每个像素可以拥有自己的权重。",
    )

    doc.add_heading("6.1 三种路由决定", level=2)
    add_bullet(doc, "视觉主导：视觉权重大于 0.5 加决策边界。")
    add_bullet(doc, "文本主导：视觉权重小于 0.5 减决策边界。")
    add_bullet(doc, "加权融合：两者差距不够大时，不强行二选一。")
    add_body(doc, "当前配置中的 temperature=0.20、min_weight=0.05、decision_margin=0.15 都是显式参数。min_weight 表示即使某一支明显更可靠，也至少保留另一支 5% 的影响，避免路由彻底坍缩成单分支。")

    doc.add_heading("6.2 为什么先图像级、后像素级", level=2)
    add_table(
        doc,
        ["层级", "做法", "优点", "风险"],
        [
            ["图像级", "一张图只生成一个视觉权重", "结构简单、易解释、先验证总方向", "同一图中不同区域不能分别选择"],
            ["像素级", "每个像素生成视觉权重", "能让划痕区域信文本、纹理区域信视觉", "容易产生噪声、计算和调参更复杂"],
        ],
        [1500, 3000, 2500, 2360],
    )
    add_callout(doc, "推进原则", "必须先让校准和图像级路由通过，才能启动像素级动态路由。否则很难判断问题来自分数尺度、权重规则还是空间噪声。", fill=PALE_YELLOW)

    doc.add_heading("7. 数据泄漏：最重要的实验红线", level=1)
    add_body(doc, "数据泄漏是指方法在开发时偷看了最终测试答案，导致结果看起来很好，但换到真正未知的数据就失效。动态路由尤其容易泄漏，因为权重和阈值可以被测试结果“调得刚刚好”。")
    add_table(
        doc,
        ["允许进入设计/路由", "禁止进入设计/路由"],
        [
            ["合成数据；VisA seed 0 冻结预测；few-shot 正常训练参考图；预先声明的源域验证数据", "测试图真值标签；测试掩码；类别测试标签；整套测试集均值/方差/分位数；看测试指标后选择的权重"],
        ],
        [4680, 4680],
    )
    add_body(doc, "最终评测器可以读取真值来计算 AUROC、AP、AUPRO 等指标，但必须与路由模块分开。也就是说，“考试结束后批改答案”可以，“考试时偷看答案”不可以。")

    doc.add_heading("8. 完整实验设计", level=1)
    doc.add_heading("8.1 开发、冻结、最终验证三段式", level=2)
    add_table(
        doc,
        ["阶段", "允许数据", "目的", "能否改设计"],
        [
            ["开发", "合成数据 + VisA seed 0 冻结缓存", "排错、比较候选、确定结构和参数", "可以，但必须记录原因"],
            ["冻结", "不再查看新的最终结果", "锁定校准、特征、路由、阈值与超参数", "冻结后不可以"],
            ["最终验证", "VisA seed 1/2；MVTec seed 0/1/2", "一次性检验泛化能力", "不可以；失败也如实报告"],
        ],
        [1400, 2860, 2860, 2240],
    )

    doc.add_heading("8.2 必须比较的基线", level=2)
    add_bullet(doc, "视觉单分支：只使用视觉正常参考证据。")
    add_bullet(doc, "文本单分支：只使用文本语义证据。")
    add_bullet(doc, "简单平均：视觉权重固定为 0.5。")
    add_bullet(doc, "固定权重：视觉权重预先登记为 0、0.25、0.5、0.75、1。")
    add_bullet(doc, "图像级动态路由：每张图自适应一个权重。")
    add_bullet(doc, "像素级动态路由：每个像素自适应权重。")

    doc.add_heading("8.3 消融实验要回答什么", level=2)
    add_table(
        doc,
        ["消融", "拿掉或替换什么", "想证明什么"],
        [
            ["无校准", "跳过分数校准", "校准是否是公平融合的必要条件"],
            ["只用熵", "去掉稳定性和冲突特征", "多种可靠性证据是否有额外价值"],
            ["去掉增强一致性", "不看亮度/对比度扰动", "参考稳定性是否有用"],
            ["去掉跨 shot 敏感度", "不比较 1/2/4-shot 波动", "few-shot 稳定性是否有用"],
            ["固定权重替代动态权重", "所有样本使用同一权重", "动态路由是否真的优于静态融合"],
            ["图像级替代像素级", "整图共用权重", "局部动态权重是否值得增加复杂度"],
        ],
        [2200, 3600, 3560],
    )

    doc.add_heading("8.4 指标和附加报告", level=2)
    add_bullet(doc, "图像级：Image AUROC、Image AP、最大 F1。")
    add_bullet(doc, "像素级：Pixel AUROC、Pixel AP、AUPRO。")
    add_bullet(doc, "统计：3 个 seed 的均值和标准差。")
    add_bullet(doc, "效率：额外运行时间、峰值显存、缓存大小。")
    add_bullet(doc, "解释：权重分布、路由比例、冲突样本、成功与失败热图。")

    doc.add_heading("9. 工作包与项目当前状态", level=1)
    add_callout(
        doc,
        "状态说明",
        "以下是 2026-07-31 的工程快照。“代码/测试完成”只表示线路能运行，不等于动态融合已经在最终数据上证明优于基线。",
        fill=PALE_YELLOW,
    )
    add_table(
        doc,
        ["工作包", "主要内容", "当前状态", "完成判定"],
        [
            ["WP1", "缓存格式、sample_id 对齐、跨 shot 一致性、AnomalyCLIP sidecar", "已完成", "12 类审计与顺序检查通过"],
            ["WP2", "正常参考校准；图像/像素分别拟合；q99 像素尾部", "真实参考线路已通过，仍需正式对照后冻结", "参数来源、SHA256、无测试数据声明齐全"],
            ["WP3", "熵、增强一致性、跨 shot 敏感度、响应集中程度", "主要接口已完成", "真实校准数值验证；层间一致性待定"],
            ["WP4", "两分支一致、分歧和冲突特征", "已完成接口", "真实缓存上的数值与形状检查通过"],
            ["WP5", "单分支、平均、固定权重对照", "代码完成；开发矩阵正在/已开始生成", "统一评测与机器可读报告齐全"],
            ["WP6", "图像级规则路由及小模型候选比较", "待正式推进", "优于合理固定基线且不过度偏向单支"],
            ["WP7", "像素级动态路由", "尚未启动正式验证", "图像级通过后再做；空间权重稳定"],
            ["WP8", "可视化、失败案例、运行代价", "待后续", "图表可复核，结论与限制一致"],
        ],
        [950, 3330, 2320, 2760],
    )
    add_body(doc, "截至该快照，真实 VisA seed 0、1-shot 正常参考校准 v6_q99 已记录为 12/12 类通过，并明确 test_predictions_used=false、test_labels_used=false。项目还建立了 visual/text/fixed/dynamic 四种统一运行模式和回归测试。")

    doc.add_heading("10. 后续实施顺序", level=1)
    implementation_num_id = create_numbering_instance(doc)
    for text in [
        "完成并验收 VisA seed 0、1-shot 的开发矩阵：视觉、文本、固定权重和当前规则路由。",
        "检查校准后的分数分布、路由比例与失败样本，确认没有单分支坍缩。",
        "冻结图像级可靠性特征与规则路由参数；只在必要时比较逻辑回归或小型 MLP。",
        "在 1/2/4-shot 上检查跨 shot 稳定性，决定该特征如何进入路由。",
        "图像级路由通过后实现像素级平滑、边界限制和空间稳定性检查。",
        "完成固定权重、图像级动态、像素级动态和特征消融。",
        "冻结全部设计与配置，生成不可修改的配置快照和输入 SHA256。",
        "只运行一次 VisA seed 1/2 与 MVTec 3-seed 最终验证。",
        "汇总均值/标准差、运行时间、显存、权重图、热图和失败案例，形成论文表格。",
    ]:
        add_number(doc, text, implementation_num_id)

    doc.add_heading("10.1 每一步都必须留下的记录", level=2)
    add_table(
        doc,
        ["文件", "作用"],
        [
            ["run.json", "记录数据集、seed、shot、分支、输入路径、SHA256、Git 状态和是否使用 GPU"],
            ["command.txt", "保存可以原样重复执行的命令"],
            ["config.yaml", "保存本次超参数快照，避免依赖隐藏默认值"],
            ["stdout.log", "保存正常输出与错误信息"],
            ["report.json / report.csv", "机器可读结果和人工查看表格"],
            ["decision.md", "说明结论、限制、失败原因和下一步"],
        ],
        [2300, 7060],
    )

    doc.add_heading("11. 主要风险与应对", level=1)
    add_table(
        doc,
        ["风险", "会发生什么", "应对方式"],
        [
            ["分数尺度不一致", "路由长期偏向数值大的分支", "先校准；保存分布；与不校准对照"],
            ["参考图太少或不典型", "视觉分支不稳定", "增强一致性、跨 shot 敏感度、3-seed"],
            ["路由坍缩", "几乎所有样本只走一个分支", "最小权重、路由比例审计、固定权重对照"],
            ["像素权重噪声", "热图出现碎点或边缘抖动", "先图像级；空间平滑；检查边界与连通区域"],
            ["过拟合 VisA seed 0", "开发集好，最终集下降", "提前冻结；seed 1/2 和 MVTec 只做一次最终验证"],
            ["数据泄漏", "指标虚高且不可泛化", "API 禁止真值字段；审计 test_*_used=false"],
            ["创新点重复", "与已有融合方法差异不清", "突出局部可靠性估计与无真值路由；客观比较 ReMP-AD、VCP-CLIP 等"],
            ["方法过于复杂", "难复现、难解释", "规则路由先行；小模型仅作为候选；保存显式配置"],
        ],
        [1700, 3480, 4180],
    )

    doc.add_heading("12. 怎样判断第二阶段真正成功", level=1)
    add_body(doc, "成功不等于“某一个指标比最好基线高一点”。至少应同时满足下面四类条件：")
    add_bullet(doc, "有效性：动态路由在多个数据集、shot 和 seed 上稳定优于单分支与合理固定权重。")
    add_bullet(doc, "定位能力：像素指标和热图确实改善，而不是只提高图像分类分数。")
    add_bullet(doc, "可信性：没有测试泄漏，冻结协议可复核，消融能说明每个关键设计的作用。")
    add_bullet(doc, "实用性：额外时间和显存可接受，接口能替换不同基线，配置和结果能复现。")
    add_callout(
        doc,
        "客观的失败判定",
        "如果动态路由只在 VisA seed 0 有效、最终验证不稳定，或者不如一个预先固定的权重，那么应如实结论为“当前可靠性特征或路由规则不足”，而不是继续看最终测试结果反复调参。",
        fill=PALE_RED,
    )

    doc.add_heading("13. 技术栈：项目用 PyTorch 还是 TensorFlow", level=1)
    add_callout(doc, "答案", "项目的深度学习基线以 PyTorch 为主，不是 TensorFlow。动态融合当前多数是对冻结预测做 NumPy 计算，因此它本身不需要长期占用 GPU。", fill=PALE_GREEN)
    add_table(
        doc,
        ["组件", "主要工具", "是否通常占 GPU"],
        [
            ["AnomalyDINO / AnomalyCLIP / PromptAD 等基线", "PyTorch、torchvision、CUDA", "是；特征提取或训练时占用"],
            ["动态融合校准与规则路由", "NumPy + 项目自定义 Python 模块", "通常否；CPU 即可"],
            ["统一评测", "NumPy、scikit-learn、图像处理工具", "通常否"],
            ["后续小型可训练路由器（若采用）", "建议仍用 PyTorch", "占用很少，可与大基线错峰"],
        ],
        [2800, 3600, 2960],
    )

    doc.add_heading("14. 常见问题", level=1)
    faq = [
        ("Q1：动态融合是不是把两个模型一起重新训练？", "不是。当前优先使用已经冻结的预测缓存，在其上做校准、可靠性估计和权重融合。这样更省 GPU，也不会破坏第一阶段基线。"),
        ("Q2：权重越接近 1 就一定越好吗？", "不是。权重只表示更信视觉分支。关键是权重是否与实际可靠性相关，并且最终指标、热图和失败案例是否支持。"),
        ("Q3：为什么不能用测试集均值做归一化？", "因为整套测试集本身包含最终考试信息。即使不直接读取标签，也会让每张测试图依赖其他测试图，影响独立性和可部署性。"),
        ("Q4：像素级动态权重一定比图像级好吗？", "不一定。它更灵活，但也更容易产生噪声和过拟合。必须通过对照实验决定是否值得。"),
        ("Q5：少样本只使用正常图，会不会看不到异常？", "文本分支提供通用异常知识；视觉分支学习目标产品的正常外观。两者互补正是本阶段的出发点。"),
        ("Q6：目前能不能写“动态融合优于基线”？", "还不能只凭工程测试这样写。需要完成冻结后的 VisA 与 MVTec 多 seed 最终验证，才能形成性能结论。"),
    ]
    for question, answer in faq:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        set_run_font(p.add_run(question), 11, bold=True, color=DARK_BLUE)
        add_body(doc, answer)

    doc.add_heading("15. 术语小词典", level=1)
    add_table(
        doc,
        ["术语", "简单解释"],
        [
            ["Few-shot / K-shot", "每个类别只给 K 张正常参考图；本项目 K=1、2、4。"],
            ["Branch / 分支", "独立给出异常判断的一套方法，如视觉分支或文本分支。"],
            ["Calibration / 校准", "把不同模型的分数换算到可比较的尺度。"],
            ["Uncertainty / 不确定性", "模型对当前判断有多犹豫或不稳定。"],
            ["Router / 路由器", "根据可靠性给两个分支分配权重的模块。"],
            ["Pixel map / 热图", "每个像素的异常程度，用于定位缺陷。"],
            ["Ablation / 消融", "拿掉某个组件再实验，判断它是否真的有用。"],
            ["Seed / 随机种子", "控制参考图抽样和随机过程；多个 seed 用于检查稳定性。"],
            ["Freeze / 冻结", "提前锁定方法和参数，之后不再根据最终测试结果修改。"],
            ["Leakage / 泄漏", "开发时使用了本不应看到的测试答案或整体统计信息。"],
        ],
        [2400, 6960],
    )

    doc.add_heading("附录 A：当前代码与配置入口", level=1)
    add_table(
        doc,
        ["文件", "内容"],
        [
            ["SECOND_STAGE_PLAN.md", "第二阶段工作包、依赖、完成条件和实时更新"],
            ["docs/dynamic_fusion_design.md", "输入输出、开发边界和实现记录"],
            ["docs/dynamic_fusion_experiment_protocol.md", "运行目录、证据文件和泄漏检查规范"],
            ["configs/dynamic_fusion.yaml", "路由、校准、特征、固定权重和最终验证配置"],
            ["src/industrial_ad/fusion/calibration.py", "正常参考稳健校准与安全加载"],
            ["src/industrial_ad/fusion/features.py", "熵、一致性、敏感度、集中度和分支冲突"],
            ["src/industrial_ad/fusion/router.py", "确定性不确定性加权路由器"],
            ["src/industrial_ad/fusion/baselines.py", "视觉/文本单分支和固定权重对照"],
        ],
        [3900, 5460],
    )

    doc.add_heading("附录 B：设计冻结清单", level=1)
    checklist = [
        "视觉与文本输入来自同一数据划分、同一 sample_id、同一评测协议。",
        "图像与像素校准器只使用正常参考图，参数文件带来源和 SHA256。",
        "所有路由参数都写入配置文件，没有隐藏默认值。",
        "开发只使用合成数据和 VisA seed 0 冻结缓存。",
        "固定权重候选在查看最终结果前已经登记。",
        "图像级路由、像素级路由和所有消融方案已经确定。",
        "测试真值只由独立评测器读取，路由 API 不含真值字段。",
        "失败运行保留记录，不覆盖通过的产物。",
        "冻结后 VisA seed 1/2 和 MVTec 结果不再反向修改设计。",
        "最终报告包含均值、标准差、效率、可视化和失败案例。",
    ]
    for item in checklist:
        add_bullet(doc, item)

    doc.add_heading("结语", level=1)
    add_body(doc, "第二阶段的价值不在于简单地把两个分数相加，而在于回答一个更有研究意义的问题：当视觉正常参考和文本语义证据发生冲突时，系统能否在不偷看测试答案的前提下，判断当前图像、当前局部区域更应该相信谁。")
    add_body(doc, "最稳妥的路线是：先保证输入对齐和分数校准，再验证图像级规则路由，最后才进入像素级动态权重；全过程保存配置、命令、输入哈希、机器可读结果和失败记录。这样得到的结果才容易解释、复现，也更适合作为论文的核心贡献。")

    doc.core_properties.title = "第二阶段动态融合完整设计说明"
    doc.core_properties.subject = "少样本工业异常检测中的不确定性路由"
    doc.core_properties.author = "少样本工业异常检测项目组"
    doc.core_properties.keywords = "工业异常检测, 少样本, 动态融合, 不确定性路由, PyTorch"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()

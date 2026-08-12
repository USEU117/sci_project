from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "outputs" / "动态融合并行开发完整计划.docx"

NAVY = "17365D"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "667085"
GREEN = "2F6B4F"
GOLD = "8A6500"
RED = "9B1C1C"
WHITE = "FFFFFF"
BLACK = "000000"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths_dxa[index] / 1440)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_font(run, size=11, bold=False, color=BLACK, italic=False) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    pf = normal.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.25

    for style_name, size, color, before, after in (
        ("Title", 28, NAVY, 0, 8),
        ("Subtitle", 13, MID_GRAY, 0, 16),
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = style_name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_font(run, 9, color=MID_GRAY)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    tail = paragraph.add_run(" 页")
    set_font(tail, 9, color=MID_GRAY)


def configure_page(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_run = header.add_run("少样本工业异常检测项目｜动态融合设计轨")
    set_font(header_run, 9, bold=True, color=MID_GRAY)
    add_page_field(section.footer.paragraphs[0])


def add_title_page(doc: Document) -> None:
    for _ in range(5):
        doc.add_paragraph()
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = kicker.add_run("第二阶段并行开发计划")
    set_font(run, 11, bold=True, color=GOLD)
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("动态融合算法设计与验证计划")
    set_font(run, 28, bold=True, color=NAVY)
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("在第一阶段基线训练完成前可并行推进的完整工作方案")
    set_font(run, 13, color=MID_GRAY)
    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(meta.add_run("项目：少样本工业异常检测｜文本—视觉证据动态融合"), 10.5, color=MID_GRAY)
    meta2 = doc.add_paragraph()
    meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(meta2.add_run("版本：2026-07-30｜状态：基线轨与设计轨并行"), 10.5, color=MID_GRAY)
    doc.add_paragraph()
    lead = doc.add_paragraph()
    lead.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lead.paragraph_format.left_indent = Inches(0.55)
    lead.paragraph_format.right_indent = Inches(0.55)
    lead.paragraph_format.space_before = Pt(18)
    lead.paragraph_format.space_after = Pt(0)
    set_font(
        lead.add_run(
            "核心原则：后台训练继续按原协议完成；动态融合只使用合成数据和冻结的 VisA seed 0 预测缓存开发。"
            "在设计冻结之前，不查看 VisA seed 1/2 与 MVTec 的最终融合结果。"
        ),
        11,
        bold=True,
        color=DARK_BLUE,
    )
    doc.add_page_break()


def add_paragraph(doc, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        set_font(p.add_run(bold_prefix), bold=True)
        set_font(p.add_run(text[len(bold_prefix):]))
    else:
        set_font(p.add_run(text))


def add_bullets(doc, items: list[str], level: int = 0) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        if level:
            p.paragraph_format.left_indent = Inches(0.375 + level * 0.25)
        set_font(p.add_run(item))


def add_numbers(doc, items: list[str]) -> None:
    for index, item in enumerate(items, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.28)
        p.paragraph_format.first_line_indent = Inches(-0.22)
        set_font(p.add_run(f"{index}. "))
        set_font(p.add_run(item))


def add_callout(doc, label: str, text: str, fill=LIGHT_BLUE, color=DARK_BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    set_font(p.add_run(f"{label}："), bold=True, color=color)
    set_font(p.add_run(text), color=BLACK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(p.add_run(header), 10, bold=True, color=NAVY)
    for row_data in rows:
        row = table.add_row()
        for index, value in enumerate(row_data):
            cell = row.cells[index]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if index > 0 else WD_ALIGN_PARAGRAPH.CENTER
            set_font(p.add_run(value), 9.5)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def heading(doc, text: str, level: int = 1) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    set_font(p.add_run(text), {1: 16, 2: 13, 3: 12}[level], True, {1: BLUE, 2: BLUE, 3: DARK_BLUE}[level])


def page_break(doc) -> None:
    doc.add_page_break()


def build() -> None:
    doc = Document()
    configure_page(doc)
    configure_styles(doc)
    add_title_page(doc)

    heading(doc, "一、计划结论与总体建议")
    add_callout(
        doc,
        "结论",
        "可以在第一阶段耗时训练尚未全部结束时，提前开展第二阶段动态融合设计。"
        "但必须把“算法开发数据”和“最终验证数据”严格分开，避免为了得到更好结果而反复查看最终测试集。",
    )
    add_paragraph(
        doc,
        "建议采用双轨推进方式。基线轨继续完成 PromptAD、ReMP-AD、AdaptCLIP 和 MVTec 完整矩阵；"
        "设计轨同步完成接口、校准、不确定性、路由器、单元测试、分析工具和实验协议。"
        "这样能够利用等待训练的时间，同时不破坏第一阶段已经固定的划分、指标和结果。"
    )
    add_table(
        doc,
        ["轨道", "主要工作", "允许使用的数据", "第一阶段结束前目标"],
        [
            ["基线轨", "继续执行官方复现和统一矩阵", "既定 manifest、官方权重、统一协议", "形成可审计的完整基线"],
            ["设计轨", "开发动态融合接口、校准和路由", "合成数据、VisA seed 0 冻结缓存、正常参考图", "形成已测试、待最终验证的设计候选"],
        ],
        [1200, 2900, 2900, 2360],
    )

    heading(doc, "二、项目当前基础")
    add_paragraph(doc, "第一阶段已经为第二阶段提供了较完整的工程基础。")
    add_bullets(
        doc,
        [
            "MVTec AD 与 VisA 数据已经校验，1/2/4-shot、seed 0/1/2 的统一抽样清单已经固定。",
            "PatchCore、WinCLIP+、AnomalyDINO 已完成 VisA 全矩阵；AnomalyCLIP 已完成 VisA 与 MVTec 全类别推理。",
            "PromptAD 的 VisA seed 0、1/2/4-shot 已完成，并继续在后台运行 seed 1/2。",
            "统一 NPZ 输出和六项指标已经稳定：图像 AUROC、AP、F1-max，像素 AUROC、AP、AUPRO。",
            "动态融合第一版已建立统一输入输出接口、置信度路由器、配置文件和冻结缓存运行脚本。",
            "合成测试与原有评测回归测试共 10 项通过。",
        ],
    )
    heading(doc, "已经发现的关键问题", 2)
    add_paragraph(
        doc,
        "第一次冻结缓存测试使用 AnomalyDINO 作为视觉分支、WinCLIP+ 作为文本引导分支。"
        "200 张 candle 图片全部被路由到文本分支。程序本身正常，但两个分支的原始分数尺度不同，"
        "直接用原始分数计算不确定性会产生偏置。因此，校准必须成为第二阶段的第一项核心研究。"
    )

    heading(doc, "三、第二阶段要解决的核心问题")
    add_numbers(
        doc,
        [
            "把不同方法的图像分数和像素图转换到可比较的尺度。",
            "估计每个分支对当前图片和当前区域有多确定。",
            "衡量视觉分支与文本分支的判断是否一致。",
            "根据不确定性和一致性，决定视觉主导、文本主导或加权融合。",
            "在图像级路由稳定后，进一步研究像素级动态权重。",
            "证明动态融合比单分支和固定权重融合更稳定，而不是只在一个 seed 上偶然变好。",
        ],
    )
    add_callout(
        doc,
        "通俗理解",
        "文本分支像“懂语义的检查员”，视觉分支像“擅长对比细节的检查员”。"
        "动态路由器是第三位裁判，它不重新检查产品，只判断这一次更应该相信谁。",
        fill=LIGHT_GRAY,
    )

    page_break(doc)
    heading(doc, "四、双轨并行的边界")
    heading(doc, "4.1 基线轨必须保持不变", 2)
    add_bullets(
        doc,
        [
            "不更改数据划分、shot 清单、seed、指标定义和现有结果。",
            "不为了适配融合而修改单个基线的预测逻辑。",
            "每次成功运行保留配置、日志、checkpoint、预测 NPZ 和注册表记录。",
            "失败运行保留错误日志，不覆盖已成功结果。",
            "同一 GPU 不同时启动两个训练任务。",
        ],
    )
    heading(doc, "4.2 设计轨允许做什么", 2)
    add_bullets(
        doc,
        [
            "使用合成数据测试形状、边界值、NaN/Inf 和样本顺序。",
            "读取已经冻结的 VisA seed 0 预测缓存。",
            "实现分数校准、固定融合、图像级路由和像素级路由。",
            "生成权重图、路由决策和分支冲突分析。",
            "用 CPU 完成大部分接口、统计和评测工作。",
        ],
    )
    heading(doc, "4.3 设计轨禁止做什么", 2)
    add_bullets(
        doc,
        [
            "路由器不能接收测试图片的真实标签和真实掩码。",
            "不能用完整测试集的均值、标准差、最大值或分位数做归一化。",
            "不能看 VisA seed 1/2 或 MVTec 最终结果后再修改结构。",
            "不能根据每个测试类别的最终得分单独设置权重。",
            "不能把目标测试异常样本用于训练、校准或阈值选择。",
        ],
    )

    heading(doc, "五、动态融合的统一输入输出")
    heading(doc, "5.1 输入", 2)
    add_table(
        doc,
        ["输入", "含义", "要求"],
        [
            ["sample_id", "样本编号", "两分支完全一致且顺序可核对"],
            ["image_score", "整张图的异常分数", "越大越异常，必须完成校准"],
            ["pixel_map", "每个像素的异常分数", "尺寸一致，越大越异常"],
            ["uncertainty", "分支的犹豫程度", "只能从允许的数据和预测中得到"],
            ["disagreement", "两分支的差异", "同时计算图像级和像素级"],
        ],
        [1500, 3300, 4560],
    )
    heading(doc, "5.2 输出", 2)
    add_bullets(
        doc,
        [
            "融合后的图像异常分数。",
            "融合后的像素异常图。",
            "视觉分支的图像级权重和像素级权重图。",
            "路由决定：视觉主导、文本主导或加权融合。",
            "用于分析的不确定性和分支差异特征。",
        ],
    )
    heading(doc, "5.3 接口验收标准", 2)
    add_bullets(
        doc,
        [
            "输入中不存在 gt、label、mask 等真值字段。",
            "样本数量、顺序或像素图尺寸不一致时立即报错。",
            "发现 NaN 或无穷大时立即报错。",
            "所有权重位于 0 到 1，视觉权重与文本权重之和为 1。",
            "输出能够保存为统一 NPZ，并由现有评测脚本直接读取。",
        ],
    )

    heading(doc, "六、第一阶段结束前可以并行完成的工作包")

    heading(doc, "WP1：预测缓存与样本对齐", 2)
    add_paragraph(doc, "目的：保证视觉和文本分支比较的是同一张图片、同一个位置。")
    add_bullets(
        doc,
        [
            "统一不同方法的 sample_id 命名。",
            "按 sample_id 重排缓存，禁止只依赖文件顺序。",
            "统一像素图尺寸和插值方式；异常图用双线性，真值掩码只在评测阶段用最近邻。",
            "为缺少 sample_id 的 AnomalyCLIP 旧缓存生成可复核的 sidecar，或重新导出缓存。",
            "生成每个类别的对齐报告：样本数、缺失项、重复项、标签一致性和形状。",
        ],
    )
    add_callout(doc, "验收", "12 个 VisA 类别均能在不使用真值参与融合的情况下完成两分支对齐。", fill=LIGHT_GRAY)

    heading(doc, "WP2：分数校准", 2)
    add_paragraph(doc, "目的：让不同分支的分数可以公平比较。")
    add_bullets(
        doc,
        [
            "记录每个分支的原始分数范围和分布。",
            "建立不做校准、固定温度缩放、正常参考分布校准三种候选。",
            "校准参数只能来自源域验证数据或目标正常参考图。",
            "图像分数和像素图分别校准，不能假设二者尺度相同。",
            "检查校准后正常样本是否接近低异常分数，不能用异常测试标签调参。",
        ],
    )
    add_callout(
        doc,
        "验收",
        "同一批合成输入不会因分支原始数值范围不同而全部路由到单一分支；校准过程不读取测试标签。",
        fill=LIGHT_GRAY,
    )

    heading(doc, "WP3：不确定性特征", 2)
    add_paragraph(doc, "目的：估计每个分支对当前判断有多犹豫。")
    add_bullets(
        doc,
        [
            "概率熵：分数越接近 0.5，通常越不确定。",
            "增强一致性：同一张正常参考图做轻微增强后，预测变化越大，越不稳定。",
            "参考敏感度：更换 1/2/4-shot 中的正常参考图后，预测变化越大，越不稳定。",
            "像素图集中程度：异常响应过度分散可能代表低可靠性。",
            "分支内部层间一致性：不同特征层结论差异大时，提高不确定性。",
        ],
    )
    add_paragraph(
        doc,
        "第一版优先使用容易解释、成本低的特征。只有这些特征稳定后，才考虑使用小型神经网络学习不确定性。"
    )

    heading(doc, "WP4：分支一致性与冲突特征", 2)
    add_bullets(
        doc,
        [
            "图像分数差值。",
            "两张异常图的平均绝对差。",
            "高响应区域的重合比例。",
            "最大异常点位置的距离。",
            "异常区域面积差。",
            "两分支是否同时判断正常、同时判断异常或发生冲突。",
        ],
    )
    add_callout(doc, "注意", "一致性只能说明两分支是否相似，不能单独说明谁是正确的。", fill="FFF6DD", color=GOLD)

    heading(doc, "WP5：固定融合基线", 2)
    add_paragraph(doc, "目的：为动态融合提供必须超过的简单参照。")
    add_bullets(
        doc,
        [
            "只使用视觉分支。",
            "只使用文本分支。",
            "视觉 0.5 + 文本 0.5。",
            "视觉 0.7 + 文本 0.3。",
            "视觉 0.3 + 文本 0.7。",
        ],
    )
    add_paragraph(
        doc,
        "这些权重必须在开发协议中提前列出。动态方法如果不能稳定超过固定融合，就不能说明路由器真正有效。"
    )

    heading(doc, "WP6：图像级动态路由", 2)
    add_paragraph(doc, "目的：为每张图片给出一个视觉权重。")
    add_numbers(
        doc,
        [
            "先实现规则路由：比较两分支校准后的不确定性。",
            "视觉明显更可靠时输出“视觉主导”。",
            "文本明显更可靠时输出“文本主导”。",
            "差距不明显时输出“加权融合”。",
            "保存每张图片的权重、决定和输入特征，便于人工检查。",
            "规则路由稳定后，再比较小型全连接网络或逻辑回归。",
        ],
    )
    add_callout(doc, "验收", "路由决定可解释、可重复；相同输入得到相同输出；不使用测试真值。", fill=LIGHT_GRAY)

    heading(doc, "WP7：像素级动态路由", 2)
    add_paragraph(doc, "目的：让同一张图的不同区域选择不同分支。")
    add_bullets(
        doc,
        [
            "输出与异常图同尺寸的视觉权重图。",
            "对每个像素执行视觉图与文本图的加权融合。",
            "限制权重在 0 到 1，并增加空间平滑，防止权重噪声。",
            "先使用确定性规则，再考虑轻量卷积模块。",
            "与图像级路由比较，确认复杂度增加是否值得。",
        ],
    )
    add_paragraph(
        doc,
        "像素级路由不是第一优先级。只有图像级路由和校准通过后才进入本工作包，避免同时引入过多变量。"
    )

    heading(doc, "WP8：可视化与解释工具", 2)
    add_bullets(
        doc,
        [
            "原图、真值掩码、视觉异常图、文本异常图和融合异常图并排显示。",
            "显示图像级视觉权重、文本权重和最终路由决定。",
            "显示像素级权重图及高权重区域。",
            "建立“融合修正单分支错误”和“融合反而变差”两类案例库。",
            "按类别统计平均视觉权重、冲突率和路由比例。",
        ],
    )

    heading(doc, "七、测试与验证计划")
    heading(doc, "7.1 单元测试", 2)
    add_table(
        doc,
        ["测试类型", "检查内容", "通过标准"],
        [
            ["形状", "N、H、W 与权重尺寸", "输出与输入一一对应"],
            ["数值", "0、1、极端值、NaN、Inf", "有限值稳定；非法值报错"],
            ["对齐", "乱序、缺失、重复 sample_id", "不能静默错配"],
            ["边界", "温度、最小权重、空类别", "参数非法时明确报错"],
            ["回归", "原统一评测测试", "现有指标不受影响"],
        ],
        [1450, 4100, 3810],
    )
    heading(doc, "7.2 冻结缓存烟雾测试", 2)
    add_bullets(
        doc,
        [
            "仅使用 VisA seed 0 的冻结缓存。",
            "先单类别，再扩展到 12 类。",
            "先检查样本对齐和数值，再看开发集指标。",
            "开发指标只用于筛选结构，不作为论文最终结果。",
            "每次运行保存配置、commit、输入缓存哈希和输出报告。",
        ],
    )
    heading(doc, "7.3 设计冻结后的最终验证", 2)
    add_bullets(
        doc,
        [
            "冻结校准方法、特征、路由结构、阈值和所有超参数。",
            "只运行一次 VisA seed 1/2 最终验证。",
            "再运行 MVTec seed 0/1/2 跨数据集验证。",
            "任何最终结果不理想时，先报告，不回到测试集调参。",
            "如确需改变设计，必须声明新版本并重新冻结。",
        ],
    )

    heading(doc, "八、消融实验设计")
    add_paragraph(doc, "消融实验用于回答“性能提升到底来自哪里”。至少包含以下比较：")
    add_numbers(
        doc,
        [
            "视觉分支单独使用。",
            "文本分支单独使用。",
            "固定 0.5 融合。",
            "预先规定的 0.7/0.3 与 0.3/0.7 融合。",
            "只使用不确定性的图像级路由。",
            "不确定性加分支一致性。",
            "加入分数校准前后对比。",
            "图像级动态路由与像素级动态路由对比。",
            "去掉像素权重平滑。",
            "规则路由与轻量学习路由对比。",
        ],
    )
    add_paragraph(
        doc,
        "结果要同时报告图像指标、像素指标、不同 seed 的均值和标准差，以及额外运行时间。"
    )

    heading(doc, "九、并行执行顺序与里程碑")
    add_table(
        doc,
        ["里程碑", "设计轨工作", "可与哪些训练并行", "完成标志"],
        [
            ["M1", "接口、配置、合成测试", "PromptAD seed 1/2", "测试全通过"],
            ["M2", "样本对齐与缓存适配", "PromptAD 完整矩阵", "VisA seed 0 十二类可对齐"],
            ["M3", "校准方案与固定融合", "ReMP-AD 环境/Gate A", "校准不依赖测试统计"],
            ["M4", "图像级规则路由", "AdaptCLIP 权重与 Gate A", "超过或解释固定融合"],
            ["M5", "像素级路由候选", "MVTec 矩阵补齐", "权重图稳定且可解释"],
            ["M6", "冻结设计与验证脚本", "第一阶段收尾", "seed 1/2 与 MVTec 尚未用于调参"],
        ],
        [900, 2600, 2600, 3260],
    )
    heading(doc, "建议的日常工作方式", 2)
    add_bullets(
        doc,
        [
            "GPU 忙时：做接口、测试、统计、文档和 CPU 缓存评测。",
            "GPU 空闲时：优先让基线轨继续，不让融合试验抢占长时间训练资源。",
            "每完成一个基线配置：登记结果并启动下一个配置。",
            "每完成一个设计工作包：运行单元测试和冻结缓存烟雾测试。",
            "每周形成一次简短检查点：完成项、失败项、设计决定和下一步。",
        ],
    )

    heading(doc, "十、风险与处理办法")
    add_table(
        doc,
        ["风险", "表现", "处理办法"],
        [
            ["分数尺度不一致", "全部路由到一个分支", "先校准，再计算不确定性"],
            ["测试数据泄漏", "最终结果过高但不可复现", "API禁用真值；最终集锁定"],
            ["样本错配", "指标异常或热图对应错误", "强制 sample_id 对齐"],
            ["路由过拟合", "seed 0好、seed 1/2差", "简化模型并严格冻结"],
            ["像素权重噪声", "权重图碎片化", "空间平滑和复杂度约束"],
            ["基线未齐", "无法公平比较", "双轨推进但最终结论等待基线完成"],
            ["GPU资源不足", "OOM或训练过慢", "设计轨CPU优先；基线串行"],
        ],
        [1700, 3300, 4360],
    )

    heading(doc, "十一、第一阶段训练结束前的交付清单")
    add_bullets(
        doc,
        [
            "动态融合统一输入输出接口。",
            "动态融合配置文件和禁止信息规则。",
            "视觉/文本缓存对齐工具及对齐报告。",
            "分数校准候选及其无泄漏说明。",
            "固定融合基线脚本。",
            "图像级规则路由器。",
            "像素级路由原型或明确的延后结论。",
            "合成数据与缓存单元测试。",
            "权重、路由决定和失败案例可视化工具。",
            "VisA seed 0 开发结果，不包含 seed 1/2 与 MVTec 最终结论。",
            "冻结设计清单和最终验证一键脚本。",
            "动态融合设计说明、实验记录和复现命令。",
        ],
    )
    add_callout(
        doc,
        "第一阶段结束前的理想状态",
        "基线矩阵基本完成；动态融合代码、校准、规则路由、测试和分析工具全部就绪。"
        "此时只需要冻结设计并运行保留的 seed 1/2 与 MVTec，即可获得可信的第二阶段结论。",
        fill=LIGHT_BLUE,
    )

    heading(doc, "十二、近期立即执行的任务")
    add_numbers(
        doc,
        [
            "继续后台运行 PromptAD seed 1/2，不改变现有协议。",
            "为 AnomalyCLIP VisA 缓存补充可复核 sample_id，作为更纯的文本引导分支。",
            "实现仅使用正常参考图的校准接口。",
            "输出视觉、文本和融合三套固定基线结果。",
            "实现图像级规则路由，并保存每张图的决定和权重。",
            "在 VisA seed 0 十二类上做冲突与失败案例分析。",
            "确定是否值得进入像素级动态路由。",
            "形成设计冻结清单，但暂不运行最终保留集。",
        ],
    )

    heading(doc, "十三、完成判定")
    add_paragraph(doc, "在第一阶段训练结束前，设计轨达到以下条件即可认为“准备完成”：")
    add_bullets(
        doc,
        [
            "接口不接收测试真值，配置和代码均可审计。",
            "两分支缓存可以稳定对齐并生成统一融合 NPZ。",
            "校准来源合法，且解决明显的单分支路由偏置。",
            "固定融合和图像级动态融合均可重复运行。",
            "10项以上单元/回归测试持续通过。",
            "所有开发决定只基于合成数据和 VisA seed 0。",
            "最终验证脚本已经写好，但 VisA seed 1/2 与 MVTec 尚未用于调参。",
            "设计局限、失败案例和下一步像素级方案均有书面记录。",
        ],
    )
    add_paragraph(
        doc,
        "满足以上条件后，第二阶段就不再处于“边训练边改结构”的状态，而是进入“冻结设计、统一验证、论文分析”的阶段。"
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()

from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).resolve().parent / "DCFnet_English_Draft_Through_Method_Revised_20260903.docx"
FIGURE_1 = (
    Path(__file__).resolve().parent
    / "figures_20260830"
    / "DCFnet_Fig1_AI_reference_20260903.png"
)


TITLE = "DCFnet: Few-Shot Industrial Anomaly Localization via a Dual-Encoder Calibrated Fusion Network"


ABSTRACT = (
    "Automated visual inspection must identify defects whose appearance, scale, and location cannot be exhaustively "
    "enumerated before deployment. The problem is particularly acute during production-line start-up, when only a few "
    "defect-free images of a new product may be available. Existing few-shot industrial anomaly detectors increasingly "
    "use strong pretrained representations, yet a single feature space may not respond uniformly to fine textures, object "
    "geometry, and semantically meaningful structural changes. This paper presents DCFnet, a dual-encoder calibrated "
    "fusion architecture for normal-only few-shot industrial anomaly detection and pixel-level localization. DCFnet uses "
    "a frozen DINOv2 ViT-B/14 encoder and a frozen CLIP ViT-L/14@336 visual encoder configured with AnomalyCLIP's DPAM "
    "attention modification. The heterogeneous patch "
    "grids are spatially aligned, independently L2-normalized, combined by a fixed equal-weight concatenation, and "
    "normalized again before a category-specific normal memory is constructed. Query patches are scored by their nearest-"
    "neighbor distance to this memory. The method has no trainable fusion parameters and uses neither text embeddings nor "
    "abnormal target-domain samples. Under a matched protocol covering MPDD, BTAD, VisA, and MVTec AD, three "
    "reference-image seeds, and 1-, 2-, and 4-shot settings, DCFnet increases mean pixel-level average precision over an "
    "otherwise identical DINO-only control by 0.0258, 0.0249, 0.0524, and 0.0320, respectively. The gain is positive in "
    "all 36 dataset-level seed/shot configurations, although it is not universal across categories or image-level metrics. "
    "These findings establish a transparent and reproducible reference for the benefits, limits, and computational cost "
    "of fixed heterogeneous visual-feature fusion in data-scarce industrial inspection."
)


INTRO_PARAS = [
    (
        "Visual quality inspection is a core component of modern manufacturing because small surface defects, geometric "
        "deformations, missing components, and contamination can compromise product reliability even when the affected "
        "region occupies only a small fraction of an image. Conventional supervised detection assumes that representative "
        "defect categories and sufficiently detailed annotations are available during training. That assumption is often "
        "untenable in industrial production: defective items are deliberately rare, new defect modes emerge after process "
        "changes, and pixel-wise labeling requires specialized labor. Industrial anomaly detection therefore learns a model "
        "of normal appearance and identifies departures from that model at test time. MVTec AD, VisA, MPDD, and BTAD have "
        "made this setting measurable across textured surfaces, manufactured objects, metal components, and complex real-"
        "world parts [1]-[4]."
    ),
    (
        "A more stringent and practically relevant setting arises when even normal data are scarce. At the start of a new "
        "production run, only one or a few verified normal images may be available, while collecting a large category-specific "
        "training set would delay deployment. Normal-only few-shot anomaly detection addresses this cold-start condition by "
        "constructing a reference model from K normal images and detecting previously unseen abnormalities without using "
        "abnormal target-domain samples. Patch-based memory methods are attractive in this regime because they require no "
        "closed defect taxonomy: local query descriptors are compared with descriptors extracted from normal references, and "
        "their deviations directly support pixel-level localization. PatchCore demonstrated the effectiveness of nearest-"
        "neighbor patch memories [6], while DINOv2 provides robust self-supervised visual descriptors that can be transferred "
        "without category-specific representation training [9]. AnomalyDINO subsequently showed that frozen DINOv2 patches "
        "form a strong basis for few-shot industrial anomaly detection [19]."
    ),
    (
        "Nevertheless, no single representation is guaranteed to be uniformly informative across the diversity of industrial "
        "defects. Texture irregularities, thin scratches, large structural changes, pose variation, and repeated components "
        "place different demands on local descriptors. A representation optimized by self-supervised visual objectives may "
        "preserve spatial correspondence and fine structure, whereas an image encoder inherited from image-text contrastive "
        "pretraining may encode broader semantic regularities. CLIP established the transferability of large-scale image-text "
        "pretraining [7], and WinCLIP adapted CLIP to zero- and few-shot anomaly classification and segmentation [8]. More "
        "recently, AnomalyCLIP learned object-agnostic normal and abnormal prompts, and PromptAD learned prompts using only "
        "normal target samples [10], [11]. These methods demonstrate the value of CLIP-derived representations, but they do "
        "not answer whether the visual patch descriptors of two frozen encoders provide complementary evidence when the "
        "downstream detector is held fixed and text is removed from inference."
    ),
    (
        "The question is nontrivial because combining two encoders is not itself new. Sea-CLIP integrates CLIP and DINOv2 "
        "representations through trainable prompts, synthetic anomalies, and an anomaly-matching decoder [27]. Other recent "
        "approaches have introduced prompt learning, multimodal attention, adaptive pooling, query-conditioned prototype "
        "refinement, and frequency-domain branches. Their capacity can improve performance, but it also makes attribution "
        "difficult: an observed gain may arise from auxiliary training data, learned fusion parameters, synthetic defects, or "
        "test-time adaptation rather than from the intrinsic complementarity of the underlying frozen image features. A "
        "controlled and training-free study is therefore needed to isolate the contribution of the second visual representation."
    ),
    (
        "Three technical difficulties must be addressed. First, heterogeneous encoders produce patch grids at different input "
        "resolutions, so spatially corresponding descriptors cannot be combined without explicit alignment. Second, the feature "
        "spaces have different scales and pretraining biases; uncalibrated concatenation can allow one branch to dominate the "
        "distance metric. Third, average benchmark scores may conceal reference-sampling instability, category-specific negative "
        "transfer, or additional inference cost. A credible solution should therefore use a clearly specified alignment and "
        "normalization procedure, freeze all fusion decisions before external evaluation, compare against a genuinely matched "
        "single-encoder control, and report both successful and adverse cases."
    ),
    (
        "To this end, we present DCFnet, a Dual-Encoder Calibrated Fusion Network for normal-only few-shot industrial anomaly "
        "localization. The term calibrated refers to deterministic spatial and magnitude calibration rather than to a learned "
        "gating network. DCFnet extracts patch features from frozen DINOv2 ViT-B/14 and a frozen CLIP ViT-L/14@336 visual "
        "encoder configured with AnomalyCLIP's DPAM attention modification. The CLIP-derived patch grid is bilinearly resized "
        "to the DINOv2 grid; each branch is L2-normalized, the "
        "two branches are concatenated with fixed equal scaling, and the fused descriptors are normalized again. A category-"
        "specific memory is then built exclusively from K normal reference images, and anomaly scores are obtained by one-"
        "nearest-neighbor search. No text embedding, abnormal target-domain sample, target-domain fine-tuning, or test-label-derived "
        "decision rule is used."
    ),
    (
        "The frozen method is evaluated on four industrial datasets under 1-, 2-, and 4-shot protocols with three reference-"
        "sampling seeds. Relative to a matched DINO-only pipeline that shares the same reference identities, memory construction, "
        "nearest-neighbor scoring, post-processing, and evaluator, DCFnet yields positive mean Pixel-AP gains in every one of "
        "the 36 dataset-level seed/shot configurations. The dataset-level mean gains are 0.0258 on MPDD, 0.0249 on BTAD, "
        "0.0524 on VisA, and 0.0320 on MVTec AD. The evidence is deliberately qualified: some categories exhibit persistent "
        "negative transfer, BTAD does not improve on every image-level metric, and the second encoder introduces non-negligible "
        "latency. Thus, the paper does not claim universal superiority or state-of-the-art performance."
    ),
]


CONTRIBUTIONS = [
    "A deterministic dual-encoder calibration and fusion pipeline that aligns heterogeneous patch grids, equalizes branch magnitudes, and combines frozen DINOv2 and DPAM-modified CLIP visual features without a learned fusion module, text input, or target-domain parameter training.",
    "A matched experimental design that isolates the effect of the second visual representation while keeping normal references, nearest-neighbor scoring, post-processing, and evaluation unchanged across four industrial datasets, three reference seeds, and three shot levels.",
    "A reproducibility- and boundary-oriented evaluation protocol that reports complete image- and pixel-level evidence, category-specific negative transfer, computational cost, leakage checks, and reconstructable artifacts instead of presenting only favorable aggregate scores.",
]


RELATED = [
    (
        "Industrial anomaly detection methods can be organized into three broad families according to how normality is "
        "represented: reconstruction or one-class learning, pretrained feature embedding with normal memories, and foundation-"
        "model-based or hybrid approaches. The families are not mutually exclusive, but the classification clarifies their "
        "principal trade-offs. Reconstruction methods learn an explicit normal-image generator and can produce intuitive residual "
        "maps, yet they require category-specific optimization and may reconstruct abnormal regions too well. Embedding and memory "
        "methods are simpler at inference and often localize defects accurately, but their effectiveness depends strongly on the "
        "chosen representation and reference coverage. Foundation-model and hybrid methods improve transfer or semantic breadth, "
        "although they may introduce prompts, auxiliary data, learned adapters, or protocol-dependent complexity."
    ),
]


RW_RECON = [
    (
        "Early visual-inspection systems relied on hand-engineered texture, edge, morphology, and statistical descriptors, "
        "followed by conventional classifiers or thresholding. Such pipelines can be efficient when imaging conditions and defect "
        "definitions are stable, but their task-specific features generalize poorly to new materials, illumination conditions, or "
        "previously unseen defect morphologies. Deep reconstruction methods replaced manual descriptors with autoencoders, "
        "generative models, and student-teacher networks trained on normal images. Their central assumption is that a model fitted "
        "to normality reconstructs or predicts normal regions accurately while producing larger residuals on anomalous content. "
        "This assumption provides a natural anomaly map, but high-capacity networks may also reconstruct defects, and separate "
        "optimization is usually required for each target category."
    ),
    (
        "Recent journal studies continue to refine this family rather than treating it as obsolete. SLSG improves feature embeddings "
        "and one-class classification through generative pretraining, anomaly simulation, and graph reasoning [17]. RealNet uses "
        "realistic synthetic anomalies and feature selection to learn defect-sensitive representations [12]. DFD decomposes images "
        "into frequency components and trains dual discriminators to expose subtle abnormalities that may be weak in the spatial "
        "domain [15]. DMMGNet combines feature discrimination mapping, augmentation, and memory-bank mean guidance for few-shot "
        "training [16]. These approaches strengthen category-specific discrimination, but their learned components and synthetic "
        "anomaly assumptions differ from the zero-training normal-reference setting considered here. Their common advantage is "
        "adaptability to specific defect statistics; their common limitation is that retraining or optimization complicates rapid "
        "deployment when a new product has only a handful of verified normal samples."
    ),
]


RW_MEMORY = [
    (
        "Embedding-based methods avoid image reconstruction and instead measure deviations in a pretrained feature space. PaDiM "
        "models the distribution of normal patch embeddings with multivariate Gaussian statistics [5]. PatchCore stores a coreset "
        "of normal local features and scores query patches by nearest-neighbor distance, demonstrating that a simple memory can "
        "achieve strong localization [6]. This paradigm is especially compatible with normal-only inspection because it neither "
        "requires defect labels nor assumes a closed anomaly taxonomy. Its weakness is equally direct: when normal references are "
        "few, the memory may inadequately cover acceptable appearance variation, and the distance metric is only as informative as "
        "the underlying feature representation."
    ),
    (
        "Few-shot work has consequently focused on expanding normal coverage or improving representation quality. FEAD transforms "
        "support features conditionally and constructs multi-frequency descriptors to compensate for sparse normal patterns [24]. "
        "The K-NG framework formulates few-shot online detection and updates a Neural Gas representation from an unlabeled stream "
        "while attempting to avoid contamination by anomalies [18]. PGAD combines global invariant and multiscale local features "
        "and explicitly calibrates differently distributed anomaly scores [31]. FastRef refines normal prototypes at test time and "
        "uses optimal transport to reduce the risk of absorbing anomalous query features [29]. These methods address reference "
        "scarcity through transformation, updating, multiscale modeling, or transductive refinement. DCFnet instead keeps the normal "
        "memory fixed and asks whether a second frozen representation provides a reproducible gain without query-conditioned adaptation."
    ),
]


RW_FOUNDATION = [
    (
        "Self-supervised visual foundation models have substantially changed the quality of transferable local descriptors. DINOv2 "
        "learns robust visual features without labels and transfers well across recognition and dense-prediction tasks [9]. "
        "AnomalyDINO applies frozen DINOv2 patches to few-shot anomaly detection and shows that straightforward patch comparison is "
        "already highly competitive [19]. UniVAD develops a training-free unified model that reasons over reference images across "
        "industrial, logical, and medical anomalies [20]. SubspaceAD fits a principal subspace to frozen DINOv2 patches and scores "
        "reconstruction residuals, providing a strong training-free alternative to direct nearest-neighbor memories [28]. DCP-SFR "
        "argues that deep structural representations may suppress shallow defect cues and refines them to preserve subtle anomalies "
        "[32]. Collectively, these studies show that feature selection, normal-space modeling, and cue preservation can be as "
        "important as the downstream anomaly scorer. They also establish a demanding visual-only baseline against which fusion must "
        "be assessed."
    ),
    (
        "Vision-language pretraining offers another route to transfer. CLIP learns aligned image and text embeddings from large-scale "
        "image-text supervision [7]. WinCLIP uses prompt ensembles, window-level visual features, and optional normal references for "
        "zero- and few-shot anomaly classification and segmentation [8]. AnomalyCLIP replaces object-specific descriptions with "
        "object-agnostic prompts learned for normality and abnormality [10], whereas PromptAD learns prompts using only normal target "
        "samples [11]. InCTRL performs in-context residual learning with few-shot sample prompts [13]. AA-CLIP introduces anomaly-"
        "aware text anchors and patch-level visual alignment [22], and FAPrompt decomposes coarse abnormality descriptions into fine-"
        "grained prompt components [23]. These methods improve semantic transfer, but most retain text-conditioned scoring or learned "
        "prompt parameters. DCFnet uses only the frozen DPAM-modified CLIP image encoder at inference; it is therefore a dual-visual-encoder "
        "method rather than a vision-text inference system."
    ),
]


RW_FUSION = [
    (
        "Recent hybrid studies are the closest methodological context. ReMP-AD combines retrieval-enhanced normal references with "
        "multimodal prompt fusion [21]. A multimodal cooperative-refinement framework jointly refines text prompts, spatial features, "
        "frequency features, and decision-level fusion [25]. Sea-CLIP combines CLIP and DINOv2 with trainable prompts, synthetic "
        "anomalies, and a learned matching decoder [27]. PAPL integrates CLIP and DINO-family hierarchical features through prompt "
        "distribution learning and adaptive pooling [30], while a CLIP-DINOv2 framework uses trained dual-modality attention and "
        "stabilized pooling for zero-shot transfer [26]. Real-IAD further shows why such methods must be tested beyond small or "
        "potentially saturated benchmarks: it introduces a substantially larger, multiview industrial dataset with broader real-"
        "world variation [14]."
    ),
    (
        "The literature therefore establishes both the promise and the ambiguity of feature fusion. Learned hybrid systems can "
        "achieve strong results, but they conflate encoder complementarity with the effects of prompt optimization, synthetic data, "
        "attention modules, or test-time adaptation. Conversely, a fixed combination is easy to audit but can suffer from spatial "
        "misalignment, incompatible feature magnitudes, and negative transfer. The present study occupies this unresolved middle "
        "ground. It does not claim the first use of CLIP- and DINO-family representations. Instead, it introduces a deterministic "
        "spatial-and-magnitude calibration pipeline and isolates the effect of adding the second frozen image representation under "
        "an otherwise matched normal-memory detector. This position also motivates reporting reference-seed stability, category-level "
        "failures, and computational cost as primary evidence rather than as supplementary observations."
    ),
]


METHOD_PROBLEM = [
    (
        "Let c denote an industrial product category and let R_c = {x_1^c, ..., x_K^c} be a support set containing K verified normal "
        "images, where K is in {1, 2, 4}. No abnormal target-domain image is available for fitting the detector. Given a query "
        "image x, the objective is to estimate a nonnegative evaluation map S_448(x) of size 448 x 448 and a nonnegative image-level "
        "anomaly score s_img(x). Larger values should indicate a greater departure "
        "from the normal appearance represented "
        "by R_c. The method addresses binary normal-versus-anomalous inspection and spatial localization; it does not predict a "
        "closed defect label such as scratch, crack, hole, or stain. Pixel masks and image labels are used only for evaluation and "
        "never for feature fusion, memory construction, threshold selection, or reference sampling."
    ),
    (
        "For a query pixel u, let y_u in {0, 1} denote its evaluation-only ground-truth status. The desired scoring function ranks "
        "anomalous pixels above normal pixels as often as possible, while remaining applicable to "
        "previously unseen defect types. Because no exhaustive anomaly distribution can be estimated from R_c, DCFnet models normality "
        "rather than learning an explicit multiclass decision boundary. The same category-specific support set is used by DCFnet and "
        "the matched DINO-only control, so their difference can be attributed to the additional frozen representation and its "
        "calibration path."
    ),
]


METHOD_OVERVIEW = [
    (
        "DCFnet contains four deterministic stages: frozen dual-encoder extraction, cross-encoder spatial calibration, branch-wise "
        "magnitude calibration and fixed fusion, and nearest-neighbor normal-memory scoring. The support and query paths share the "
        "same feature extractor but serve different roles. Support images populate the category-specific memory once; a query image "
        "is never inserted into that memory. Each query patch is compared with the fixed support memory, and the resulting patch "
        "scores are bilinearly upsampled and then Gaussian-smoothed to obtain the pixel-level anomaly map. Figure 1 specifies this "
        "separation explicitly."
    ),
]


METHOD_ENCODERS = [
    (
        "The first branch is a frozen DINOv2 ViT-B/14 encoder [9]. An input is resized according to the fixed pipeline with a target "
        "short-side resolution of 448 pixels and cropped so that each spatial dimension is divisible by the patch size of 14. The "
        "resulting grid is denoted by h_D x w_D; its shorter dimension is 32 patches, and each patch descriptor has dimension "
        "d_D = 768. "
        "DINOv2 is selected because its self-supervised objective produces transferable dense visual descriptors and because DINO-based "
        "patch comparison is a strong few-shot baseline [19]. The branch is not fine-tuned on any of the four target datasets."
    ),
    (
        "The second branch is a frozen CLIP ViT-L/14@336 visual encoder configured with AnomalyCLIP's diagonally prominent attention "
        "map (DPAM) modification [10]. It receives a 518 x 518 input and produces a 37 x 37 grid of 768-dimensional local descriptors. "
        "The extractor requests the outputs of visual layers 6, 12, 18, and 24 with DPAM_layer set to 20 and retains the final returned "
        "patch tensor, corresponding to layer 24. Only image features are retained. No text prompt, text-encoder output, normal-versus-"
        "abnormal language similarity, or dynamic vision-language routing is computed during DCFnet inference. The branch is used to "
        "test whether CLIP-pretrained local descriptors processed by a fixed DPAM attention path provide information complementary to "
        "self-supervised DINOv2 descriptors. This motivation does not assume that complementarity holds for every category; that "
        "question is evaluated rather than built into the method."
    ),
]


METHOD_FUSION = [
    (
        "Let F_D(x) and F_C(x) denote the 768-dimensional DINOv2 and DPAM-modified CLIP patch tensors on grids of size "
        "h_D x w_D and h_C x w_C, respectively. Direct concatenation is invalid because the two grid sizes differ. We therefore "
        "apply a bilinear grid-resizing operator R(·) to F_C(x), with align_corners set to false, producing F_C-to-D(x) = R(F_C(x)) "
        "on the h_D x w_D grid. This spatial "
        "calibration preserves the DINO grid as the reference lattice and avoids discrete nearest-cell assignment when the two patch "
        "centers do not coincide."
    ),
    (
        "Spatial alignment does not remove differences in feature magnitude. For every patch p, each branch is independently L2-"
        "normalized before the calibrated branches are concatenated with fixed alpha = 0.5 and normalized again:"
    ),
]


NORMALIZATION_EQUATION = "normalization"


FUSION_EQUATION = "fusion"


METHOD_FUSION_AFTER = [
    (
        "The equal coefficients are frozen protocol constants, not learned gates. Because the complete concatenated vector is "
        "renormalized, the construction should be interpreted as equal branch scaling before a joint distance computation, rather "
        "than as a probabilistic average of two anomaly scores. The operation yields a 1536-dimensional descriptor at every DINO-grid "
        "location. In this manuscript, calibrated fusion therefore has a precise and limited meaning: resolution alignment plus "
        "branch-wise magnitude normalization before fixed feature concatenation."
    ),
]


METHOD_MEMORY = [
    (
        "For category c, the fused normal memory is the union of all calibrated patch descriptors extracted from its K support images:"
    ),
]


MEMORY_EQUATION = "memory"


METHOD_MEMORY_2 = [
    (
        "The descriptors are stored in an exact FAISS IndexFlatL2 index [33]. For query patch p, the squared Euclidean distance to "
        "its nearest normal descriptor defines the raw anomaly score:"
    ),
]


SCORE_EQUATION = "score"


METHOD_MEMORY_3 = [
    (
        "All fused descriptors are unit-normalized; consequently, one half of their squared Euclidean distance equals cosine distance. "
        "The factor 1/2 therefore preserves a convenient distance interpretation without changing score ordering. The patch scores are "
        "arranged on the DINO grid, bilinearly resized to 448 x 448, and then filtered with a Gaussian kernel of sigma = 4 to produce "
        "S_448(x). The image-level score is obtained by maximum pooling, s_img(x) = max_u S_448,u(x), so that a strongly anomalous local region can "
        "trigger image-level inspection. This aggregation is fixed for DCFnet and the matched control."
    ),
]


METHOD_CONTROL = [
    (
        "The matched DINO-only control removes the DPAM-modified CLIP branch and the concatenation step but retains the identical "
        "DINOv2 descriptors, support-image identities, K values, random seeds, exact one-nearest-neighbor search, distance convention, "
        "upsampling, Gaussian smoothing, and evaluation resolution. This design is essential: comparing DCFnet only with published "
        "numbers from methods using different checkpoints, data roles, or post-processing would not isolate the effect of feature "
        "fusion. The matched control instead changes a single factor: whether the calibrated second visual representation is present."
    ),
    (
        "DCFnet introduces zero trainable parameters in the target-domain pipeline. This property eliminates optimization time but does "
        "not make inference free: both frozen encoders must execute, and the fused memory has twice the descriptor dimensionality of the "
        "DINO-only memory. If P = h_D w_D patches are retained per reference image, the category memory contains approximately KP "
        "vectors of dimension 1536, while exact search scales linearly with the stored memory size. Accordingly, the paper treats "
        "training cost, inference latency, and memory consumption as distinct quantities and does not claim zero computational overhead."
    ),
    (
        "Finally, the pipeline follows a strict leakage contract. Test labels, pixel masks, test-set statistics, and category-specific "
        "test outcomes are unavailable when defining alpha, the distance rule, or post-processing. MPDD is used for development; the "
        "frozen configuration is then evaluated on BTAD and MVTec AD as external validation domains. VisA is reported as a frozen "
        "in-domain validation set because of the lineage of the AnomalyCLIP checkpoint and is not described as independent external "
        "generalization. This explicit data-role distinction prevents the pretrained second branch from being presented as if it had no "
        "relationship to every evaluation domain."
    ),
]


REFERENCES = [
    "[1] P. Bergmann, M. Fauser, D. Sattlegger, and C. Steger, 'MVTec AD - A comprehensive real-world dataset for unsupervised anomaly detection,' in Proc. IEEE/CVF Conf. Comput. Vis. Pattern Recognit. (CVPR), 2019, pp. 9592-9600.",
    "[2] Y. Zou, J. Jeong, L. Pemula, D. Zhang, and O. Dabeer, 'SPot-the-Difference self-supervised pre-training for anomaly detection and segmentation,' in Proc. Eur. Conf. Comput. Vis. (ECCV), 2022, pp. 392-408.",
    "[3] S. Jezek, M. Jonak, R. Burget, P. Dvorak, and M. Skotak, 'Deep learning-based defect detection of metal parts: Evaluating current methods in complex conditions,' in Proc. 13th Int. Congr. Ultra Modern Telecommun. Control Syst. Workshops, 2021, pp. 66-71, doi: 10.1109/ICUMT54235.2021.9631567.",
    "[4] P. Mishra, R. Verk, D. Fornasier, C. Piciarelli, and G. L. Foresti, 'VT-ADL: A vision transformer network for image anomaly detection and localization,' in Proc. IEEE 30th Int. Symp. Ind. Electron., 2021, pp. 1-6, doi: 10.1109/ISIE45552.2021.9576231.",
    "[5] T. Defard, A. Setkov, A. Loesch, and R. Audigier, 'PaDiM: A patch distribution modeling framework for anomaly detection and localization,' in Pattern Recognition. ICPR International Workshops and Challenges, 2021, pp. 475-489, doi: 10.1007/978-3-030-68799-1_35.",
    "[6] K. Roth, L. Pemula, J. Zepeda, B. Scholkopf, T. Brox, and P. Gehler, 'Towards total recall in industrial anomaly detection,' in Proc. IEEE/CVF Conf. Comput. Vis. Pattern Recognit. (CVPR), 2022, pp. 14318-14328.",
    "[7] A. Radford et al., 'Learning transferable visual models from natural language supervision,' in Proc. 38th Int. Conf. Mach. Learn., vol. 139, 2021, pp. 8748-8763.",
    "[8] J. Jeong, Y. Zou, T. Kim, D. Zhang, A. Ravichandran, and O. Dabeer, 'WinCLIP: Zero-/few-shot anomaly classification and segmentation,' in Proc. IEEE/CVF Conf. Comput. Vis. Pattern Recognit. (CVPR), 2023, pp. 19606-19616.",
    "[9] M. Oquab et al., 'DINOv2: Learning robust visual features without supervision,' Trans. Mach. Learn. Res., 2024.",
    "[10] Q. Zhou, G. Pang, Y. Tian, S. He, and J. Chen, 'AnomalyCLIP: Object-agnostic prompt learning for zero-shot anomaly detection,' in Proc. Int. Conf. Learn. Represent. (ICLR), 2024.",
    "[11] X. Li, Z. Zhang, X. Tan, C. Chen, Y. Qu, Y. Xie, and L. Ma, 'PromptAD: Learning prompts with only normal samples for few-shot anomaly detection,' in Proc. IEEE/CVF Conf. Comput. Vis. Pattern Recognit. (CVPR), 2024, pp. 16838-16848.",
    "[12] X. Zhang, M. Xu, and X. Zhou, 'RealNet: A feature selection network with realistic synthetic anomaly for anomaly detection,' in Proc. IEEE/CVF Conf. Comput. Vis. Pattern Recognit. (CVPR), 2024, pp. 16699-16708.",
    "[13] J. Zhu and G. Pang, 'Toward generalist anomaly detection via in-context residual learning with few-shot sample prompts,' in Proc. IEEE/CVF Conf. Comput. Vis. Pattern Recognit. (CVPR), 2024, pp. 17826-17836.",
    "[14] C. Wang et al., 'Real-IAD: A real-world multi-view dataset for benchmarking versatile industrial anomaly detection,' in Proc. IEEE/CVF Conf. Comput. Vis. Pattern Recognit. (CVPR), 2024, pp. 22883-22892.",
    "[15] Y. Bai, J. Zhang, Z. Chen, Y. Dong, Y. Cao, and G. Tian, 'Dual-path frequency discriminators for few-shot anomaly detection,' Knowl.-Based Syst., vol. 302, Art. no. 112397, 2024, doi: 10.1016/j.knosys.2024.112397.",
    "[16] A. Luo, G. Wen, Y. Cheng, S. Mei, H. Dong, and X. Liu, 'DMMGNet: A discrimination mapping and memory bank mean guidance-based network for high-performance few-shot industrial anomaly detection,' Neurocomputing, vol. 610, Art. no. 128622, 2024, doi: 10.1016/j.neucom.2024.128622.",
    "[17] M. Yang, J. Liu, Z. Yang, and Z. Wu, 'SLSG: Industrial image anomaly detection with improved feature embeddings and one-class classification,' Pattern Recognit., vol. 156, Art. no. 110862, 2024, doi: 10.1016/j.patcog.2024.110862.",
    "[18] S. Wei, X. Wei, Z. Ma, S. Dong, S. Zhang, and Y. Gong, 'Few-shot online anomaly detection and segmentation,' Knowl.-Based Syst., vol. 300, Art. no. 112168, 2024, doi: 10.1016/j.knosys.2024.112168.",
    "[19] S. Damm, M. Laszkiewicz, J. Lederer, and A. Fischer, 'AnomalyDINO: Boosting patch-based few-shot anomaly detection with DINOv2,' in Proc. IEEE/CVF Winter Conf. Appl. Comput. Vis. (WACV), 2025, pp. 1319-1329.",
    "[20] Z. Gu, B. Zhu, G. Zhu, Y. Chen, M. Tang, and J. Wang, 'UniVAD: A training-free unified model for few-shot visual anomaly detection,' in Proc. IEEE/CVF Conf. Comput. Vis. Pattern Recognit. (CVPR), 2025, pp. 15194-15203.",
    "[21] H. Ma, G. Yang, D. Zhao, Y. Ji, and W. Zuo, 'ReMP-AD: Retrieval-enhanced multi-modal prompt fusion for few-shot industrial visual anomaly detection,' in Proc. IEEE/CVF Int. Conf. Comput. Vis. (ICCV), 2025, pp. 20425-20434, doi: 10.1109/ICCV51701.2025.01899.",
    "[22] W. Ma et al., 'AA-CLIP: Enhancing zero-shot anomaly detection via anomaly-aware CLIP,' in Proc. IEEE/CVF Conf. Comput. Vis. Pattern Recognit. (CVPR), 2025, pp. 4744-4754.",
    "[23] J. Zhu, Y.-S. Ong, C. Shen, and G. Pang, 'Fine-grained abnormality prompt learning for zero-shot anomaly detection,' in Proc. IEEE/CVF Int. Conf. Comput. Vis. (ICCV), 2025, pp. 22241-22251.",
    "[24] Z. Hu, X. Zeng, Y. Li, Z. Yin, E. Meng, L. Zhu, and X. Kong, 'Few-shot anomaly detection with adaptive feature transformation and descriptor construction,' Chin. J. Aeronaut., vol. 38, no. 3, Art. no. 103098, 2025, doi: 10.1016/j.cja.2024.06.007.",
    "[25] L. Xu, D. Han, G. Li, M. Zhou, J. Wan, and M. Li, 'Multimodal feature cooperative refinement for few-shot anomaly detection,' Adv. Eng. Inform., vol. 68, pt. C, Art. no. 103792, 2025, doi: 10.1016/j.aei.2025.103792.",
    "[26] J. Jiang, Z. He, A. Wan, K. AL-Bukhaiti, and K. Wang, 'Zero-shot industrial anomaly detection via CLIP-DINOv2 multimodal fusion and stabilized attention pooling,' Electronics, vol. 14, no. 24, Art. no. 4785, 2025, doi: 10.3390/electronics14244785.",
    "[27] X. Guo, Z. Chen, C. D. Castillo, H. Wang, and X. Liu, 'Sea-CLIP: Mining semantic-aware representations for few-shot anomaly detection with CLIP,' in Proc. IEEE/CVF Winter Conf. Appl. Comput. Vis. (WACV), 2026, pp. 3689-3699.",
    "[28] C. Lendering, E. Akdag, and E. Bondarau, 'SubspaceAD: Training-free few-shot anomaly detection via subspace modeling,' in Proc. IEEE/CVF Conf. Comput. Vis. Pattern Recognit. (CVPR), 2026, pp. 28557-28566.",
    "[29] Y. Li, L. Tian, Y. Dai, W. Chen, L. Bao, and X. Liu, 'FastRef: Fast prototype refinement for few-shot industrial anomaly detection,' in Proc. IEEE/CVF Conf. Comput. Vis. Pattern Recognit. (CVPR), 2026, pp. 43040-43049.",
    "[30] R. Ma, C. Li, J. Chen, Y. Feng, and J. Xie, 'PAPL: Particle-based adaptive prompt learning for zero-shot industrial anomaly detection,' Pattern Recognit., vol. 178, Art. no. 113489, 2026, doi: 10.1016/j.patcog.2026.113489.",
    "[31] J. Zhou, W. Wong, and F. Liao, 'One-shot unsupervised industrial anomaly detection: Enhanced performance under extreme data scarcity,' Pattern Recognit., vol. 173, Art. no. 112759, 2026, doi: 10.1016/j.patcog.2025.112759.",
    "[32] L. Jiang, Y. Huang, Z. Xu, Y. Xu, H.-S. Wong, and S. Wu, 'Defect cue-preserved structural feature refinement for few-shot anomaly detection,' in Proc. IEEE/CVF Conf. Comput. Vis. Pattern Recognit. (CVPR), 2026, pp. 35607-35616.",
    "[33] J. Johnson, M. Douze, and H. Jegou, 'Billion-scale similarity search with GPUs,' IEEE Trans. Big Data, vol. 7, no. 3, pp. 535-547, 2021, doi: 10.1109/TBDATA.2019.2921572.",
]


# IEEE-style reference numbers follow first appearance in the manuscript.
# The source list above preserves the working-bibliography order; this map
# deterministically converts it to citation order at document-generation time.
CITATION_RENUMBER = {
    1: 1, 2: 2, 3: 3, 4: 4, 6: 5, 9: 6, 19: 7, 7: 8, 8: 9,
    10: 10, 11: 11, 27: 12, 17: 13, 12: 14, 15: 15, 16: 16,
    5: 17, 24: 18, 18: 19, 31: 20, 29: 21, 20: 22, 28: 23,
    32: 24, 13: 25, 22: 26, 23: 27, 21: 28, 25: 29, 30: 30,
    26: 31, 14: 32, 33: 33,
}


def renumber_citations(text):
    return re.sub(
        r"\[(\d+)\]",
        lambda m: f"[{CITATION_RENUMBER[int(m.group(1))]}]",
        text,
    )


def set_run_font(run, name="Times New Roman", size=11, bold=None, italic=None, color="000000"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(paragraph, fill="F4F6F9"):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def set_paragraph_border(paragraph, color="A7B4C2", size="8", space="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:space"), space)
        el.set(qn("w:color"), color)
        p_bdr.append(el)
    p_pr.append(p_bdr)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    set_run_font(run, size=9, color="666666")


def add_body(doc, text, first_line=True, after=6):
    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if first_line:
        p.paragraph_format.first_line_indent = Inches(0.25)
    run = p.add_run(renumber_citations(text))
    set_run_font(run)
    return p


def math_run(text):
    run = OxmlElement("m:r")
    value = OxmlElement("m:t")
    value.text = text
    run.append(value)
    return run


def math_script(base, sub=None, sup=None):
    if sub is not None and sup is not None:
        script = OxmlElement("m:sSubSup")
    elif sub is not None:
        script = OxmlElement("m:sSub")
    else:
        script = OxmlElement("m:sSup")
    base_el = OxmlElement("m:e")
    if isinstance(base, (list, tuple)):
        for node in base:
            base_el.append(node)
    else:
        base_el.append(math_run(base))
    script.append(base_el)
    if sub is not None:
        sub_el = OxmlElement("m:sub")
        sub_el.append(math_run(sub))
        script.append(sub_el)
    if sup is not None:
        sup_el = OxmlElement("m:sup")
        sup_el.append(math_run(sup))
        script.append(sup_el)
    return script


def math_fraction(numerator, denominator):
    fraction = OxmlElement("m:f")
    num = OxmlElement("m:num")
    num.append(math_run(numerator))
    den = OxmlElement("m:den")
    den.append(math_run(denominator))
    fraction.append(num)
    fraction.append(den)
    return fraction


def add_equation(doc, equation_kind):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(7)
    math_para = OxmlElement("m:oMathPara")
    math = OxmlElement("m:oMath")

    if equation_kind == "normalization":
        nodes = [
            math_script("g", "D,p"), math_run(" = "), math_script("f", "D,p"),
            math_run(" / max(‖"), math_script("f", "D,p"), math_script("‖", "2"),
            math_run(", ε),    "), math_script("g", "C,p"), math_run(" = "),
            math_script("f", "C→D,p"), math_run(" / max(‖"),
            math_script("f", "C→D,p"), math_script("‖", "2"), math_run(", ε)."),
        ]
    elif equation_kind == "fusion":
        nodes = [
            math_script("f", "A,p"), math_run(" = Norm([α "),
            math_script("g", "D,p"), math_run(" ; (1 − α) "),
            math_script("g", "C,p"), math_run("]),    α = 0.5,    dim("),
            math_script("f", "A,p"), math_run(") = 1536."),
        ]
    elif equation_kind == "memory":
        nodes = [
            math_script("M", "c"), math_run(" = "),
            math_script("⋃", "i=1", "K"), math_run(" { "),
            math_script("f", "A,p"), math_run("("),
            math_script("x", "i", "c"), math_run(") | p ∈ "),
            math_script("G", "D"), math_run(" }."),
        ]
    elif equation_kind == "score":
        norm_expression = [
            math_run("‖"), math_script("f", "A,p"),
            math_run("(x) − m‖"),
        ]
        nodes = [
            math_script("a", "p"), math_run("(x) = "),
            math_fraction("1", "2"), math_run(" "),
            math_script("min", "m∈M_c"), math_run(" "),
            math_script(norm_expression, "2", "2"), math_run("."),
        ]
    else:
        raise ValueError(f"Unknown equation kind: {equation_kind}")

    for node in nodes:
        math.append(node)
    math_para.append(math)
    p._p.append(math_para)
    return p


def add_heading(doc, text, level):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_run_font(r, size={1: 15, 2: 12.5, 3: 11.5}[level], bold=True)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.208
    r = p.add_run(renumber_citations(text))
    set_run_font(r)
    return p


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(11)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for level, size, before, after in [(1, 15, 16, 8), (2, 12.5, 12, 6), (3, 11.5, 8, 4)]:
        style = styles[f"Heading {level}"]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_name in ("List Bullet", "List Number"):
        style = styles[list_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208

    if "Figure Caption" not in styles:
        cap = styles.add_style("Figure Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        cap = styles["Figure Caption"]
    cap.font.name = "Times New Roman"
    cap._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    cap._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    cap.font.size = Pt(9)
    cap.font.italic = False
    cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    cap.paragraph_format.space_before = Pt(4)
    cap.paragraph_format.space_after = Pt(8)


def build():
    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    hp = section.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hr = hp.add_run("DCFnet | Working manuscript draft through Method")
    set_run_font(hr, size=9, color="666666")
    add_page_number(section.footer.paragraphs[0])

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(TITLE)
    set_run_font(r, size=18, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("Author information withheld in this working draft")
    set_run_font(r, size=10, italic=True, color="666666")

    add_heading(doc, "Abstract", 1)
    add_body(doc, ABSTRACT, first_line=False, after=5)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.2
    k = p.add_run("Keywords: ")
    set_run_font(k, size=10.5, bold=True)
    r = p.add_run("industrial anomaly detection; few-shot learning; anomaly localization; visual foundation models; feature fusion; normal memory bank")
    set_run_font(r, size=10.5)

    add_heading(doc, "1. Introduction", 1)
    for para in INTRO_PARAS:
        add_body(doc, para)
    lead = add_body(doc, "The contributions of this work are summarized as follows:", first_line=False, after=4)
    lead.paragraph_format.keep_with_next = True
    for item in CONTRIBUTIONS:
        add_bullet(doc, item)

    add_heading(doc, "2. Related Work", 1)
    for para in RELATED:
        add_body(doc, para)

    add_heading(doc, "2.1 Reconstruction and One-Class Learning", 2)
    for para in RW_RECON:
        add_body(doc, para)

    add_heading(doc, "2.2 Feature Embedding and Normal-Memory Methods", 2)
    for para in RW_MEMORY:
        add_body(doc, para)

    add_heading(doc, "2.3 Foundation Models for Few-Shot Anomaly Detection", 2)
    for para in RW_FOUNDATION:
        add_body(doc, para)

    add_heading(doc, "2.4 Hybrid Fusion and the Position of This Work", 2)
    for para in RW_FUSION:
        add_body(doc, para)

    add_heading(doc, "3. Detailed DCFnet", 1)
    add_heading(doc, "3.1 Problem Statement", 2)
    for para in METHOD_PROBLEM:
        add_body(doc, para)

    add_heading(doc, "3.2 Overview Structure", 2)
    for para in METHOD_OVERVIEW:
        add_body(doc, para)

    if not FIGURE_1.is_file():
        raise FileNotFoundError(f"Figure 1 not found: {FIGURE_1}")
    figure_p = doc.add_paragraph()
    figure_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    figure_p.paragraph_format.space_before = Pt(8)
    figure_p.paragraph_format.space_after = Pt(4)
    figure_p.paragraph_format.keep_with_next = True
    figure_run = figure_p.add_run()
    figure_shape = figure_run.add_picture(str(FIGURE_1), width=Inches(6.45))
    figure_shape._inline.docPr.set("title", "DCFnet architecture")
    figure_shape._inline.docPr.set(
        "descr",
        "Two-lane workflow of DCFnet. Normal support images construct a fixed category memory, "
        "while a query image is processed separately by frozen DINOv2 and DPAM-modified CLIP "
        "visual encoders before exact nearest-neighbor anomaly scoring.",
    )

    cap = doc.add_paragraph(style="Figure Caption")
    cap.paragraph_format.keep_together = True
    cap.paragraph_format.keep_with_next = True
    cr = cap.add_run(
        "Fig. 1. Overall workflow of DCFnet. Normal support images construct a fixed category memory, whereas the query image follows a separate inference path. Both paths share frozen dual-encoder extraction, CLIP-to-DINO spatial alignment, branch-wise magnitude normalization, fixed equal-scaling concatenation, and global normalization. Query patches are compared with the fixed memory before bilinear upsampling and Gaussian smoothing produce the final anomaly map."
    )
    set_run_font(cr, size=9)

    add_heading(doc, "3.3 Frozen Dual-Encoder Visual Representation", 2)
    for para in METHOD_ENCODERS:
        add_body(doc, para)

    add_heading(doc, "3.4 Spatial Calibration and Fixed Feature Fusion", 2)
    for para in METHOD_FUSION:
        add_body(doc, para)
    add_equation(doc, NORMALIZATION_EQUATION)
    add_equation(doc, FUSION_EQUATION)
    for para in METHOD_FUSION_AFTER:
        add_body(doc, para)

    add_heading(doc, "3.5 Normal Reference Memory and Anomaly Scoring", 2)
    for para in METHOD_MEMORY:
        add_body(doc, para)
    add_equation(doc, MEMORY_EQUATION)
    for para in METHOD_MEMORY_2:
        add_body(doc, para)
    add_equation(doc, SCORE_EQUATION)
    for para in METHOD_MEMORY_3:
        add_body(doc, para)

    add_heading(doc, "3.6 Matched Control, Complexity, and Protocol Integrity", 2)
    for para in METHOD_CONTROL:
        add_body(doc, para)

    add_heading(doc, "References", 1)
    ordered_references = sorted(
        (renumber_citations(ref) for ref in REFERENCES),
        key=lambda value: int(re.match(r"\[(\d+)\]", value).group(1)),
    )
    for ref in ordered_references:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.05
        r = p.add_run(ref)
        set_run_font(r, size=8.5)

    props = doc.core_properties
    props.title = TITLE
    props.subject = "Working English manuscript draft through the Method section"
    props.author = ""
    props.keywords = "industrial anomaly detection; few-shot; DINOv2; AnomalyCLIP; DCFnet"
    props.comments = "Prepared from the frozen A1 method specification and verified literature sources."

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
